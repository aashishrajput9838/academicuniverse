import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("=== PARAGRAPHS 0 TO 15 IN CURRENT PaperV5_Ollama_Primary.docx ===")
for i in range(min(15, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f"P{i} [{p.style.name}] (align={p.alignment}): '{p.text[:100]}'")
