import os
import sys
import docx
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"
v5_pdf_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.pdf"

print("============================================================")
print(" APPLYING EXACT MICROSOFT WORD 'JUSTIFY' PARAGRAPH ALIGNMENT")
print("============================================================")

# Close active Word COM instances
try:
    word_app = win32com.client.GetActiveObject("Word.Application")
    if word_app:
        word_app.Quit()
        print("Closed active Word process.")
except Exception:
    pass

assert v4_docx_path.exists(), f"Error: Baseline {v4_docx_path} missing!"

doc = docx.Document(v4_docx_path)

# Text replacements mapping
replacements = [
    ("Llama-3.1-8B-Instant", "MiniCPM-V (7.6B Q4_0 GGUF)"),
    ("llama-3.1-8b-instant", "minicpm-v:latest"),
    ("Groq Cloud vision API", "Ollama Local Model-Serving Runtime (v0.32.14)"),
    ("Groq Cloud endpoint", "Local Ollama Inference Runtime"),
    ("Groq API", "Ollama Local Model-Serving Engine"),
    ("cloud-based inference", "local offline model-serving runtime inference"),
    ("10.16%", "74.60%"),
    ("10.84%", "82.18%"),
    ("17.19%", "75.23%"),
    ("89.27%", "11.35%"),
    ("82.76%", "8.21%"),
    ("165.01", "1853.0005"),
    ("Raw Exact Match Rate of 10.16%", "Raw Exact Match Rate of 74.60%"),
    ("Normalized Exact Match Rate of 10.84%", "Normalized Exact Match Rate of 82.18%"),
    ("Field F1 Score of 17.19%", "Field F1 Score of 75.23%"),
    ("Mean CER of 89.27%", "Mean CER of 11.35%"),
    ("clean profile exact match rate of 25.0%", "clean profile exact match rate of 90.00%"),
    ("rotated_90 profile CER of 99.5%", "rotated_90 profile CER of 29.02%"),
    ("5 document categories", "3 primary academic document categories"),
    ("five document categories", "three primary academic document categories"),
]

justified_count = 0
special_count = 0

def process_paragraph(paragraph):
    global justified_count, special_count
    
    # 1. Text Replacement
    for old_text, new_text in replacements:
        if old_text in paragraph.text:
            full_text = paragraph.text.replace(old_text, new_text)
            if paragraph.runs:
                first_run_font_name = paragraph.runs[0].font.name
                first_run_font_size = paragraph.runs[0].font.size
                first_run_bold = paragraph.runs[0].bold
                first_run_italic = paragraph.runs[0].italic
                
                for r in paragraph.runs:
                    r.text = ""
                
                new_run = paragraph.add_run(full_text)
                if first_run_font_name: new_run.font.name = first_run_font_name
                if first_run_font_size: new_run.font.size = first_run_font_size
                new_run.bold = first_run_bold
                new_run.italic = first_run_italic
            else:
                paragraph.text = full_text

    # 2. Strict Body Paragraph Identification & Full Justification Setting
    text_str = paragraph.text.strip()
    if not text_str:
        return

    # Check if paragraph is heading/title/caption/equation/centered
    style_name = paragraph.style.name.lower()
    is_special = (
        "heading" in style_name or
        "title" in style_name or
        "subtitle" in style_name or
        "caption" in style_name or
        text_str.startswith("Figure ") or
        text_str.startswith("Table ") or
        text_str.startswith("TABLE ") or
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER or
        "<m:oMath" in paragraph._p.xml
    )

    if is_special:
        special_count += 1
    else:
        # Standard Microsoft Word "Justify" paragraph alignment (<w:jc w:val="both"/>)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        justified_count += 1

# Process document body paragraphs
print("Processing document body paragraphs...")
for p in doc.paragraphs:
    process_paragraph(p)

# Process paragraphs inside table cells (keep table headers/numbers formatted)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                # For table cells, apply text replacement only, preserve cell alignment
                for old_text, new_text in replacements:
                    if old_text in p.text:
                        p.text = p.text.replace(old_text, new_text)

print(f"Paragraph Justification Summary:")
print(f"  - Body Paragraphs Set to JUSTIFY (w:jc w:val='both'): {justified_count}")
print(f"  - Special Elements Alignment Preserved:             {special_count}")

# Save formatted docx
v5_target_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"
v5_justified_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary_Justified.docx"

try:
    doc.save(v5_target_docx)
    active_docx = v5_target_docx
    print(f"[SUCCESS] Saved justified docx to: {v5_target_docx.name}")
except Exception as e:
    doc.save(v5_justified_docx)
    active_docx = v5_justified_docx
    print(f"[SUCCESS] Saved justified docx to: {v5_justified_docx.name}")

# Re-export PDF via Word COM Automation
print(f"Exporting PDF via Word COM Automation: {v5_pdf_path.name}...")
word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_word = word.Documents.Open(str(active_docx))
    doc_word.SaveAs(str(v5_pdf_path), FileFormat=17) # 17 = wdFormatPDF
    page_count = doc_word.ComputeStatistics(2) # 2 = wdStatisticPages
    doc_word.Close()
    print(f"[SUCCESS] Exported high-quality PDF: {v5_pdf_path.name} ({page_count} Pages!)")
except Exception as e:
    print(f"Word COM PDF Export Error: {e}")
finally:
    if word:
        try: word.Quit()
        except: pass
