import os
import docx
import win32com.client
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_TOC.docx"
test_pdf_path = workspace / "docs" / "paper" / "PaperV5_Test_TOC.pdf"

print("=== TESTING TOC GENERATION PIPELINE ===")

doc = docx.Document(v4_docx_path)

# Locate insertion point after Index Terms (P5)
target_idx = None
for i, p in enumerate(doc.paragraphs[:15]):
    if p.text.strip().startswith("Index Terms"):
        target_idx = i
        break

assert target_idx is not None, "Error: Index Terms paragraph not found!"

target_p = doc.paragraphs[target_idx]

# Insert CONTENTS heading
p_contents_hdr = target_p.insert_paragraph_before("CONTENTS")
p_contents_hdr.paragraph_format.space_before = Pt(14)
p_contents_hdr.paragraph_format.space_after = Pt(6)
p_contents_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
if p_contents_hdr.runs:
    r = p_contents_hdr.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.bold = True

# Insert TOC Field paragraph
p_toc = target_p.insert_paragraph_before()
p_toc.paragraph_format.space_after = Pt(12)
pPr = p_toc._p.get_or_add_pPr()
pStyle = docx.oxml.OxmlElement('w:pStyle')
pStyle.set(docx.oxml.ns.qn('w:val'), 'TOCHeading')
pPr.append(pStyle)

r_toc = p_toc.add_run()
fldChar1 = docx.oxml.OxmlElement('w:fldChar')
fldChar1.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
instrText = docx.oxml.OxmlElement('w:instrText')
instrText.set(docx.oxml.ns.qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = docx.oxml.OxmlElement('w:fldChar')
fldChar2.set(docx.oxml.ns.qn('w:fldCharType'), 'separate')
fldChar3 = docx.oxml.OxmlElement('w:fldChar')
fldChar3.set(docx.oxml.ns.qn('w:fldCharType'), 'end')

r_toc._r.append(fldChar1)
r_toc._r.append(instrText)
r_toc._r.append(fldChar2)
r_toc._r.append(fldChar3)

doc.save(test_docx_path)
print(f"[SUCCESS] Saved test docx with TOC field to {test_docx_path.name}")

# Test Word COM TOC update & PDF export
abs_docx = os.path.abspath(str(test_docx_path))
abs_pdf = os.path.abspath(str(test_pdf_path))

word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_word = word.Documents.Open(abs_docx)
    
    # Update all fields in document (including TOC)
    doc_word.Fields.Update()
    for s in doc_word.Sections:
        s.Range.Fields.Update()
    
    doc_word.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
    page_count = doc_word.ComputeStatistics(2)
    doc_word.Close()
    print(f"[SUCCESS] Exported PDF with updated TOC: {test_pdf_path.name} ({page_count} Pages!)")
except Exception as e:
    print(f"Word COM TOC Export Note: {e}")
finally:
    if word:
        try: word.Quit()
        except: pass
