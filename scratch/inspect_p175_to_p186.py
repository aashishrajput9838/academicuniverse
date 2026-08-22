import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

doc = docx.Document(v5_docx_path)

print("=== INSPECTING PARAGRAPHS P175 TO P186 ===")

for i in range(175, min(188, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f"P{i:3d}: '{p.text.strip()}'")
