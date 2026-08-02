import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def clean_math_text(text):
    """
    Converts LaTeX math expressions into clean, publication-quality mathematical notation.
    Replaces all backslash commands, superscripts, subscripts, and LaTeX operators.
    """
    if not text:
        return text

    # Handle display math brackets if present
    text = text.replace('$$', '')
    
    # Superscripts and Subscripts mappings
    sub_map = {'0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄', '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉', 
               'v': 'ᵥ', 'k': 'ₖ', 't': 'ₜ', '0': '₀', '1': '₁', '2': '₂'}
    sup_map = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴', '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹', 
               '-': '⁻', '+': '⁺', '1': '¹'}

    # Specific common LaTeX replacements
    replacements = [
        (r'\\mu_v', 'μᵥ'),
        (r'\\mu', 'μ'),
        (r'\\mathcal\{H\}', 'H'),
        (r'\\lambda', 'λ'),
        (r'\\epsilon', 'ε'),
        (r'\\Delta', 'Δ'),
        (r'\\Sigma', 'Σ'),
        (r'\\sum_\{k=1\}\^\{K\}', 'Σ(k=1..K)'),
        (r'\\sum', 'Σ'),
        (r'\\xrightarrow\{\\text\{depends on\}\}', ' → '),
        (r'\\xrightarrow', ' → '),
        (r'\\in', '∈'),
        (r'\\ge', '≥'),
        (r'\\le', '≤'),
        (r'\\cdot', '·'),
        (r'\\bar\{S\}_k', 'S̄ₖ'),
        (r'\\bar\{C\}_k', 'C̄ₖ'),
        (r'\\bar\{S\}', 'S̄'),
        (r'\\bar\{C\}', 'C̄'),
        (r'\\min', 'min'),
        (r'\\max', 'max'),
        (r'\\sqrt\{N\}', '√N'),
        (r'\\frac\{S\(t_2\) - S\(t_1\)\}\{\\Delta t\}', '(S(t₂) − S(t₁)) / Δt'),
        (r'\\frac\{\\sum_\{k=1\}\^\{K\} \\left\( \\bar\{S\}_k \\cdot \\bar\{C\}_k \\cdot w_k \\right\)\}\{\\sum_\{k=1\}\^\{K\} \\left\( \\bar\{C\}_k \\cdot w_k \\right\) \+ \\epsilon\}', 'Σ(S̄ₖ · C̄ₖ · wₖ) / [ Σ(C̄ₖ · wₖ) + ε ]'),
        (r'\\min\\left\(99, \\frac\{V_\{total\} \\cdot \(0\.6 \+ 0\.4 \\cdot O_\{ratio\}\) \\cdot D_\{avg\} \\cdot K_\{complexity\}\}\{\\max\(1, \\sqrt\{N\}\)\} \+ 3N\\right\)', 'min(99, [V_total · (0.6 + 0.4 · O_ratio) · D_avg · K_complexity] / max(1, √N) + 3N)'),
        (r'W_\{\\text\{AU\\_DIC\}\}', 'W_AU_DIC'),
        (r'W_\{\\text\{AU\_DIC\}\}', 'W_AU_DIC'),
        (r'W_\{\\text\{source\}\}', 'W_source'),
        (r'V_\{\\text\{total\}\}', 'V_total'),
        (r'O_\{\\text\{ratio\}\}', 'O_ratio'),
        (r'D_\{\\text\{avg\}\}', 'D_avg'),
        (r'K_\{\\text\{complexity\}\}', 'K_complexity'),
        (r'S_\{\\text\{decayed\}\}', 'S_decayed'),
        (r't_\{\\text\{last\}\}', 't_last'),
        (r'\\text\{month\}\^\{-1\}', 'month⁻¹'),
        (r'\\text\{pts/mo\}', 'pts/mo'),
        (r'\\text\{months\}', 'months'),
        (r'\\text\{Slow\}', 'Slow'),
        (r'\\text\{Moderate\}', 'Moderate'),
        (r'\\text\{Baseline\}', 'Baseline'),
        (r'\\text\{Fast\}', 'Fast'),
        (r'\\text\{Aggressive\}', 'Aggressive'),
        (r'\\text\{Eq\. (\d+)\}', r'[Eq. \1]'),
        (r'\\left\(', '('),
        (r'\\right\)', ')'),
        (r'\\left\[', '['),
        (r'\\right\]', ']'),
        (r'S_0', 'S₀'),
        (r'S\(t_1\)', 'S(t₁)'),
        (r'S\(t_2\)', 'S(t₂)'),
        (r't_1', 't₁'),
        (r't_2', 't₂'),
        (r't_\{1/2\}', 't₁⁄₂'),
        (r't_\{last\}', 't_last'),
        (r'e\^\{-\\lambda \(t - t_\{last\}\)\}', 'e^(-λ(t − t_last))'),
        (r'10\^\{-6\}', '10⁻⁶'),
        (r'\\quad', ' '),
    ]

    for pat, repl in replacements:
        text = re.sub(pat, repl, text)

    # General regex cleanups for remaining LaTeX tags
    text = text.replace(r'\mu_v', 'μᵥ')
    text = text.replace(r'\mu', 'μ')
    text = text.replace(r'\lambda', 'λ')
    text = text.replace(r'\mathcal{H}', 'H')
    text = text.replace(r'\mathcal', '')
    text = text.replace(r'\epsilon', 'ε')
    text = text.replace(r'\Delta', 'Δ')
    text = text.replace(r'\frac', '')
    text = text.replace(r'\sum', 'Σ')
    text = text.replace(r'\xrightarrow', ' → ')
    text = text.replace(r'\cdot', '·')
    text = text.replace(r'\bar', '')
    text = re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    text = re.sub(r'\_\{([^}]+)\}', r'_\1', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)  # Remove inline $ signs
    text = text.replace('\\_', '_')
    text = text.replace('\\', '')
    
    return text

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_corrected_docx():
    md_path = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\902d6e6c-b35a-4d3e-a983-59d2e321b9fc\paper2_academic_universe_growth_intelligence.md"
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final_Corrected.docx"
    fig1_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure1_Architecture_Pipeline.png"
    fig2_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure2_Decay_Sensitivity_Curves.png"

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    display_eq_count = 0
    inline_eq_count = 0
    table_eq_count = 0

    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def process_table(t_lines):
        nonlocal table_eq_count
        rows = []
        for line in t_lines:
            if '|' in line:
                parts = [p.strip() for p in line.split('|')[1:-1]]
                if parts and not all(set(p) <= {'-', ':', ' '} for p in parts):
                    rows.append(parts)
        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                if c_idx < len(table.columns):
                    cell = table.cell(r_idx, c_idx)
                    cleaned_val = clean_math_text(val)
                    if '\\' in val or '$' in val:
                        table_eq_count += 1
                    cell.text = cleaned_val
                    set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

                    if r_idx == 0:
                        set_cell_background(cell, "1E293B")
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    else:
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "F8FAFC")
                        for run in p.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()

    while i < len(lines):
        line = lines[i].rstrip('\r\n')

        # Code Blocks -> Replace ASCII diagrams with High-Res PNG Figures
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                block_content = '\n'.join(code_block_lines)
                if 'GitHub API' in block_content or 'Evidence Intelligence Layer' in block_content:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(10)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.add_run().add_picture(fig1_path, width=Inches(6.2))
                elif 'Proficiency S(t)' in block_content or 'λ=0.01' in block_content:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(10)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.add_run().add_picture(fig2_path, width=Inches(5.8))
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    run = p.add_run(clean_math_text(block_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and not line.startswith('>'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            process_table(table_lines)
            table_lines = []

        # Blockquotes
        if line.startswith('> '):
            quote_text = clean_math_text(line[2:].strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(clean_math_text(line[2:]))
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_math_text(line[3:]))
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_math_text(line[4:]))
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(clean_math_text(line[5:]))
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif line.startswith('$$') and line.endswith('$$'):
            # Centered Display Math Object (Cambria Math)
            display_eq_count += 1
            eq_text = clean_math_text(line[2:-2].strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(eq_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif line.strip() != '':
            if '$' in line:
                inline_eq_count += line.count('$') // 2
            
            cleaned_line = clean_math_text(line)
            p = doc.add_paragraph()

            # Captions
            if cleaned_line.startswith('Figure') or cleaned_line.startswith('Table') or cleaned_line.startswith('*Figure') or cleaned_line.startswith('*Table'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(10)

            # Format paragraph content
            r = p.add_run(cleaned_line)
            if cleaned_line.startswith('*Figure') or cleaned_line.startswith('*Table'):
                r.italic = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        i += 1

    if in_table:
        process_table(table_lines)

    # Reference Section
    ref_heading = False
    for paragraph in doc.paragraphs:
        if '## References' in paragraph.text or 'References' in paragraph.text:
            ref_heading = True
            break
            
    if ref_heading:
        enhanced_refs = [
            "5. Corbett, A. T., & Anderson, J. R. (1994). Knowledge tracing: Modeling the acquisition of procedural knowledge. User Modeling and User-Adapted Interaction, 4(4), 253-278.",
            "6. Piech, C., Bassen, J., Huang, J., Ganguli, S., Sahami, M., Guibas, L. J., & Sohl-Dickstein, J. (2015). Deep knowledge tracing. Advances in Neural Information Processing Systems (NeurIPS), 28, 505-513.",
            "7. Bodily, R., & Verbert, K. (2017). Review of research on student-facing learning analytics dashboards. Journal of Learning Analytics, 4(3), 67-89.",
            "8. Ferguson, R. (2012). Learning analytics: Drivers, developments and challenges. International Journal of Technology Enhanced Learning, 4(5-6), 304-317.",
            "9. Jivet, I., Scheffel, M., Specht, M., & Drachsler, H. (2018). License to evaluate: Preparing learning analytics dashboards for educational practice. Proceedings of the 8th International Conference on Learning Analytics & Knowledge (LAK '18), 31-40."
        ]
        doc.add_paragraph()
        p_note = doc.add_paragraph()
        r_note = p_note.add_run("Foundational Literature Additions (IEEE / ACM Learning Analytics Corpus):")
        r_note.bold = True
        r_note.font.size = Pt(11)
        r_note.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        for ref in enhanced_refs:
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.left_indent = Inches(0.25)
            p_ref.paragraph_format.space_after = Pt(4)
            r = p_ref.add_run(ref)
            r.font.size = Pt(10)

    doc.save(docx_path)
    print(f"Successfully compiled corrected DOCX to {docx_path}")
    print(f"Display Equations Converted: {display_eq_count}")
    print(f"Inline Equations Converted: {inline_eq_count}")
    print(f"Table Cells Math Converted: {table_eq_count}")

if __name__ == '__main__':
    build_corrected_docx()
