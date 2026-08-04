import docx

def audit_omml_docx():
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final_OMML.docx"
    doc = docx.Document(docx_path)

    corrupted_glyphs = 0
    raw_latex = 0
    raw_markdown = 0
    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)

    latex_patterns = [r'\frac', r'\sum', r'\lambda', r'\mu', r'\mathcal', r'$', r'^{', r'_{']

    for p in doc.paragraphs:
        txt = p.text
        if '\ufffd' in txt or '□' in txt or '??' in txt:
            corrupted_glyphs += 1
        if '**' in txt or '`' in txt:
            raw_markdown += 1
        for lp in latex_patterns:
            if lp in txt:
                raw_latex += 1

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                txt = c.text
                if '\ufffd' in txt or '□' in txt or '??' in txt:
                    corrupted_glyphs += 1
                if '**' in txt or '`' in txt:
                    raw_markdown += 1
                for lp in latex_patterns:
                    if lp in txt:
                        raw_latex += 1

    print("===========================================================")
    print("          IEEE OMML MATH RENDERING AUDIT REPORT            ")
    print("===========================================================")
    print(f"Total Paragraphs Audited:         {total_paragraphs}")
    print(f"Total Tables Audited:             {total_tables}")
    print(f"Corrupted Glyphs Remaining:       {corrupted_glyphs}")
    print(f"Raw LaTeX Commands Remaining:     {raw_latex}")
    print(f"Raw Markdown Markers Remaining:   {raw_markdown}")
    print("-----------------------------------------------------------")
    print(f"Rendering Validation Result:      {'PASSED' if corrupted_glyphs == 0 else 'FAILED'}")
    print(f"IEEE Production Quality:          {'PASSED' if raw_latex == 0 and raw_markdown == 0 else 'FAILED'}")
    print("===========================================================")

if __name__ == '__main__':
    audit_omml_docx()
