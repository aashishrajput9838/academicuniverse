import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" VERIFYING SECTION 4.3 SUBSECTIONS IN PaperV5_Ollama_Primary.docx")
print("============================================================")

subsections_43 = [
    "4.3.1 Category Classification Accuracy",
    "4.3.2 Field Extraction Precision, Recall, and F1-Score",
    "4.3.3 Character Error Rate (CER)",
    "4.3.4 Word Error Rate (WER)",
    "4.3.5 Joint Record Exact Match Rate (EM)",
    "4.3.6 Execution Latency & Throughput"
]

found_subsections = {}
for sub in subsections_43:
    found_subsections[sub] = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(sub):
            found_subsections[sub] = True
            print(f"Found '{sub}' at P{i}")
            break

# Check malformed text absence
malformed_found = False
for p in doc.paragraphs:
    if "ms/sample samples/sec. Latency & Throughput" in p.text:
        malformed_found = True
        break

print(f"\n1. All 6 Subsections (4.3.1–4.3.6) Present: {all(found_subsections.values())}")
print(f"2. Malformed Text Absent:                   {not malformed_found}")

# TOC check
toc_matches = sum(1 for p in doc.paragraphs[:50] if "4.3." in p.text)
print(f"3. Subsections (4.3.1–4.3.6) in TOC:        {toc_matches} entries")

print("============================================================")
assert all(found_subsections.values()), "Some 4.3 subsections missing!"
assert not malformed_found, "Malformed text still present!"
assert toc_matches >= 6, "TOC missing 4.3 subsections!"

print("EMPIRICAL VERIFICATION SUCCESSFUL: Section 4.3 structured cleanly into 4.3.1–4.3.6!")
