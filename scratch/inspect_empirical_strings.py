import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

doc = docx.Document(v5_docx_path)

print("=== EMPIRICAL STRINGS CHECK IN DOCX ===")
full_text = " ".join([p.text for p in doc.paragraphs])

print(f"'Ollama' in doc:    {'Ollama' in full_text}")
print(f"'MiniCPM' in doc:   {'MiniCPM' in full_text}")
print(f"'75.23%' in doc:    {'75.23%' in full_text}")
print(f"'82.18%' in doc:    {'82.18%' in full_text}")

for i, p in enumerate(doc.paragraphs):
    if "75.23%" in p.text or "Ollama" in p.text or "MiniCPM" in p.text:
        print(f"P{i}: '{p.text[:100]}'")
