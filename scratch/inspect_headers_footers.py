import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

def inspect_doc_hf(doc, label):
    print(f"=== INSPECTING HEADERS & FOOTERS IN {label} ===")
    for s_idx, section in enumerate(doc.sections):
        print(f"Section {s_idx}:")
        header = section.header
        print(f"  Header is_linked_to_previous: {header.is_linked_to_previous}")
        for p in header.paragraphs:
            print(f"    Header P: repr={repr(p.text)}")
        
        first_page_header = section.first_page_header
        for p in first_page_header.paragraphs:
            print(f"    First Page Header P: repr={repr(p.text)}")
            
        footer = section.footer
        print(f"  Footer is_linked_to_previous: {footer.is_linked_to_previous}")
        for p in footer.paragraphs:
            print(f"    Footer P: repr={repr(p.text)}")

doc4 = docx.Document(v4_docx_path)
inspect_doc_hf(doc4, "PaperV4_Final_Submission.docx")

doc5 = docx.Document(v5_docx_path)
inspect_doc_hf(doc5, "PaperV5_Ollama_Primary.docx")
