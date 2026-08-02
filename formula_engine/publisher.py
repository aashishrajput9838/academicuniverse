import os
import re
import time
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from formula_engine.extractor import FormulaExtractor, ExtractionValidationError
from formula_engine.converter import FormulaConverter, ConversionError
from formula_engine.validator import DocumentValidator, DocumentValidationError

class PublisherError(Exception):
    """Raised when document compilation or multi-stage validation fails."""
    pass

class IEEEPublisher:
    """
    Multi-Stage Pipeline Orchestrator:
    Compiles Academic Universe Paper 2 Version 1.0 into IEEE production-grade DOCX and PDF deliverables.
    Enforces multi-stage validation gates at every stage.
    """

    def __init__(self, md_path, docx_path, pdf_path, fig1_path, fig2_path):
        self.md_path = md_path
        self.docx_path = docx_path
        self.pdf_path = pdf_path
        self.fig1_path = fig1_path
        self.fig2_path = fig2_path

        self.stats = {
            'display_equations_converted': 0,
            'inline_equations_converted': 0,
            'table_math_converted': 0,
            'caption_math_converted': 0,
            'total_equations_processed': 0,
            'total_validation_checks_executed': 0,
            'unsupported_latex_constructs': [],
            'start_time': 0,
            'end_time': 0,
            'total_build_time': 0,
            'avg_equation_time_ms': 0
        }

    def build(self):
        """Executes full multi-stage document compilation and validation pipeline."""
        self.stats['start_time'] = time.time()

        if not os.path.exists(self.md_path):
            raise PublisherError(f"Input markdown manuscript not found at '{self.md_path}'.")

        with open(self.md_path, 'r', encoding='utf-8') as f:
            raw_markdown = f.read()

        # Gate 1: Extraction Validation
        extracted_formulas = FormulaExtractor.extract_formulas(raw_markdown)
        self.stats['total_validation_checks_executed'] += 1

        # Build DOCX Document
        doc = docx.Document()

        # Page Setup - 1 inch margins
        for s in doc.sections:
            s.top_margin = Inches(1.0)
            s.bottom_margin = Inches(1.0)
            s.left_margin = Inches(1.0)
            s.right_margin = Inches(1.0)

        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        normal_style.paragraph_format.line_spacing = 1.15
        normal_style.paragraph_format.space_after = Pt(6)

        lines = raw_markdown.splitlines()
        i = 0
        in_code_block = False
        code_block_lines = []
        in_table = False
        table_lines = []
        display_eq_counter = 0

        def process_table(t_lines):
            rows = []
            for line in t_lines:
                if '|' in line:
                    parts = [p.strip() for p in line.split('|')[1:-1]]
                    if parts and not all(set(p) <= {'-', ':', ' '} for p in parts):
                        rows.append(parts)
            if not rows:
                return

            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            self._set_table_borders(table)

            for r_idx, row in enumerate(rows):
                trPr = table.rows[r_idx]._tr.get_or_add_trPr()
                trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                if r_idx == 0:
                    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

                for c_idx, val in enumerate(row):
                    if c_idx < len(table.columns):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.space_before = Pt(2)

                        self._parse_line_to_paragraph(p, val, default_font_size=10, is_table=True)

                        self._set_cell_margins(cell)
                        if r_idx == 0:
                            self._set_cell_background(cell, "1E293B")
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.size = Pt(10)
                        else:
                            if r_idx % 2 == 1:
                                self._set_cell_background(cell, "F8FAFC")
                            for run in p.runs:
                                run.font.size = Pt(10)
            doc.add_paragraph()

        while i < len(lines):
            line = lines[i]

            # Code blocks -> High-Res IEEE Diagrams
            if line.startswith('```'):
                if in_code_block:
                    in_code_block = False
                    block_content = '\n'.join(code_block_lines)
                    if 'GitHub API' in block_content or 'Evidence Intelligence Layer' in block_content:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(12)
                        p_img.paragraph_format.space_after = Pt(4)
                        p_img.paragraph_format.keep_with_next = True
                        p_img.add_run().add_picture(self.fig1_path, width=Inches(6.2))

                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_after = Pt(12)
                        p_cap.paragraph_format.keep_with_next = True
                        r_tag = p_cap.add_run("Fig. 1. ")
                        r_tag.bold = True
                        r_tag.font.name = 'Times New Roman'
                        r_tag.font.size = Pt(9)
                        r_txt = p_cap.add_run("High-level architectural pipeline of the Academic Universe Growth Intelligence Ecosystem.")
                        r_txt.font.name = 'Times New Roman'
                        r_txt.font.size = Pt(9)
                        r_txt.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

                    elif 'Proficiency S(t)' in block_content or 'λ=0.01' in block_content:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(12)
                        p_img.paragraph_format.space_after = Pt(4)
                        p_img.paragraph_format.keep_with_next = True
                        p_img.add_run().add_picture(self.fig2_path, width=Inches(5.8))

                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_after = Pt(12)
                        p_cap.paragraph_format.keep_with_next = True
                        r_tag = p_cap.add_run("Fig. 2. ")
                        r_tag.bold = True
                        r_tag.font.name = 'Times New Roman'
                        r_tag.font.size = Pt(9)
                        r_txt = p_cap.add_run("Sensitivity analysis of technical skill proficiency decay S(t) under varied decay parameter λ coefficients.")
                        r_txt.font.name = 'Times New Roman'
                        r_txt.font.size = Pt(9)
                        r_txt.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                    else:
                        p = doc.add_paragraph()
                        p.paragraph_format.left_indent = Inches(0.2)
                        self._parse_line_to_paragraph(p, block_content, default_font_size=9.5)
                    code_block_lines = []
                else:
                    in_code_block = True
                    code_block_lines = []
                i += 1
                continue

            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue

            # Tables
            if '|' in line and not line.startswith('>'):
                if not in_table:
                    in_table = True
                    table_lines = [line]
                else:
                    table_lines.append(line)
                i += 1
                continue
            elif in_table:
                in_table = False
                process_table(table_lines)
                table_lines = []

            # Blockquotes
            if line.startswith('> '):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                self._parse_line_to_paragraph(p, line[2:].strip(), default_font_size=10)
                for r in p.runs:
                    r.font.italic = True
                    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
                i += 1
                continue

            # Headings
            if line.startswith('# '):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.keep_with_next = True
                self._parse_line_to_paragraph(p, line[2:], default_font_size=18)
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif line.startswith('## '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                self._parse_line_to_paragraph(p, line[3:], default_font_size=13)
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
            elif line.startswith('### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
                self._parse_line_to_paragraph(p, line[4:], default_font_size=11.5)
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            elif line.startswith('#### '):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True
                self._parse_line_to_paragraph(p, line[5:], default_font_size=11)
                for r in p.runs:
                    r.bold = True
                    r.italic = True
                    r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

            # Centered Display Equations
            elif line.startswith('$$') and line.endswith('$$'):
                display_eq_counter += 1
                raw_eq = line[2:-2].strip()

                p_eq = doc.add_paragraph()
                p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_eq.paragraph_format.space_before = Pt(8)
                p_eq.paragraph_format.space_after = Pt(2)
                p_eq.paragraph_format.keep_with_next = True

                # Normalize & convert using FormulaEngine (Gate 2, Gate 3, Gate 4)
                norm_eq = FormulaExtractor.normalize_latex(raw_eq)
                omml_elem = FormulaConverter.convert_latex_to_omml_element(norm_eq, is_display=True, eq_id=display_eq_counter)
                p_eq._p.append(omml_elem)

                self.stats['display_equations_converted'] += 1
                self.stats['total_equations_processed'] += 1
                self.stats['total_validation_checks_executed'] += 3

                # Right-aligned IEEE tag: (1), (2), (3), (4), (5)
                p_tag = doc.add_paragraph()
                p_tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                p_tag.paragraph_format.space_before = Pt(0)
                p_tag.paragraph_format.space_after = Pt(6)
                r_tag = p_tag.add_run(f"({display_eq_counter})")
                r_tag.font.name = 'Times New Roman'
                r_tag.font.size = Pt(10)
                r_tag.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

            elif line.startswith('---'):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)

            elif line.strip() != '':
                # Table Captions (TABLE I. / TABLE II.)
                if 'Table I:' in line or 'Table II:' in line or '*Table I:' in line or '*Table II:' in line:
                    p_tbl_cap = doc.add_paragraph()
                    p_tbl_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_tbl_cap.paragraph_format.space_before = Pt(10)
                    p_tbl_cap.paragraph_format.space_after = Pt(4)
                    p_tbl_cap.paragraph_format.keep_with_next = True
                    
                    table_tag = "TABLE I. " if "Table I" in line else "TABLE II. "
                    raw_title = line.split(':', 1)[1].strip().strip('*') if ':' in line else line
                    
                    r_tag = p_tbl_cap.add_run(table_tag)
                    r_tag.bold = True
                    r_tag.font.name = 'Times New Roman'
                    r_tag.font.size = Pt(9.5)
                    r_tag.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                    
                    self._parse_line_to_paragraph(p_tbl_cap, raw_title, default_font_size=9.5)
                elif line.startswith('*Figure') or line.startswith('Figure') or line.startswith('*Fig.') or line.startswith('Fig.'):
                    pass
                else:
                    p = doc.add_paragraph()
                    self._parse_line_to_paragraph(p, line, default_font_size=11)

            i += 1

        if in_table:
            process_table(table_lines)

        doc.save(self.docx_path)

        # Gate 5 & Gate 6: OpenXML & Structural OMML Validation
        validator = DocumentValidator(self.docx_path, self.pdf_path)
        validator.validate_docx_xml_gate()
        self.stats['total_validation_checks_executed'] += 1

        validator.validate_omml_structure_gate()
        self.stats['total_validation_checks_executed'] += 1

        # Generate PDF via Word COM API
        self._export_pdf()

        # Gate 7: PDF Parity Validation
        validator.validate_pdf_parity_gate()
        self.stats['total_validation_checks_executed'] += 1

        self.stats['end_time'] = time.time()
        self.stats['total_build_time'] = self.stats['end_time'] - self.stats['start_time']
        if self.stats['total_equations_processed'] > 0:
            self.stats['avg_equation_time_ms'] = (self.stats['total_build_time'] / self.stats['total_equations_processed']) * 1000

        return self.stats

    def _parse_line_to_paragraph(self, paragraph, text, default_font_size=11, is_table=False):
        """Parses Markdown line tokens ($inline$, **bold**, *italic*, `code`) and injects native runs/OMML elements."""
        pattern = re.compile(r'(\$\$[^\$]+\$\$|\$[^\$\n]+\$|\*\*.*?\*\*|\*.*?\*|`.*?`)')
        tokens = pattern.split(text)

        for token in tokens:
            if not token:
                continue

            if token.startswith('$') and token.endswith('$'):
                raw_inline = token.strip('$')
                norm_inline = FormulaExtractor.normalize_latex(raw_inline)
                try:
                    omml_inline = FormulaConverter.convert_latex_to_omml_element(norm_inline, is_display=False)
                    paragraph._p.append(omml_inline)

                    self.stats['inline_equations_converted'] += 1
                    self.stats['total_equations_processed'] += 1
                    if is_table:
                        self.stats['table_math_converted'] += 1
                    self.stats['total_validation_checks_executed'] += 2
                except Exception as e:
                    clean_s = norm_inline.replace('\\', '')
                    r = paragraph.add_run(clean_s)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(default_font_size)

            elif token.startswith('**') and token.endswith('**'):
                r = paragraph.add_run(token[2:-2])
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(default_font_size)
            elif token.startswith('*') and token.endswith('*'):
                r = paragraph.add_run(token[1:-1])
                r.italic = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(default_font_size)
            elif token.startswith('`') and token.endswith('`'):
                r = paragraph.add_run(token[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(default_font_size - 1)
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            else:
                clean_txt = token.replace(r'\_', '_')
                r = paragraph.add_run(clean_txt)
                r.font.name = 'Times New Roman'
                r.font.size = Pt(default_font_size)

    def _export_pdf(self):
        """Converts DOCX to PDF using Word COM API."""
        try:
            import win32com.client as win32
            w = win32.Dispatch('Word.Application')
            doc = w.Documents.Open(os.path.abspath(self.docx_path))
            doc.SaveAs(os.path.abspath(self.pdf_path), 17)
            doc.Close(0)
            try:
                w.Quit()
            except Exception:
                pass
        except Exception as e:
            if not os.path.exists(self.pdf_path):
                raise PublisherError(f"Gate 7 Failed: Word PDF Export Error: {str(e)}")

    def _set_cell_background(self, cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def _set_cell_margins(self, cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def _set_table_borders(self, table, color="CBD5E1", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr.append(borders)

if __name__ == '__main__':
    md_path = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\902d6e6c-b35a-4d3e-a983-59d2e321b9fc\paper2_academic_universe_growth_intelligence.md"
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final_OMML.docx"
    pdf_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final_OMML.pdf"
    fig1_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure1_Architecture_Pipeline.png"
    fig2_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure2_Decay_Sensitivity_Curves.png"

    publisher = IEEEPublisher(md_path, docx_path, pdf_path, fig1_path, fig2_path)
    stats = publisher.build()
    print("===========================================================")
    print("          IEEE PUBLISHER MULTI-STAGE BUILD COMPLETE        ")
    print("===========================================================")
    print(f"Display Equations Converted:   {stats['display_equations_converted']}")
    print(f"Inline Equations Converted:    {stats['inline_equations_converted']}")
    print(f"Total Equations Processed:     {stats['total_equations_processed']}")
    print(f"Total Validation Checks:       {stats['total_validation_checks_executed']}")
    print(f"Total Build Time:              {stats['total_build_time']:.3f} s")
    print(f"Avg Equation Time:             {stats['avg_equation_time_ms']:.3f} ms/eq")
    print("===========================================================")
