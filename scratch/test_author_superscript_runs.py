import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_AuthorSuperscript.docx"

doc = docx.Document(v4_docx_path)
p1 = doc.paragraphs[1]

# Clear existing runs in P1
p1.text = ""
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Build structured runs with native Word XML superscript properties
r1 = p1.add_run("Kushagra Singh Bhadauria")
r1.font.name = "Times New Roman"
r1.font.size = Pt(11.5)
r1.bold = True

r1_sup = p1.add_run("1")
r1_sup.font.name = "Times New Roman"
r1_sup.font.size = Pt(9)
r1_sup.font.superscript = True
r1_sup.bold = True

r2_sep = p1.add_run(", ")
r2_sep.font.name = "Times New Roman"
r2_sep.font.size = Pt(11.5)
r2_sep.bold = True

r2 = p1.add_run("Aashish Rajput")
r2.font.name = "Times New Roman"
r2.font.size = Pt(11.5)
r2.bold = True

r2_sup = p1.add_run("1")
r2_sup.font.name = "Times New Roman"
r2_sup.font.size = Pt(9)
r2_sup.font.superscript = True
r2_sup.bold = True

r3_sep = p1.add_run(", and ")
r3_sep.font.name = "Times New Roman"
r3_sep.font.size = Pt(11.5)
r3_sep.bold = True

r3 = p1.add_run("Avdesh Kumar Sah")
r3.font.name = "Times New Roman"
r3.font.size = Pt(11.5)
r3.bold = True

r3_sup = p1.add_run("1")
r3_sup.font.name = "Times New Roman"
r3_sup.font.size = Pt(9)
r3_sup.font.superscript = True
r3_sup.bold = True

r_br = p1.add_run("\n")

r_affil_sup = p1.add_run("1")
r_affil_sup.font.name = "Times New Roman"
r_affil_sup.font.size = Pt(8.5)
r_affil_sup.font.superscript = True

r_affil_text = p1.add_run("Department of Computer Science and Engineering, Sharda University, Greater Noida, Uttar Pradesh, India\nEmails: 2023361009.kushagra@ug.sharda.ac.in, 2023329421.aashish@ug.sharda.ac.in, 2023265132.avdesh@ug.sharda.ac.in")
r_affil_text.font.name = "Times New Roman"
r_affil_text.font.size = Pt(9.5)

print("=== VERIFYING P1 RUNS ===")
print(f"P1 Alignment: {p1.alignment}")
for i, r in enumerate(p1.runs):
    print(f"  Run {i}: text={repr(r.text)}, superscript={r.font.superscript}, size={r.font.size}")

doc.save(test_docx_path)
print(f"[SUCCESS] Saved test docx to {test_docx_path.name}")
