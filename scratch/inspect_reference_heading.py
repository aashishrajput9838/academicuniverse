import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

print("=== INSPECTING REFERENCE SECTION HEADING IN DOCX BASELINES ===")

doc4 = docx.Document(v4_docx_path)
for i, p in enumerate(doc4.paragraphs):
    if "[1]" in p.text or "References" in p.text or "REFERENCES" in p.text:
        print(f"V4 P{i} [{p.style.name}]: '{p.text[:100]}'")

print("\n--- Inspecting V5 DOCX ---")
doc5 = docx.Document(v5_docx_path)
for i, p in enumerate(doc5.paragraphs):
    if "[1]" in p.text or "References" in p.text or "REFERENCES" in p.text:
        print(f"V5 P{i} [{p.style.name}]: '{p.text[:100]}'")
