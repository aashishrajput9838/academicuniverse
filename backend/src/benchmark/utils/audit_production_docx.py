import docx

def audit_docx():
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final.docx"
    doc = docx.Document(docx_path)

    raw_markdown_asterisks = 0
    raw_backticks = 0
    replacement_chars = 0
    fig1_captions = 0
    fig2_captions = 0

    for p in doc.paragraphs:
        txt = p.text
        if '**' in txt:
            raw_markdown_asterisks += txt.count('**')
        if '`' in txt:
            raw_backticks += txt.count('`')
        if '\ufffd' in txt:
            replacement_chars += txt.count('\ufffd')
        if 'Fig. 1.' in txt or 'Figure 1' in txt:
            fig1_captions += 1
        if 'Fig. 2.' in txt or 'Figure 2' in txt:
            fig2_captions += 1

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                txt = c.text
                if '**' in txt:
                    raw_markdown_asterisks += txt.count('**')
                if '`' in txt:
                    raw_backticks += txt.count('`')
                if '\ufffd' in txt:
                    replacement_chars += txt.count('\ufffd')

    print("===========================================================")
    print("         IEEE PRODUCTION AUDIT & QUALITY CONTROL           ")
    print("===========================================================")
    print(f"Raw Markdown Asterisks ('**'): {raw_markdown_asterisks}")
    print(f"Raw Code Backticks ('`'):      {raw_backticks}")
    print(f"Unicode Replacement Chars (''): {replacement_chars}")
    print(f"Fig. 1 Caption Occurrences:    {fig1_captions} (Target: 1)")
    print(f"Fig. 2 Caption Occurrences:    {fig2_captions} (Target: 1)")
    print("===========================================================")

if __name__ == '__main__':
    audit_docx()
