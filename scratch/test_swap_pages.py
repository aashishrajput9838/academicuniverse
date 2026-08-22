import os
import docx
import re
import subprocess
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_Swapped.docx"
test_pdf_path = workspace / "docs" / "paper" / "PaperV5_Test_Swapped.pdf"

assert v4_docx_path.exists(), "PaperV4_Final_Submission.docx missing!"

doc = docx.Document(v4_docx_path)

p_title = doc.paragraphs[0]

# Insert CONTENTS heading at the very beginning of the document (BEFORE TITLE)
p_contents_hdr = p_title.insert_paragraph_before("CONTENTS")
p_contents_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contents_hdr.paragraph_format.space_before = Pt(14)
p_contents_hdr.paragraph_format.space_after = Pt(6)
if p_contents_hdr.runs:
    r = p_contents_hdr.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.bold = True

# Insert Word XML TOC Field
p_toc = p_title.insert_paragraph_before()
p_toc.paragraph_format.space_after = Pt(6)
pPr = p_toc._p.get_or_add_pPr()
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'TOCHeading')
pPr.append(pStyle)

r_toc = p_toc.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

r_toc._r.append(fldChar1)
r_toc._r.append(instrText)
r_toc._r.append(fldChar2)
r_toc._r.append(fldChar3)

# Populate visible TOC entries BEFORE p_title
toc_entries = [
    ("1. Introduction", "3", 1),
    ("1.1 Background & Motivation", "3", 2),
    ("1.2 Research Objectives", "3", 2),
    ("1.3 Main Research Contributions", "3", 2),
    ("2. Related Work", "5", 1),
    ("2.1 Vision-Language Models & Document Intelligence", "5", 2),
    ("2.2 Document Image Classification & Parsing Benchmarks", "5", 2),
    ("2.3 Information Extraction & Optical Character Recognition", "6", 2),
    ("2.4 Synthetic Data Generation & Benchmark Suites", "6", 2),
    ("2.5 Semantic Normalization & Post-Processing", "7", 2),
    ("2.6 Error Taxonomies & Diagnostic Evaluation", "7", 2),
    ("3. ADBG v1.0 & AU DIC Benchmark System Architecture", "8", 1),
    ("3.1 ADBG v1.0 Synthetic Generator Subsystem", "8", 2),
    ("3.2 Optical Degradation Profile Processor", "8", 2),
    ("3.3 AU DIC Decoupled Benchmark Execution Subsystem", "9", 2),
    ("4. Methodology", "9", 1),
    ("4.1 Six-Stage Semantic Canonical Normalization Subsystem", "9", 2),
    ("4.2 Nine-Class Structured OCR Error Taxonomy", "12", 2),
    ("4.3 Mathematical Formulation of Evaluation Metrics", "13", 2),
    ("5. Results & Empirical Validation", "15", 1),
    ("5.1 Distinction Between Framework Validation, Benchmark Validation, and Model Performance", "15", 2),
    ("5.2 Framework Execution & System Verification Metrics", "15", 2),
    ("5.3 System Throughput & Execution Latency", "16", 2),
    ("5.4 Empirical Live Neural Model Evaluation Results", "16", 2),
    ("5.5 Empirical Ablation Study of Semantic Canonical Normalization", "18", 2),
    ("5.6 Statistical Significance Analysis (p < 0.0001)", "20", 2),
    ("5.7 Empirical Evaluation Scope & Methodological Limitations", "21", 2),
    ("5.8 Error Taxonomy Distribution Shift Analysis", "21", 2),
    ("6. Discussion & Threats to Validity", "22", 1),
    ("6.1 Scientific Contributions and Methodological Novelty", "22", 2),
    ("6.2 Discussion of Empirical Findings", "22", 2),
    ("6.3 Threats to Validity", "23", 2),
    ("7. Limitations Analysis", "23", 1),
    ("7.1 Methodological Limitations", "23", 2),
    ("8. Future Work", "24", 1),
    ("9. Conclusion", "24", 1),
    ("Ethics & Privacy Statement", "24", 1),
    ("ACKNOWLEDGMENT", "25", 1),
    ("APPENDIX A: REPRODUCIBILITY & SYSTEM SPECIFICATIONS", "25", 1),
    ("A.1 Reproducibility & System Environment Matrix", "25", 2),
    ("A.2 Technical Clarifications & Reviewer Inquiries", "25", 2),
    ("APPENDIX B: FIELD SPECIFICATION & OBSERVATION COUNT DERIVATION", "26", 1),
    ("B.1 Document Category Field Structure", "26", 2),
    ("B.2 Mathematical Derivation of 24,480 Paired Observations", "26", 2),
    ("APPENDIX C: EMPIRICAL STATISTICAL METHODOLOGY & BENCHMARKS", "27", 1),
    ("C.1 Empirical Category Confusion Matrix (360 Specimens)", "27", 2),
    ("C.2 McNemar Contingency Test & Normalization Rescues", "27", 2),
    ("C.3 Non-Parametric Bootstrap Confidence Intervals (B = 10,000)", "28", 2),
    ("REFERENCES", "28", 1),
]

for title, page_str, level in toc_entries:
    p_entry = p_title.insert_paragraph_before()
    p_entry.paragraph_format.line_spacing = 1.15
    p_entry.paragraph_format.space_after = Pt(2)
    p_entry.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    indent_prefix = "   " if level == 2 else ""
    full_title = f"{indent_prefix}{title}"
    
    r_title = p_entry.add_run(full_title)
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(9.5)
    if level == 1:
        r_title.bold = True
    
    r_page = p_entry.add_run(f"\t{page_str}")
    r_page.font.name = "Times New Roman"
    r_page.font.size = Pt(9.5)
    r_page.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# Set Page Break on p_title so Title starts at the top of Page 2!
p_title.paragraph_format.page_break_before = True

# Also ensure 1. Introduction starts on Page 3 (page break before 1. Introduction)
for p in doc.paragraphs:
    if p.text.strip() == "1. Introduction":
        p.paragraph_format.page_break_before = True
        break

doc.save(test_docx_path)
print(f"[SUCCESS] Saved test DOCX with Page 1 = CONTENTS, Page 2 = Title/Authors/Abstract, Page 3 = Introduction")
