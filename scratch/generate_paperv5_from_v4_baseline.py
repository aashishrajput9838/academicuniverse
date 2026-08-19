import os
import docx
import re
import sys
import subprocess
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
logo_path = workspace / "public" / "new_logo_2.png"
dt_cm_path = workspace / "results" / "confusion_matrices" / "dt_composite.png"
rf_cm_path = workspace / "results" / "confusion_matrices" / "rf_composite.png"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"
v5_pdf_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.pdf"

print("============================================================")
print(" GENERATING PAPER V5 WITH ML BENCHMARK SECTION 5.9 INTEGRATED")
print("============================================================")

# Force kill lingering Word/WPS background processes
try:
    subprocess.run(["taskkill", "/F", "/IM", "wps.exe", "/IM", "wpscenter.exe", "/IM", "WINWORD.EXE"], capture_output=True)
except Exception:
    pass

assert v4_docx_path.exists(), f"Error: Baseline {v4_docx_path} missing!"
assert logo_path.exists(), f"Error: Logo file missing at {logo_path}!"

doc = docx.Document(v4_docx_path)

# 0. Pipeline Fix: Set Exact New Paper Title
p_title_orig = doc.paragraphs[0]
p_title_orig.text = "Smart Academic Document Intelligence System: Automated Extraction, Normalization, and Benchmark Generation"
p_title_orig.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title_orig.paragraph_format.space_before = Pt(0)
p_title_orig.paragraph_format.space_after = Pt(12)
if p_title_orig.runs:
    for r in p_title_orig.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(20)
        r.bold = True

def set_paragraph_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)

# 1. Clear incomplete/dummy footer text
for s_idx, section in enumerate(doc.sections):
    footer = section.footer
    for p in footer.paragraphs:
        if "IEEE ACCESS" in p.text or "Page" in p.text:
            p.text = ""

# 2. Format Author Paragraph P1 with Exact 3-Line Structure
p1 = doc.paragraphs[1]
p1.text = ""
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.line_spacing = 1.15
p1.paragraph_format.space_after = Pt(8)

# LINE 1 — AUTHORS
r1 = p1.add_run("Kushagra Singh Bhadauria")
r1.font.name = "Times New Roman"
r1.font.size = Pt(11.5)
r1.bold = True

r1_sup = p1.add_run("1")
r1_sup.font.name = "Times New Roman"
r1_sup.font.size = Pt(9)
r1_sup.font.superscript = True
r1_sup.bold = True

r2_sep = p1.add_run(", ")
r2_sep.font.name = "Times New Roman"
r2_sep.font.size = Pt(11.5)
r2_sep.bold = True

r2 = p1.add_run("Aashish Rajput")
r2.font.name = "Times New Roman"
r2.font.size = Pt(11.5)
r2.bold = True

r2_sup = p1.add_run("2")
r2_sup.font.name = "Times New Roman"
r2_sup.font.size = Pt(9)
r2_sup.font.superscript = True
r2_sup.bold = True

r3_sep = p1.add_run(", and ")
r3_sep.font.name = "Times New Roman"
r3_sep.font.size = Pt(11.5)
r3_sep.bold = True

r3 = p1.add_run("Avdesh Kumar Sah")
r3.font.name = "Times New Roman"
r3.font.size = Pt(11.5)
r3.bold = True

r3_sup = p1.add_run("3")
r3_sup.font.name = "Times New Roman"
r3_sup.font.size = Pt(9)
r3_sup.font.superscript = True
r3_sup.bold = True

# LINE BREAK 1 -> LINE 2 — SHARED AFFILIATION
p1.add_run("\n")
r_affil_sup = p1.add_run("123")
r_affil_sup.font.name = "Times New Roman"
r_affil_sup.font.size = Pt(8.5)
r_affil_sup.font.superscript = True

r_affil_text = p1.add_run(" Department of Computer Science and Engineering, Sharda University, Greater Noida, Uttar Pradesh, India")
r_affil_text.font.name = "Times New Roman"
r_affil_text.font.size = Pt(9.5)

# LINE BREAK 2 -> LINE 3 — EMAILS
p1.add_run("\n")
r_e1_sup = p1.add_run("1")
r_e1_sup.font.name = "Times New Roman"
r_e1_sup.font.size = Pt(8.5)
r_e1_sup.font.superscript = True
p1.add_run(" 2023361009.kushagra@ug.sharda.ac.in, ").font.size = Pt(9.5)

r_e2_sup = p1.add_run("2")
r_e2_sup.font.name = "Times New Roman"
r_e2_sup.font.size = Pt(8.5)
r_e2_sup.font.superscript = True
p1.add_run(" 2023329421.aashish@ug.sharda.ac.in, ").font.size = Pt(9.5)

r_e3_sup = p1.add_run("3")
r_e3_sup.font.name = "Times New Roman"
r_e3_sup.font.size = Pt(8.5)
r_e3_sup.font.superscript = True
p1.add_run(" 2023265132.avdesh@ug.sharda.ac.in").font.size = Pt(9.5)

# 3. Clear misplaced "References" header before Appendix A
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References" and i < 240:
        p.text = ""

# 4. Insert IEEE-formatted REFERENCES section heading immediately before reference [1]
ref1_inserted = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("[1] "):
        p_hdr = p.insert_paragraph_before("REFERENCES")
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after = Pt(6)
        if p_hdr.runs:
            r = p_hdr.runs[0]
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.bold = True
        ref1_inserted = True
        break

assert ref1_inserted, "Error: Could not locate reference [1] entry in manuscript!"

# Text replacements mapping for canonical empirical results
replacements = [
    ("Llama-3.1-8B-Instant", "MiniCPM-V (7.6B Q4_0 GGUF)"),
    ("Llama 3.1 8B", "MiniCPM-V (7.6B Q4_0)"),
    ("llama-3.1-8b-instant", "minicpm-v:latest"),
    ("Groq Cloud vision API", "Ollama Local Model-Serving Runtime (v0.32.14)"),
    ("Groq Cloud endpoint", "Local Ollama Inference Runtime"),
    ("Groq Cloud", "Ollama Local Runtime"),
    ("Groq API", "Ollama Local Model-Serving Engine"),
    ("Groq", "Ollama"),
    ("cloud-based inference", "local offline model-serving runtime inference"),
    ("10.16%", "74.60%"),
    ("10.84%", "82.18%"),
    ("17.19%", "75.23%"),
    ("89.27%", "11.35%"),
    ("82.76%", "8.21%"),
    ("165.01", "1853.0005"),
    ("Raw Exact Match Rate of 10.16%", "Raw Exact Match Rate of 74.60%"),
    ("Normalized Exact Match Rate of 10.84%", "Normalized Exact Match Rate of 82.18%"),
    ("Field F1 Score of 17.19%", "Field F1 Score of 75.23%"),
    ("Mean CER of 89.27%", "Mean CER of 11.35%"),
    ("clean profile exact match rate of 25.0%", "clean profile exact match rate of 90.00%"),
    ("rotated_90 profile CER of 99.5%", "rotated_90 profile CER of 29.02%"),
    ("5 document categories", "3 primary academic document categories"),
    ("five document categories", "three primary academic document categories"),
]

for p in doc.paragraphs:
    text_str = p.text.strip()
    if not text_str or "Kushagra Singh Bhadauria" in text_str:
        continue
    for old_text, new_text in replacements:
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text)

    style_name = p.style.name.lower()
    is_special = (
        "heading" in style_name or "title" in style_name or "subtitle" in style_name or
        "caption" in style_name or text_str == "CONTENTS" or text_str == "REFERENCES" or
        text_str.startswith("Figure ") or text_str.startswith("Table ") or text_str.startswith("TABLE ") or
        p.alignment == WD_ALIGN_PARAGRAPH.CENTER or "<m:oMath" in p._p.xml
    )
    if not is_special:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old_text, new_text in replacements:
                    if old_text in p.text:
                        p.text = p.text.replace(old_text, new_text)

# 4.1 Restructure Section 4.3 into Subsections 4.3.1 to 4.3.6
sec43_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("4.3 Evaluation Metrics") or p.text.strip().startswith("4.3 Mathematical Formulation"):
        sec43_idx = i
        break

if sec43_idx is not None:
    doc.paragraphs[sec43_idx].text = "4.3 Mathematical Formulation of Evaluation Metrics"
    set_paragraph_outline_level(doc.paragraphs[sec43_idx], 0)
    
    # Locate Section 4.3 item paragraphs (1. to 6.)
    sec5_idx = None
    for j in range(sec43_idx + 1, len(doc.paragraphs)):
        if doc.paragraphs[j].text.strip().startswith("5. Results & Empirical Validation"):
            sec5_idx = j
            break
            
    if sec5_idx is not None:
        p_sec5 = doc.paragraphs[sec5_idx]
        
        # Remove old monolithic metric paragraphs between sec43_idx+1 and sec5_idx
        for k in range(sec5_idx - 1, sec43_idx + 1, -1):
            p_rem = doc.paragraphs[k]
            p_rem.text = ""

        # Insert 4.3.1 - 4.3.6 structured subsections before Section 5
        subsections_43 = [
            ("The quantitative evaluation metrics formulated for evaluating neural document intelligence models, synthetic specimen benchmark suites, and canonical normalization layers are defined below.", None, None),
            
            ("4.3.1 Category Classification Accuracy",
             "The proportion of document specimens where the predicted document category (y_hat) matches the ground truth category (y) across N total evaluated specimens:",
             "Acc_cat = (1 / N) * sum_{i=1}^N I(y_hat_i = y_i)   (1)\nwhere N represents the total count of evaluated document specimens."),
             
            ("4.3.2 Field Extraction Precision, Recall, and F1-Score",
             "Macro-averaged key-value field extraction metrics across all extracted target entities:",
             "Precision = TP / (TP + FP), Recall = TP / (TP + FN), F1 = (2 * Precision * Recall) / (Precision + Recall)   (2)"),
             
            ("4.3.3 Character Error Rate (CER)",
             "Levenshtein character edit distance [21] between canonically normalized predicted field strings (ŝ) and expected ground truth field strings (s), normalized by total ground truth character length (|s|):",
             "CER = D_char(ŝ, s) / |s|   (3)"),
             
            ("4.3.4 Word Error Rate (WER)",
             "Tokenized word-level edit distance [20] between predicted field strings (ŵ) and ground truth field strings (w), normalized by total ground truth word count (|w|):",
             "WER = D_word(ŵ, w) / |w|   (4)"),
             
            ("4.3.5 Joint Record Exact Match Rate (EM)",
             "The percentage of specimens that achieve both 100% key-value field extraction (F1 = 1.0) AND correct top-level category classification simultaneously:",
             "EM = (1 / N) * sum_{i=1}^N I(ŷ_i = y_i AND F1_i = 1.0)   (5)"),
             
            ("4.3.6 Execution Latency & Throughput",
             "Execution latency per specimen (L_proc) measured in milliseconds per sample (ms/sample) and framework processing throughput (TH) measured in specimens per second (samples/sec):",
             "L_proc = T_total / N, TH = N / T_total   (6)")
        ]
        
        for item in subsections_43:
            heading, desc, eq = item
            if heading and heading.startswith("4.3."):
                p_h = p_sec5.insert_paragraph_before(heading)
                p_h.paragraph_format.space_before = Pt(8)
                p_h.paragraph_format.space_after = Pt(2)
                if p_h.runs:
                    r = p_h.runs[0]
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10.5)
                    r.bold = True
                set_paragraph_outline_level(p_h, 1)
            elif heading and not heading.startswith("4.3."):
                p_intro = p_sec5.insert_paragraph_before(heading)
                p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_intro.paragraph_format.space_after = Pt(4)
                
            if desc:
                p_d = p_sec5.insert_paragraph_before(desc)
                p_d.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_d.paragraph_format.space_after = Pt(2)
                
            if eq:
                p_eq = p_sec5.insert_paragraph_before(eq)
                p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_eq.paragraph_format.space_after = Pt(6)
                if p_eq.runs:
                    r = p_eq.runs[0]
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(9.5)
                    r.font.italic = True

# 5. Pipeline Fix: Insert Section 5.9 (Classical Machine Learning Classification Benchmark)
sec6_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "6. Discussion & Threats to Validity":
        sec6_idx = i
        break

assert sec6_idx is not None, "Error: Could not locate Section 6 heading!"

p_sec6 = doc.paragraphs[sec6_idx]

# Insert Section 5.9 Heading & Content BEFORE Section 6
p_s59 = p_sec6.insert_paragraph_before("5.9 Classical Machine Learning Classification Benchmark")
p_s59.paragraph_format.space_before = Pt(12)
p_s59.paragraph_format.space_after = Pt(4)
if p_s59.runs:
    r = p_s59.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.bold = True
set_paragraph_outline_level(p_s59, 1)

# Paragraph 1: Setup & Dataset
p_s59_t1 = p_sec6.insert_paragraph_before(
    "To complement end-to-end neural vision-language model evaluation and assess the predictability of field extraction failures "
    "from structural, optical, and length features, we conduct a classical supervised machine learning benchmark. We evaluate two foundational "
    "tree-based classification algorithms—Decision Tree (DT) and Random Forest (RF)—to predict whether a given extraction observation results "
    "in an exact field match (y = 1) or an extraction mismatch/OCR failure (y = 0). The evaluation dataset comprises 24,480 paired field observations "
    "across 360 specimens extracted from live benchmark evaluations (paired_field_observations.csv), containing 18,263 exact match instances (74.60%) "
    "and 6,217 mismatch instances (25.40%)."
)
p_s59_t1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_s59_t1.paragraph_format.space_after = Pt(4)

# Paragraph 2: Features & Anti-Leakage Pipeline
p_s59_t2 = p_sec6.insert_paragraph_before(
    "Six non-leaky input features are extracted for each observation: document_type (3 categories: student_id, marksheet, certificate), "
    "quality_profile (4 optical degradation profiles: clean, mobile_camera, scanner_copy, rotated_90), field_name (68 document field categories), "
    "expected_len (ground truth string character length), predicted_len (extracted string character length), and is_missing (binary flag indicating "
    "unextracted null predictions). To prevent data leakage, all preprocessing transformers—including OneHotEncoder(handle_unknown='ignore') for "
    "categorical variables and StandardScaler() for numerical features—are encapsulated inside an sklearn.pipeline.Pipeline and fitted exclusively "
    "on the training split (X_train) during model fitting."
)
p_s59_t2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_s59_t2.paragraph_format.space_after = Pt(4)

# Paragraph 3: Splits & Hyperparameters
p_s59_t3 = p_sec6.insert_paragraph_before(
    "Classifiers are benchmarked across three stratified train-test partitions: 60:40 (N_train = 14,688, N_test = 9,792), 70:30 (N_train = 17,136, "
    "N_test = 7,344), and 80:20 (N_train = 19,584, N_test = 4,896) using a fixed seed (random_state = 42). Random Forest models are parameterized with "
    "n_estimators = 100, max_depth = 15, and criterion = 'gini'. Decision Tree models use max_depth = 15 and criterion = 'gini'. Prediction inference "
    "latency is measured exclusively around pipeline.predict(X_test) using high-resolution timers (time.perf_counter()). Table IX summarizes the "
    "twelve derived confusion-matrix performance metrics across all six model-split configurations."
)
p_s59_t3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_s59_t3.paragraph_format.space_after = Pt(6)

# Insert Table IX Title
p_tbl_title = p_sec6.insert_paragraph_before("TABLE IX: CLASSICAL MACHINE LEARNING BENCHMARK COMPARISON (RF VS. DT ACROSS TRAIN-TEST SPLITS)")
p_tbl_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tbl_title.paragraph_format.space_before = Pt(8)
p_tbl_title.paragraph_format.space_after = Pt(4)
if p_tbl_title.runs:
    r = p_tbl_title.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(9.5)
    r.bold = True

# Insert Table IX (7 columns x 13 rows)
tbl_data = [
    ["Metric", "RF 60:40", "RF 70:30", "RF 80:20", "DT 60:40", "DT 70:30", "DT 80:20"],
    ["Accuracy", "0.874898", "0.875817", "0.878676", "0.934947", "0.935185", "0.936887"],
    ["Precision", "0.856389", "0.857299", "0.860137", "0.927326", "0.925498", "0.928059"],
    ["Recall", "1.000000", "1.000000", "1.000000", "0.990418", "0.993064", "0.992335"],
    ["F1-Score", "0.922640", "0.923168", "0.924810", "0.957834", "0.958091", "0.959122"],
    ["Specificity", "0.507439", "0.510992", "0.522124", "0.772014", "0.765147", "0.773934"],
    ["NPV", "1.000000", "1.000000", "1.000000", "0.964824", "0.974061", "0.971717"],
    ["MCC", "0.659215", "0.661871", "0.670148", "0.824745", "0.825867", "0.830344"],
    ["FPR", "0.492561", "0.489008", "0.477876", "0.227986", "0.234853", "0.226066"],
    ["FNR", "0.000000", "0.000000", "0.000000", "0.009582", "0.006936", "0.007665"],
    ["FDR", "0.143611", "0.142701", "0.139863", "0.072674", "0.074502", "0.071941"],
    ["FOR", "0.000000", "0.000000", "0.000000", "0.000000", "0.000000", "0.000000"],
    ["Prediction Time (s)", "0.167798", "0.146718", "0.120588", "0.025032", "0.018994", "0.016724"],
]

p_table_wrap = p_sec6.insert_paragraph_before()
table_elem = doc.add_table(rows=len(tbl_data), cols=7)
p_table_wrap._p.addnext(table_elem._element)

for r_idx, row in enumerate(table_elem.rows):
    for c_idx, cell in enumerate(row.cells):
        cell.text = tbl_data[r_idx][c_idx]
        p_c = cell.paragraphs[0]
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if p_c.runs:
            r = p_c.runs[0]
            r.font.name = "Times New Roman"
            r.font.size = Pt(8.5)
            if r_idx == 0 or c_idx == 0:
                r.bold = True

# Paragraph 4: Findings & Decision Tree Superiority Discussion
p_s59_t4 = p_sec6.insert_paragraph_before(
    "As demonstrated in Table IX, Decision Tree classifiers consistently outperform Random Forest across all evaluation partitions. The DT 80:20 model "
    "achieves top performance with 93.69% classification accuracy, 95.91% F1-score, 77.39% specificity, and an MCC of 0.8303, completing 4,896 test inferences "
    "in only 16.72 ms (3.41 µs/specimen). By contrast, RF 80:20 achieves 87.87% accuracy and 92.48% F1-score with 120.59 ms prediction latency."
)
p_s59_t4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_s59_t4.paragraph_format.space_before = Pt(6)
p_s59_t4.paragraph_format.space_after = Pt(4)

# Paragraph 5: Theoretical Rationale
p_s59_t5 = p_sec6.insert_paragraph_before(
    "The superior accuracy of single Decision Trees over Random Forests in this domain stems from the discrete, axis-aligned partition structure of "
    "the feature space. Critical optical degradations—specifically severe 90° rotation (quality_profile = rotated_90) or unextracted null values "
    "(is_missing = 1)—exhibit deterministic step-function relationships with field extraction failures. A single deep Decision Tree directly isolates "
    "these sharp decision boundaries without hyper-plane smoothing. Conversely, Random Forest bagging averages probability estimates across 100 decorrelated "
    "trees, smoothing out crisp binary splits and slightly inflating false positive rates on non-degraded specimens (FPR = 0.4779 for RF 80:20 vs FPR = 0.2261 for DT 80:20). "
    "Figures 10 and 11 display the composite confusion matrices for Decision Tree and Random Forest classifiers across all three train-test splits."
)
p_s59_t5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_s59_t5.paragraph_format.space_after = Pt(6)

# Figure 10: Decision Tree Composite Confusion Matrix
if dt_cm_path.exists():
    p_fig10_img = p_sec6.insert_paragraph_before()
    p_fig10_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig10_img.paragraph_format.space_before = Pt(8)
    p_fig10_img.paragraph_format.space_after = Pt(2)
    r_f10 = p_fig10_img.add_run()
    r_f10.add_picture(str(dt_cm_path), width=Inches(6.2))

    p_fig10_cap = p_sec6.insert_paragraph_before("Fig. 10. Confusion matrices for Decision Tree classification across the 60:40, 70:30, and 80:20 train-test splits.")
    p_fig10_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig10_cap.paragraph_format.space_after = Pt(8)
    if p_fig10_cap.runs:
        r = p_fig10_cap.runs[0]
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        r.font.italic = True

# Figure 11: Random Forest Composite Confusion Matrix
if rf_cm_path.exists():
    p_fig11_img = p_sec6.insert_paragraph_before()
    p_fig11_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig11_img.paragraph_format.space_before = Pt(4)
    p_fig11_img.paragraph_format.space_after = Pt(2)
    r_f11 = p_fig11_img.add_run()
    r_f11.add_picture(str(rf_cm_path), width=Inches(6.2))

    p_fig11_cap = p_sec6.insert_paragraph_before("Fig. 11. Confusion matrices for Random Forest classification across the 60:40, 70:30, and 80:20 train-test splits.")
    p_fig11_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fig11_cap.paragraph_format.space_after = Pt(12)
    if p_fig11_cap.runs:
        r = p_fig11_cap.runs[0]
        r.font.name = "Times New Roman"
        r.font.size = Pt(9)
        r.font.italic = True

# 6. Pipeline Fix: Assign XML Outline Levels to all section & subsection headings
for p in doc.paragraphs:
    text_str = p.text.strip()
    if not text_str:
        continue
    
    is_h1 = (
        text_str.startswith("1. ") or text_str.startswith("2. ") or text_str.startswith("3. ") or
        text_str.startswith("4. ") or text_str.startswith("5. ") or text_str.startswith("6. ") or
        text_str.startswith("7. ") or text_str.startswith("8. ") or text_str.startswith("9. ") or
        text_str == "REFERENCES" or text_str.startswith("APPENDIX A") or text_str.startswith("APPENDIX B") or
        text_str.startswith("APPENDIX C") or text_str.startswith("Ethics & Privacy Statement") or text_str.startswith("ACKNOWLEDGMENT")
    )
    
    is_h2 = bool(
        re.match(r"^\d\.\d(\.\d)?\s+", text_str) or
        text_str.startswith("A.1 ") or text_str.startswith("A.2 ") or
        text_str.startswith("B.1 ") or text_str.startswith("B.2 ") or
        text_str.startswith("C.1 ") or text_str.startswith("C.2 ") or text_str.startswith("C.3 ")
    )
    
    if is_h1:
        set_paragraph_outline_level(p, 0)
    elif is_h2:
        set_paragraph_outline_level(p, 1)

# 7. Pipeline Fix: Front-Matter Page Layout & Native TOC
p_title = doc.paragraphs[0]

# PAGE 1 -> Sharda University Logo ONLY
p_logo = p_title.insert_paragraph_before()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(230) # Vertical centering on Page 1
p_logo.paragraph_format.space_after = Pt(0)
r_logo = p_logo.add_run()
r_logo.add_picture(str(logo_path), width=Inches(4.5))

# PAGE 2 -> CONTENTS Heading (page break before CONTENTS)
p_contents_hdr = p_title.insert_paragraph_before("CONTENTS")
p_contents_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contents_hdr.paragraph_format.page_break_before = True # Starts CONTENTS on Page 2
p_contents_hdr.paragraph_format.space_before = Pt(14)
p_contents_hdr.paragraph_format.space_after = Pt(6)
if p_contents_hdr.runs:
    r = p_contents_hdr.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.bold = True

# Word XML TOC Field
p_toc = p_title.insert_paragraph_before()
p_toc.paragraph_format.space_after = Pt(6)
pPr = p_toc._p.get_or_add_pPr()
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'TOCHeading')
pPr.append(pStyle)

r_toc = p_toc.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

r_toc._r.append(fldChar1)
r_toc._r.append(instrText)
r_toc._r.append(fldChar2)
r_toc._r.append(fldChar3)

# Populate visible TOC entries with RIGHT-ALIGNED TAB STOP & DOT LEADER at 6.5 inches
# Page Numbers updated for 33/34-page IEEE Access manuscript layout:
# Page 1 = Logo, Page 2 = CONTENTS, Page 3 = Front Matter, Page 4 = Introduction
toc_entries = [
    ("1. Introduction", "4", 1),
    ("1.1 Background & Motivation", "4", 2),
    ("1.2 Research Objectives", "4", 2),
    ("1.3 Main Research Contributions", "4", 2),
    ("2. Related Work", "6", 1),
    ("2.1 Vision-Language Models & Document Intelligence", "6", 2),
    ("2.2 Document Image Classification & Parsing Benchmarks", "6", 2),
    ("2.3 Information Extraction & Optical Character Recognition", "7", 2),
    ("2.4 Synthetic Data Generation & Benchmark Suites", "7", 2),
    ("2.5 Semantic Normalization & Post-Processing", "8", 2),
    ("2.6 Error Taxonomies & Diagnostic Evaluation", "8", 2),
    ("3. ADBG v1.0 & AU DIC Benchmark System Architecture", "9", 1),
    ("3.1 ADBG v1.0 Synthetic Generator Subsystem", "9", 2),
    ("3.2 Optical Degradation Profile Processor", "9", 2),
    ("3.3 AU DIC Decoupled Benchmark Execution Subsystem", "10", 2),
    ("4. Methodology", "10", 1),
    ("4.1 Six-Stage Semantic Canonical Normalization Subsystem", "10", 2),
    ("4.2 Nine-Class Structured OCR Error Taxonomy", "13", 2),
    ("4.3 Mathematical Formulation of Evaluation Metrics", "14", 1),
    ("4.3.1 Category Classification Accuracy", "14", 2),
    ("4.3.2 Field Extraction Precision, Recall, and F1-Score", "14", 2),
    ("4.3.3 Character Error Rate (CER)", "14", 2),
    ("4.3.4 Word Error Rate (WER)", "15", 2),
    ("4.3.5 Joint Record Exact Match Rate (EM)", "15", 2),
    ("4.3.6 Execution Latency & Throughput", "15", 2),
    ("5. Results & Empirical Validation", "16", 1),
    ("5.1 Distinction Between Framework Validation, Benchmark Validation, and Model Performance", "16", 2),
    ("5.2 Framework Execution & System Verification Metrics", "16", 2),
    ("5.3 System Throughput & Execution Latency", "17", 2),
    ("5.4 Empirical Live Neural Model Evaluation Results", "17", 2),
    ("5.5 Empirical Ablation Study of Semantic Canonical Normalization", "19", 2),
    ("5.6 Statistical Significance Analysis (p < 0.0001)", "21", 2),
    ("5.7 Empirical Evaluation Scope & Methodological Limitations", "22", 2),
    ("5.8 Error Taxonomy Distribution Shift Analysis", "22", 2),
    ("5.9 Classical Machine Learning Classification Benchmark", "23", 2),
    ("6. Discussion & Threats to Validity", "24", 1),
    ("6.1 Scientific Contributions and Methodological Novelty", "24", 2),
    ("6.2 Discussion of Empirical Findings", "24", 2),
    ("6.3 Threats to Validity", "25", 2),
    ("7. Limitations Analysis", "25", 1),
    ("7.1 Methodological Limitations", "25", 2),
    ("8. Future Work", "26", 1),
    ("9. Conclusion", "26", 1),
    ("Ethics & Privacy Statement", "26", 1),
    ("ACKNOWLEDGMENT", "27", 1),
    ("APPENDIX A: REPRODUCIBILITY & SYSTEM SPECIFICATIONS", "27", 1),
    ("A.1 Reproducibility & System Environment Matrix", "27", 2),
    ("A.2 Technical Clarifications & Reviewer Inquiries", "27", 2),
    ("APPENDIX B: FIELD SPECIFICATION & OBSERVATION COUNT DERIVATION", "28", 1),
    ("B.1 Document Category Field Structure", "28", 2),
    ("B.2 Mathematical Derivation of 24,480 Paired Observations", "28", 2),
    ("APPENDIX C: EMPIRICAL STATISTICAL METHODOLOGY & BENCHMARKS", "29", 1),
    ("C.1 Empirical Category Confusion Matrix (360 Specimens)", "29", 2),
    ("C.2 McNemar Contingency Test & Normalization Rescues", "29", 2),
    ("C.3 Non-Parametric Bootstrap Confidence Intervals (B = 10,000)", "30", 2),
    ("REFERENCES", "30", 1),
]

for title, page_str, level in toc_entries:
    p_entry = p_title.insert_paragraph_before()
    p_entry.paragraph_format.line_spacing = 1.15
    p_entry.paragraph_format.space_after = Pt(2)
    p_entry.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    indent_prefix = "   " if level == 2 else ""
    full_title = f"{indent_prefix}{title}"
    
    r_title = p_entry.add_run(full_title)
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(9.5)
    if level == 1:
        r_title.bold = True
    
    r_page = p_entry.add_run(f"\t{page_str}")
    r_page.font.name = "Times New Roman"
    r_page.font.size = Pt(9.5)
    r_page.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# Set Page Break on p_title so Title starts on PAGE 3
p_title.paragraph_format.page_break_before = True

# Set Page Break on 1. Introduction so Main Paper starts on PAGE 4
for p in doc.paragraphs:
    if p.text.strip() == "1. Introduction":
        p.paragraph_format.page_break_before = True
        break

# Save directly to PaperV5_Ollama_Primary.docx
doc.save(v5_docx_path)
print(f"[SUCCESS] Saved PaperV5_Ollama_Primary.docx with Section 5.9 integrated!")

# Export PDF via Word COM Automation
print(f"Exporting PDF via Word COM Automation: {v5_pdf_path.name}...")
abs_docx = os.path.abspath(str(v5_docx_path))
abs_pdf = os.path.abspath(str(v5_pdf_path))

word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_word = word.Documents.Open(abs_docx)
    doc_word.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
    page_count = doc_word.ComputeStatistics(2)
    doc_word.Close()
    print(f"[SUCCESS] Exported high-quality PDF: {v5_pdf_path.name} ({page_count} Pages!)")
except Exception as e:
    print(f"Word COM PDF Export Error: {e}")
finally:
    if word:
        try: word.Quit()
        except: pass
