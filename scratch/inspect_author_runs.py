import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

print("=== INSPECTING RUNS IN AUTHOR PARAGRAPH P1 ===")

doc = docx.Document(v4_docx_path)
p1 = doc.paragraphs[1]
print("V4 P1 Raw Text:", repr(p1.text))
print(f"V4 P1 Run Count: {len(p1.runs)}")
for i, r in enumerate(p1.runs):
    print(f"  Run {i} (superscript={r.font.superscript}, font_name={r.font.name}, font_size={r.font.size}): repr={repr(r.text)}")

print("\n--- V5 P1 ---")
doc5 = docx.Document(v5_docx_path)
p1_5 = doc5.paragraphs[1]
print("V5 P1 Raw Text:", repr(p1_5.text))
print(f"V5 P1 Run Count: {len(p1_5.runs)}")
for i, r in enumerate(p1_5.runs):
    print(f"  Run {i} (superscript={r.font.superscript}, font_name={r.font.name}, font_size={r.font.size}): repr={repr(r.text)}")
