import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def style_table_header(row, bg_color="1F497D", text_color="FFFFFF"):
    for cell in row.cells:
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9.5)
                run.font.name = "Calibri"

def style_table_cells(table, font_size=9):
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            continue
        bg_color = "F2F5F9" if r_idx % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Calibri"

def format_document(doc_path):
    print(f"Processing: {doc_path}")
    doc = Document(doc_path)
    
    # 1. Identify section indices
    p_texts = [p.text.strip() for p in doc.paragraphs]
    
    # Locate where messy pre-ack appendix starts (e.g. 'Appendix A: Reproducibility Specifications')
    draft_app_start = -1
    ack_idx = -1
    ref_idx = -1
    blob_app_idx = -1
    
    for i, t in enumerate(p_texts):
        if draft_app_start == -1 and t.startswith("Appendix A: Reproducibility"):
            draft_app_start = i
        elif "ACKNOWLEDGMENT" in t and len(t) < 50:
            ack_idx = i
        elif t.startswith("ACKNOWLEDGMENT\n"):
            ack_idx = i
        elif t == "References" or t.startswith("References"):
            ref_idx = i
        elif t.startswith("APPENDIX: REPRODUCIBILITY"):
            blob_app_idx = i
            
    print(f"Indices -> draft_app_start: {draft_app_start}, ack_idx: {ack_idx}, ref_idx: {ref_idx}, blob_app_idx: {blob_app_idx}")
    
    # Remove messy draft appendix before acknowledgment (if draft_app_start < ack_idx)
    if draft_app_start != -1 and ack_idx != -1 and draft_app_start < ack_idx:
        print(f"Deleting draft appendix paragraphs from {draft_app_start} to {ack_idx-1}")
        for _ in range(ack_idx - draft_app_start):
            p = doc.paragraphs[draft_app_start]
            p._element.getparent().remove(p._element)
            
    # Re-evaluate indices after deletion
    p_texts = [p.text.strip() for p in doc.paragraphs]
    blob_app_idx = -1
    for i, t in enumerate(p_texts):
        if t.startswith("APPENDIX: REPRODUCIBILITY"):
            blob_app_idx = i
            break
            
    if blob_app_idx != -1:
        print(f"Removing unformatted blob at index {blob_app_idx}")
        p = doc.paragraphs[blob_app_idx]
        p._element.getparent().remove(p._element)
        
    # Also clean up any loose '---' paragraphs
    for p in list(doc.paragraphs):
        if p.text.strip() == "---":
            p._element.getparent().remove(p._element)

    # Add structured Appendices at the end of the document
    add_structured_appendices(doc)
    
    doc.save(doc_path)
    print(f"Successfully saved structured manuscript to: {doc_path}")

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125) # Dark Blue
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = RGBColor(59, 89, 152)
    return p

def add_paragraph(doc, text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Calibri"
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(38, 38, 38)
    r_text = p.add_run(text)
    r_text.font.name = "Calibri"
    r_text.font.size = Pt(10)
    r_text.font.italic = italic
    r_text.font.color.rgb = RGBColor(51, 51, 51)
    return p

def add_structured_appendices(doc):
    # Divider Rule
    add_paragraph(doc, "")
    
    # APPENDIX A
    add_heading_1(doc, "APPENDIX A: REPRODUCIBILITY & SYSTEM SPECIFICATIONS")
    
    add_heading_2(doc, "A.1 Reproducibility & System Environment Matrix")
    add_paragraph(doc, "To ensure 100% scientific reproducibility across independent compute environments, Table A.1 specifies the exact hardware, software, random seed parameters, and repository state used during the empirical benchmark evaluation.")
    
    # Table A.1
    t1 = doc.add_table(rows=1, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t1.rows[0]
    hdr.cells[0].text = "Reproducibility Parameter"
    hdr.cells[1].text = "Verified System Configuration / Value"
    style_table_header(hdr)
    
    env_data = [
        ("Repository & Commit Hash", "https://github.com/aashishrajput9838/academicuniverse.git (Commit: 88140d1)"),
        ("Benchmark Suite Version", "AU DIC Benchmark v1.0 (Release Candidate 1 - RC1)"),
        ("Evaluated Neural Engine", "Groq Cloud AI (llama-3.1-8b-instant, temp=0.2, max_tokens=8192)"),
        ("Master Random Seed", "SeedManager deterministic seed = 42"),
        ("Operating System & Arch", "Windows 11 Professional (x86_64)"),
        ("Compute Infrastructure", "Intel Core i7 (HP EliteBook 840 G8) | 16 GB DDR4 RAM (CPU-only inference)"),
        ("Runtime Environment", "Python 3.14.x | Node.js v18.x | npm v9.x"),
        ("Statistical Libraries", "scipy >= 1.11 | pandas >= 2.0 | numpy >= 1.24"),
        ("Dataset SHA-256 Hash", "17c136ef76dd0f82"),
        ("Execution Timestamp & Log", "2026-08-05T20:50:48.067Z (backend/benchmark_reports/run_1785959173886/)"),
        ("Dataset Licensing & Access", "MIT License (Open Research Access)")
    ]
    
    for param, val in env_data:
        row = t1.add_row()
        row.cells[0].text = param
        row.cells[1].text = val
    style_table_cells(t1)
    
    add_heading_2(doc, "A.2 Technical Clarifications & Reviewer Inquiries")
    
    add_paragraph(doc, "How does CanonicalNormalizer handle unexpected date or numeric strings that do not match standard regex patterns?", bold_prefix="Inquiry A.2.1: ")
    add_paragraph(doc, "When a field value does not match defined date or numerical patterns, CanonicalNormalizer falls back to StringNormalizer.normalize(val, true) [25]. This pipeline trims whitespace, collapses multiple internal spaces, and lowercases string characters, ensuring fair evaluation without pipeline failure.", bold_prefix="Answer: ")
    
    add_paragraph(doc, "How does the framework scale when evaluating datasets exceeding 10,000+ specimens?", bold_prefix="Inquiry A.2.2: ")
    add_paragraph(doc, "BenchmarkRunner implements O(N) linear dataset processing with worker pool concurrency (concurrency: 4) and automatic checkpointing (checkpoint.json) [31]. In large-scale evaluation runs, BenchmarkRunner updates checkpoint.json after every batch increment. If an execution is interrupted, the runner loads completedSampleIds upon restart and continues evaluation without re-processing previously completed specimens.", bold_prefix="Answer: ")
    
    add_paragraph(doc, "If candidate values still differ after normalization, why is NORMALIZATION_ERROR categorized separately rather than treated as a generic field mismatch?", bold_prefix="Inquiry A.2.3: ")
    add_paragraph(doc, "The purpose of NORMALIZATION_ERROR is to prevent semantic mismatches from being incorrectly attributed to superficial formatting variations [25], [28], [37]. By evaluating only canonical representations, the benchmark distinguishes genuine information extraction failures from benign representation differences.", bold_prefix="Answer: ")

    # APPENDIX B
    add_heading_1(doc, "APPENDIX B: FIELD SPECIFICATION & OBSERVATION COUNT DERIVATION")
    
    add_heading_2(doc, "B.1 Document Category Field Structure")
    add_paragraph(doc, "The benchmark evaluates 360 document specimens partitioned equally across 3 document categories (120 Certificates, 120 Marksheets, and 120 Student ID Cards). Every specimen shares 18 canonical identity/metadata fields:")
    add_paragraph(doc, "studentName, rollNumber, enrollmentNumber, degreeName, branchName, batchYears, cgpa, issueDate, documentType, universityName, universityCode, universityTagline, fatherName, motherName, dateOfBirth, email, phone, bloodGroup.", bold_prefix="Identity / Metadata Fields (18): ")
    
    add_heading_2(doc, "B.2 Mathematical Derivation of 24,480 Paired Observations")
    add_paragraph(doc, "Subject-level array fields follow the structured format subject[i].{code, credits, grade} for i in 0..N, where N represents the number of course records per document. Table B.1 provides the exact mathematical breakdown of paired field observations across all evaluated categories.")
    
    # Table B.1
    t2 = doc.add_table(rows=1, cols=4)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = t2.rows[0]
    hdr2.cells[0].text = "Document Category"
    hdr2.cells[1].text = "Specimens (N)"
    hdr2.cells[2].text = "Fields / Specimen"
    hdr2.cells[3].text = "Total Paired Field Observations"
    style_table_header(hdr2)
    
    deriv_data = [
        ("Degree Certificate", "120", "33 (18 Metadata + 15 Subject)", "3,960"),
        ("Semester Marksheet", "120", "138 (18 Metadata + 120 Subject)", "16,560"),
        ("Student ID Card", "120", "33 (18 Metadata + 15 Subject)", "3,960"),
        ("Total / Weighted Mean", "360 Specimens", "68.0 Mean Fields", "24,480 Observations")
    ]
    for cat, n, f, tot in deriv_data:
        row = t2.add_row()
        row.cells[0].text = cat
        row.cells[1].text = n
        row.cells[2].text = f
        row.cells[3].text = tot
    style_table_cells(t2)

    # APPENDIX C
    add_heading_1(doc, "APPENDIX C: EMPIRICAL STATISTICAL METHODOLOGY & BENCHMARKS")
    
    add_heading_2(doc, "C.1 Empirical Category Confusion Matrix (360 Specimens)")
    add_paragraph(doc, "Table C.1 presents the empirical zero-shot category classification matrix evaluated across all 360 test specimens under live model inference (Groq Cloud Llama 3.1 8B Instant).")
    
    # Table C.1
    t3 = doc.add_table(rows=1, cols=5)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr3 = t3.rows[0]
    hdr3.cells[0].text = "Ground Truth \\ Predicted"
    hdr3.cells[1].text = "Certificate"
    hdr3.cells[2].text = "Marksheet"
    hdr3.cells[3].text = "Student ID"
    hdr3.cells[4].text = "Total Category Accuracy"
    style_table_header(hdr3)
    
    cm_data = [
        ("Certificate", "120", "0", "0", "120 / 120 (100.00%)"),
        ("Marksheet", "0", "120", "0", "120 / 120 (100.00%)"),
        ("Student ID", "0", "0", "120", "120 / 120 (100.00%)"),
        ("Total Specimens", "120", "120", "120", "360 / 360 (100.00%)")
    ]
    for gt, c, m, sid, acc in cm_data:
        row = t3.add_row()
        row.cells[0].text = gt
        row.cells[1].text = c
        row.cells[2].text = m
        row.cells[3].text = sid
        row.cells[4].text = acc
    style_table_cells(t3)
    
    add_heading_2(doc, "C.2 McNemar Contingency Test & Normalization Rescues")
    add_paragraph(doc, "To evaluate whether semantic canonical normalization produces a statistically significant improvement over un-normalized exact string matching, we perform a McNemar paired contingency test across all 24,480 field observations.")
    
    # Table C.2
    t4 = doc.add_table(rows=1, cols=3)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr4 = t4.rows[0]
    hdr4.cells[0].text = "Paired Matching Outcome"
    hdr4.cells[1].text = "Normalized Match SUCCESS"
    hdr4.cells[2].text = "Normalized Match FAIL"
    style_table_header(hdr4)
    
    mcn_data = [
        ("Raw Match SUCCESS", "a = 2,487 (Concordant Positive)", "c = 0 (Discordant Negative)"),
        ("Raw Match FAIL", "b = 167 (Discordant Rescued)", "d = 21,826 (Concordant Negative)")
    ]
    for row_hdr, col1, col2 in mcn_data:
        row = t4.add_row()
        row.cells[0].text = row_hdr
        row.cells[1].text = col1
        row.cells[2].text = col2
    style_table_cells(t4)
    
    add_paragraph(doc, "McNemar Test Equation (Continuity-Corrected):", bold_prefix="Statistical Test Formula: ")
    add_paragraph(doc, "chi^2 = (|b - c| - 1)^2 / (b + c) = (|167 - 0| - 1)^2 / (167 + 0) = (166)^2 / 167 = 165.01 (p < 0.0001)")
    add_paragraph(doc, "The test confirms that canonical normalization yields a statistically significant improvement (p < 0.0001). Canonical normalization rescued 167 field observations (0.68%) from false-negative formatting mismatches with zero regression (c = 0).")
    
    add_heading_2(doc, "C.3 Non-Parametric Bootstrap Confidence Intervals (B = 10,000)")
    add_paragraph(doc, "Table C.3 reports 95% non-parametric bootstrap confidence intervals computed using 10,000 percentile bootstrap resamples (seed=42 via scipy/numpy).")
    
    # Table C.3
    t5 = doc.add_table(rows=1, cols=4)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr5 = t5.rows[0]
    hdr5.cells[0].text = "Evaluation Metric"
    hdr5.cells[1].text = "Point Estimate"
    hdr5.cells[2].text = "95% Bootstrap CI (Lower, Upper)"
    hdr5.cells[3].text = "Resample Count (B)"
    style_table_header(hdr5)
    
    boot_data = [
        ("Raw Exact Match Rate", "10.16%", "[9.78%, 10.54%]", "10,000"),
        ("Normalized Field Match Rate", "10.84%", "[10.45%, 11.24%]", "10,000"),
        ("Category Classification Accuracy", "100.00%", "[100.00%, 100.00%]", "10,000")
    ]
    for m, pe, ci, b_cnt in boot_data:
        row = t5.add_row()
        row.cells[0].text = m
        row.cells[1].text = pe
        row.cells[2].text = ci
        row.cells[3].text = b_cnt
    style_table_cells(t5)

if __name__ == "__main__":
    docs = [
        r"docs\paper\PaperV4_Final_Submission.docx",
        r"docs\paper\PaperV4_Final_Submission.docx.tmp"
    ]
    for d in docs:
        if os.path.exists(d):
            format_document(d)
