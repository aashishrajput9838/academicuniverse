import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" VERIFYING SECTION 5.9 INTEGRATION IN PaperV5_Ollama_Primary.docx")
print("============================================================")

p0 = doc.paragraphs[0]
has_logo = "<w:drawing>" in p0._p.xml or "graphic" in p0._p.xml

s59_p = None
s58_p = None
s6_p = None
toc_has_s59 = False

for i, p in enumerate(doc.paragraphs):
    text_str = p.text.strip()
    if text_str.startswith("5.8 Error Taxonomy"):
        s58_p = p
        s58_idx = i
    elif text_str.startswith("5.9 Classical Machine Learning"):
        if i < 40:
            toc_has_s59 = True
        else:
            s59_p = p
            s59_idx = i
    elif text_str.startswith("6. Discussion & Threats"):
        s6_p = p
        s6_idx = i

assert s58_p is not None, "Error: Section 5.8 missing!"
assert s59_p is not None, "Error: Section 5.9 missing in body!"
assert s6_p is not None, "Error: Section 6 missing!"

print(f"1. Logo on Page 1 (P0 Drawing XML):             {has_logo}")
print(f"2. TOC contains Section 5.9 Entry:              {toc_has_s59}")
print(f"3. Section 5.8 Paragraph Index:                P{s58_idx}")
print(f"4. Section 5.9 Paragraph Index:                P{s59_idx}")
print(f"5. Section 6.0 Paragraph Index:                P{s6_idx}")

order_ok = (s58_idx < s59_idx < s6_idx)
print(f"6. Order Verification (5.8 < 5.9 < 6.0):        {order_ok}")

# Verify Table IX
has_tbl_ix = False
for t in doc.tables:
    if len(t.columns) == 7 and "RF 60:40" in t.rows[0].cells[1].text:
        has_tbl_ix = True
        break

print(f"7. Table IX (7-column ML comparison) present:   {has_tbl_ix}")

# Verify Drawings (Images)
drawings_count = sum(1 for p in doc.paragraphs if "<w:drawing>" in p._p.xml)
print(f"8. Total Figure Drawings in Document Body:      {drawings_count}")

print("============================================================")
assert order_ok, "Section ordering incorrect!"
assert toc_has_s59, "TOC missing Section 5.9 entry!"
assert has_tbl_ix, "Table IX missing!"
assert drawings_count >= 10, "Figures missing!"

print("EMPIRICAL VERIFICATION SUCCESSFUL: Section 5.9 integrated perfectly!")
