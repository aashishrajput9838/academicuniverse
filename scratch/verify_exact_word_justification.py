import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx)

justified_count = 0
center_count = 0
left_count = 0

for p in doc.paragraphs:
    if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        justified_count += 1
    elif p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        center_count += 1
    elif p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
        left_count += 1

print("============================================================")
print(" VERIFYING EXACT WORD JUSTIFY PARAGRAPH ALIGNMENT")
print("============================================================")
print(f"Total Paragraphs:                 {len(doc.paragraphs)}")
print(f"Fully Justified Body Paragraphs:  {justified_count} (WD_ALIGN_PARAGRAPH.JUSTIFY)")
print(f"Centered Headings/Titles/Captions:{center_count}")
print(f"Tables Count:                     {len(doc.tables)}")
print(f"Inline Shape Diagrams:            {len(doc.inline_shapes)}")
print("============================================================")
