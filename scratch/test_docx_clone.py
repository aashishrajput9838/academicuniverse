import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_Clone.docx"

print(f"=== TESTING DIRECT DOCX CLONING FROM {v4_docx_path.name} ===")

assert v4_docx_path.exists(), "PaperV4_Final_Submission.docx missing!"

doc = docx.Document(v4_docx_path)

print(f"Loaded {v4_docx_path.name}:")
print(f"  - Paragraphs count: {len(doc.paragraphs)}")
print(f"  - Tables count:     {len(doc.tables)}")
print(f"  - Sections count:   {len(doc.sections)}")
print(f"  - Inline Shapes:    {len(doc.inline_shapes)}")

# Inspect first few paragraphs
for i, p in enumerate(doc.paragraphs[:10]):
    print(f"    P{i} [{p.style.name}]: {p.text[:80]}")

# Inspect tables
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.cell(0, 0).text.replace('\n', ' ')[:40] if rows > 0 and cols > 0 else ""
    print(f"    Table {i+1}: {rows}x{cols} | Header: '{first_cell}'")

doc.save(test_docx_path)
print(f"[SUCCESS] Saved cloned test docx to: {test_docx_path.name}")
