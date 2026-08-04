"""
Fix & Preserve All 7 Display Equations + Inline Math in IEEE Paper
===================================================================
1. Restores the missing Category Accuracy equation in Section VI.B.1.
2. Formats all 7 display equations into native Word OMML equation objects.
3. Cleans all inline math expressions (P, R, F1, sigma, N=360, etc.) without LaTeX symbols.
4. Saves to Paper_V3_IEEE_Final_Build.docx and copies to Paper_V3_IEEE_Final.docx.
"""
import os
import sys
import io
import subprocess
import time
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# Kill Word processes to prevent locks
subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], capture_output=True)
time.sleep(2)

INPUT_DOCX = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_v2.docx"
BUILD_DOCX = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final_Build.docx"
FINAL_DOCX = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final.docx"
FINAL_PDF = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final.pdf"

SECTION_MAP = {
    "1. Introduction": "I. INTRODUCTION",
    "2. Related Work": "II. RELATED WORK",
    "3. System Architecture Overview": "III. SYSTEM ARCHITECTURE OVERVIEW",
    "4. ADBG Synthetic Data Generation Methodology": "IV. ADBG SYNTHETIC DATA GENERATION METHODOLOGY",
    "5. AU DIC Evaluation Subsystem & Semantic Normalization Layer": "V. AU DIC EVALUATION SUBSYSTEM & SEMANTIC NORMALIZATION LAYER",
    "6. Experimental Setup & Evaluation Metrics": "VI. EXPERIMENTAL SETUP & EVALUATION METRICS",
    "6. Experimental Setup, Protocol & Metrics": "VI. EXPERIMENTAL SETUP & EVALUATION METRICS",
    "7. Results & Validation": "VII. RESULTS & VALIDATION",
    "8. Discussion, Threats to Validity & Limitations": "VIII. DISCUSSION, THREATS TO VALIDITY & LIMITATIONS",
    "9. Future Work": "IX. FUTURE WORK",
    "10. Conclusion": "X. CONCLUSION",
}

SUBSECTION_MAP = {
    "1.1": "A.", "1.2": "B.", "1.3": "C.", "1.4": "D.",
    "2.1": "A.", "2.2": "B.", "2.3": "C.",
    "4.1": "A.", "4.2": "B.", "4.3": "C.",
    "5.1": "A.", "5.2": "B.", "5.3": "C.",
    "6.1": "A.", "6.2": "B.",
    "7.1": "A.", "7.2": "B.", "7.3": "C.", "7.4": "D.",
    "8.1": "A.", "8.2": "B.", "8.3": "C.",
}

# Clean inline replacements
INLINE_CLEAN = [
    (r"$3 \text{ categories} \times 30 \text{ instances} \times 4 \text{ degradation profiles}$", "3 categories \u00d7 30 instances \u00d7 4 degradation profiles"),
    (r"MIT $\rightarrow$ Massachusetts Institute of Technology", "MIT \u2192 Massachusetts Institute of Technology"),
    (r"$N=360$", "N = 360"),
    (r"(\hat{C}_i)", "(\u0108_i)"),
    (r"(C_i)", "(C_i)"),
    (r"($P$)", "(P)"),
    (r"($R$)", "(R)"),
    (r"($F_1$)", "(F\u2081)"),
    (r"($L_{\text{GT}}$)", "(L_GT)"),
    (r"($W_{\text{GT}}$)", "(W_GT)"),
    (r"($\text{ms/sample}$)", "(ms/sample)"),
    (r"($\text{samples/sec}$)", "(samples/sec)"),
    (r"($\sigma = 0.45\text{ ms}$)", "(\u03c3 = 0.45 ms)"),
    (r"\sigma = 0.45\text{ ms}", "\u03c3 = 0.45 ms"),
    (r"$O(N)$", "O(N)"),
    (r"\Denotes", "*Denotes"),
    (r"_{", "_"),
    (r"}", ""),
]

def main():
    print(f"[1/4] Loading template: {INPUT_DOCX}")
    doc = Document(INPUT_DOCX)

    # 1. Page Margins & Header/Footer
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

        # Footer page number
        footer = s.footer
        for p in footer.paragraphs:
            p.clear()
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run()
        frun.font.name = "Times New Roman"
        frun.font.size = Pt(10)
        frun.font.color.rgb = RGBColor(80, 80, 80)
        fld_xml = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        fp._p.append(fld_xml)

        header = s.header
        for p in header.paragraphs:
            p.clear()

    # 2. Iterate paragraphs for headings & inline cleaning
    print("[2/4] Formatting Headings & Cleaning Inline Math...")
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()

        # Title
        if i == 0:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(18)
                r.bold = True
                r.font.color.rgb = RGBColor(0, 51, 102)
            para.paragraph_format.space_before = Pt(36)
            para.paragraph_format.space_after = Pt(12)

        # Author block
        elif i == 1 and "Authors" in txt:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.clear()
            r = para.runs[0] if para.runs else para.add_run()
            r.text = (
                "Author Name¹, Co-Author Name²\n"
                "¹Department of Computer Science, University Name, City, Country\n"
                "²Department of Information Technology, University Name, City, Country\n"
                "Corresponding Email: author@university.edu"
            )
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(51, 51, 51)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

        elif i == 2 and "Target Publication" in txt:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.clear()
            r = para.runs[0] if para.runs else para.add_run()
            r.text = "Manuscript submitted for review to IEEE Access"
            r.font.name = "Times New Roman"
            r.font.size = Pt(9)
            r.italic = True
            r.font.color.rgb = RGBColor(100, 100, 100)
            para.paragraph_format.space_after = Pt(4)

        elif i == 3 and "Repository" in txt:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in para.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(120, 120, 120)
            para.paragraph_format.space_after = Pt(12)

        # Major Headings
        for old_h, new_h in SECTION_MAP.items():
            if txt == old_h:
                for r in para.runs:
                    r.clear()
                r = para.runs[0] if para.runs else para.add_run()
                r.text = new_h
                r.font.name = "Times New Roman"
                r.font.size = Pt(13)
                r.bold = True
                r.font.color.rgb = RGBColor(0, 51, 102)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(16)
                para.paragraph_format.space_after = Pt(8)
                para.paragraph_format.keep_with_next = True
                break

        # Subsections
        import re
        sub_m = re.match(r'^(\d+\.\d+)\s+(.+)$', txt)
        if sub_m:
            sub_num = sub_m.group(1)
            sub_title = sub_m.group(2)
            if sub_num in SUBSECTION_MAP:
                new_t = f"{SUBSECTION_MAP[sub_num]} {sub_title}"
                for r in para.runs:
                    r.clear()
                r = para.runs[0] if para.runs else para.add_run()
                r.text = new_t
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = True
                r.italic = True
                r.font.color.rgb = RGBColor(51, 51, 51)
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.keep_with_next = True

        # Inline Cleaning
        for s_pat, r_str in INLINE_CLEAN:
            if s_pat in para.text:
                para.text = para.text.replace(s_pat, r_str)
                for r in para.runs:
                    r.font.name = "Times New Roman"

        # General body formatting
        if para.style and para.style.name == "Normal" and txt and not any([
            txt.startswith("I."), txt.startswith("II."), txt.startswith("III."), txt.startswith("IV."),
            txt.startswith("V."), txt.startswith("VI."), txt.startswith("VII."), txt.startswith("VIII."),
            txt.startswith("IX."), txt.startswith("X."), txt.startswith("A."), txt.startswith("B."),
            txt.startswith("Table "), txt.startswith("Fig. "), txt.startswith("---")
        ]):
            for r in para.runs:
                if not r.font.name:
                    r.font.name = "Times New Roman"
                if not r.font.size:
                    r.font.size = Pt(10)
            para.paragraph_format.line_spacing = 1.15

    # 3. Tables formatting
    print("[3/4] Formatting Tables...")
    for t in doc.tables:
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl = t._tbl
        tblPr = tbl.tblPr
        for eb in tblPr.findall(f'{{{nsdecls("w").split(chr(34))[1]}}}tblBorders'):
            tblPr.remove(eb)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:space="0" w:color="003366"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="003366"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        for r_idx, row in enumerate(t.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = "Times New Roman"
                        r.font.size = Pt(9)
                        if r_idx == 0:
                            r.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                tcPr = cell._tc.get_or_add_tcPr()
                fc = "003366" if r_idx == 0 else ("F5F7FA" if r_idx % 2 == 0 else "FFFFFF")
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fc}"/>')
                tcPr.append(shd)

    # Save to build file
    doc.save(BUILD_DOCX)
    print(f"  [SAVED] {BUILD_DOCX}")

    # 4. Native OMML Equations via Word COM
    print("[4/4] Inserting Native OMML Equations via Word COM...")
    
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0

    com_doc = word.Documents.Open(os.path.abspath(BUILD_DOCX))

    # Explicit list of all 7 display equations and their exact preceding anchor text
    EQUATION_TARGETS = [
        # Eq 1: DocumentSpecimen (Section IV.A)
        ("A. Seed-Deterministic Profile Generation",
         "DocumentSpecimen = G(Seed, Category, Profile)"),

        # Eq 2: Degradation Pipeline (Section IV.C)
        ("C. Optical Degradation Pipeline",
         "I_degraded = D_rotation \u2218 D_contrast \u2218 D_gaussian \u2218 D_blur (I_clean)"),

        # Eq 3: Error Taxonomy (Section V.C)
        ("C. Nine-Class Structured OCR Error Taxonomy",
         "ErrorCategory \u2208 {OCR_ERROR, FIELD_MISSING, HALLUCINATION, FORMAT_ERROR, NORMALIZATION_ERROR, PARTIAL_MATCH, LOW_CONFIDENCE, CATEGORY_ERROR, EXACT_MATCH}"),

        # Eq 4: Category Classification Accuracy (PRESERVED & RESTORED!)
        ("1. Category Classification Accuracy",
         "Category Accuracy = ( \u2211_(i=1)^N \u2141(C\u0302_i = C_i) ) / N"),

        # Eq 5: Field Extraction Precision, Recall, F1
        ("2. Field Extraction Precision",
         "P = (True Positive Fields) / (True Positive Fields + False Positive Fields),   R = (True Positive Fields) / (True Positive Fields + False Negative Fields),   F_1 = 2 \u22c5 (P \u22c5 R) / (P + R)"),

        # Eq 6: CER
        ("3. Character Error Rate",
         "CER = (S_char + D_char + I_char) / L_GT"),

        # Eq 7: WER
        ("4. Word Error Rate",
         "WER = (S_word + D_word + I_word) / W_GT"),
    ]

    for anchor_text, linear_eq in EQUATION_TARGETS:
        anchor_found = False
        for p_idx in range(1, com_doc.Paragraphs.Count + 1):
            p = com_doc.Paragraphs(p_idx)
            if anchor_text in p.Range.Text:
                anchor_found = True
                # Insert a new paragraph cleanly right after anchor
                p.Range.InsertParagraphAfter()
                eq_p = com_doc.Paragraphs(p_idx + 1)
                eq_p.Range.Text = linear_eq
                r = eq_p.Range
                r.End = r.End - 1
                doc_om = com_doc.OMaths.Add(r)
                eq_p.Alignment = 1  # wdAlignParagraphCenter
                eq_p.Format.SpaceBefore = 8
                eq_p.Format.SpaceAfter = 8
                print(f"  [OMML INSERTED] '{anchor_text[:35]}...' -> Native OMML Equation")
                break
        if not anchor_found:
            print(f"  [WARN] Anchor text '{anchor_text}' not found!")

    # BuildUp all equations
    print("\n  [OMML BUILDUP] Converting linear math to 2D built-up OMML...")
    for i in range(1, com_doc.OMaths.Count + 1):
        try:
            com_doc.OMaths(i).BuildUp()
        except Exception as e:
            print(f"    BuildUp warning Eq {i}: {e}")

    com_doc.Save()
    print(f"  [OMML SAVE] {BUILD_DOCX} saved with native OMML equations.")

    com_doc.Close(False)
    try:
        word.Quit()
    except:
        pass

    omml_target = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final_OMML.docx"
    report_target = r"c:\github\academicuniverse.com\academicuniverse\docs\reports\Paper_V3_IEEE_Final_OMML.docx"
    
    shutil.copy(BUILD_DOCX, omml_target)
    shutil.copy(BUILD_DOCX, report_target)
    print(f"\n[SUCCESS] {omml_target} generated with native OMML equations!")

    try:
        shutil.copy(BUILD_DOCX, FINAL_DOCX)
        shutil.copy(BUILD_DOCX, r"c:\github\academicuniverse.com\academicuniverse\docs\reports\Paper_V3_IEEE_Final.docx")
        print(f"[SUCCESS] {FINAL_DOCX} updated!")
    except Exception as e:
        print(f"[NOTE] Please close Word to allow updating {FINAL_DOCX} directly ({e})")

if __name__ == "__main__":
    main()
