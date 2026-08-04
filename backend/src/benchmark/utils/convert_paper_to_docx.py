import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re
import os
import win32com.client

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_ieee_docx(md_path, docx_path):
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add Page Number to Footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Page ")
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # XML for Page Number
        f_xml = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        f_p._p.append(f_xml)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    def format_inline_markdown(paragraph, text):
        # Format bold, italic, LaTeX inline math, code
        # Tokenize by bold, italic, code, math
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$.*?\$)', text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                r = paragraph.add_run(token[2:-2])
                r.bold = True
            elif token.startswith('*') and token.endswith('*'):
                r = paragraph.add_run(token[1:-1])
                r.italic = True
            elif token.startswith('`') and token.endswith('`'):
                r = paragraph.add_run(token[1:-1])
                r.font.name = 'Consolas'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(180, 40, 40)
            elif token.startswith('$') and token.endswith('$'):
                r = paragraph.add_run(token[1:-1])
                r.font.name = 'Cambria Math'
                r.italic = True
                r.font.color.rgb = RGBColor(0, 51, 102)
            else:
                paragraph.add_run(token)

    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                
                table = doc.add_table(rows=1, cols=1)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = table.cell(0, 0)
                set_cell_background(cell, "F4F6F9")
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after = Pt(2)
                c_run = cp.add_run('\n'.join(code_lines))
                c_run.font.name = 'Consolas'
                c_run.font.size = Pt(9)
                c_run.font.color.rgb = RGBColor(30, 30, 30)
                
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and ('---' in line or i + 1 < len(lines) and '|' in lines[i+1]):
            # Start collecting table lines
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\r\n'))
                i += 1
            
            # Process Markdown table
            rows_data = []
            caption_text = ""
            for tline in table_lines:
                if '---' in tline:
                    continue
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)
            
            if rows_data:
                num_rows = len(rows_data)
                num_cols = len(rows_data[0])
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Style table XML for clean borders
                tblPr = table._tbl.tblPr
                borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="8" w:space="0" w:color="003366"/><w:bottom w:val="single" w:sz="8" w:space="0" w:color="003366"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/><w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
                tblPr.append(borders)

                for r_idx, row in enumerate(rows_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
                            cp = cell.paragraphs[0]
                            cp.paragraph_format.space_before = Pt(2)
                            cp.paragraph_format.space_after = Pt(2)
                            
                            if r_idx == 0:
                                set_cell_background(cell, "003366")
                                format_inline_markdown(cp, val)
                                for run in cp.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(9.5)
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, "FFFFFF")
                                else:
                                    set_cell_background(cell, "F8F9FA")
                                format_inline_markdown(cp, val)
                                for run in cp.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(9)
                                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line[2:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(20)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[3:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[4:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True
            r.font.color.rgb = RGBColor(51, 51, 51)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[5:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)
            r.bold = True
            r.italic = True
        # Blockquotes / Alerts
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            format_inline_markdown(p, line[2:].strip())
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = RGBColor(70, 70, 70)
        # Bullet list items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            format_inline_markdown(p, line[2:].strip())
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)
        # Numbered list items
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            content = re.sub(r'^\d+\.\s', '', line)
            format_inline_markdown(p, content)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10.5)
        # Display Math
        elif line.startswith('$$') and line.endswith('$$'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[2:-2].strip())
            r.font.name = 'Cambria Math'
            r.font.size = Pt(11)
            r.italic = True
            r.font.color.rgb = RGBColor(0, 51, 102)
        # Horizontal Rule
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run('_________________________________________________________________________________')
            r.font.color.rgb = RGBColor(200, 200, 200)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Table captions / Bold Table titles
        elif line.startswith('**Table ') or line.startswith('**Figure '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            format_inline_markdown(p, line.strip())
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
        # Standard Body Paragraphs
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            format_inline_markdown(p, line.strip())
            for run in p.runs:
                if not run.font.name:
                    run.font.name = 'Times New Roman'
                if not run.font.size:
                    run.font.size = Pt(10.5)

        i += 1

    doc.save(docx_path)
    print(f"Docx generated successfully: {docx_path}")

def convert_docx_to_pdf(docx_path, pdf_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17) # 17 = wdFormatPDF
        doc.Close()
        print(f"PDF generated successfully: {pdf_path}")
    except Exception as e:
        print(f"PDF generation error: {e}")
    finally:
        word.Quit()

if __name__ == "__main__":
    md_file = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\bb9b3069-0e60-4209-b2b8-d0321ac491db\Paper_V3.md"
    docx_file = r"c:\github\academicuniverse.com\academicuniverse\Paper_V3_IEEE.docx"
    pdf_file = r"c:\github\academicuniverse.com\academicuniverse\Paper_V3_IEEE.pdf"
    
    create_ieee_docx(md_file, docx_file)
    convert_docx_to_pdf(docx_file, pdf_file)
