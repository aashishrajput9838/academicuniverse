import docx
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" VERIFYING TOC NEW PAGE BREAK & RIGHT TAB STOP ALIGNMENT")
print("============================================================")

contents_idx = None
intro_idx = None

for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "CONTENTS":
        contents_idx = i
    elif p.text.strip() == "1. Introduction":
        intro_idx = i

assert contents_idx is not None, "Error: CONTENTS heading not found!"
assert intro_idx is not None, "Error: 1. Introduction heading not found!"

contents_p = doc.paragraphs[contents_idx]

# 1. Page Break Before Check
has_page_break_before = (contents_p.paragraph_format.page_break_before is True or "pageBreakBefore" in contents_p._p.xml)
print(f"1. Page Break Before CONTENTS:         {has_page_break_before}")

# 2. Right Tab Stop & Dot Leader Check across ALL TOC entries
toc_entries = []
tab_stop_ok_count = 0

for i in range(contents_idx + 1, intro_idx):
    p = doc.paragraphs[i]
    if "\t" in p.text:
        toc_entries.append(p)
        for ts in p.paragraph_format.tab_stops:
            if abs(ts.position.inches - 6.5) < 0.05 and ts.alignment == WD_TAB_ALIGNMENT.RIGHT and ts.leader == WD_TAB_LEADER.DOTS:
                tab_stop_ok_count += 1

print(f"2. Total TOC Entry Paragraphs Found:   {len(toc_entries)}")
print(f"3. Right Tab Stops at 6.5 in w/ DOTS:  {tab_stop_ok_count}")

print("\nSample Formatted TOC Entries from DOCX:")
for p in toc_entries[:5]:
    title_part, page_part = p.text.split("\t")
    print(f"   [Title]: '{title_part[:40]:<40}' -> [Tab + Leader] -> [Page]: '{page_part}'")

print("============================================================")
assert has_page_break_before, "Page break before CONTENTS missing!"
assert len(toc_entries) == 49, f"Expected 49 TOC entries, found {len(toc_entries)}"
assert tab_stop_ok_count == 49, f"Expected 49 right tab stops, found {tab_stop_ok_count}"

print("ALL TOC PAGINATION & TAB-STOP ALIGNMENT CHECKS PASSED 100%!")
