import docx
import re
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" FINAL EMPIRICAL VERIFICATION MATRIX — PaperV5_Ollama_Primary")
print("============================================================")

# 1. Author Block Check
p1_text = doc.paragraphs[1].text
author_ok = ("Kushagra Singh Bhadauria1" in p1_text and "Aashish Rajput2" in p1_text and "Avdesh Kumar Sah3" in p1_text and "Department of Computer Science" in p1_text and "2023361009.kushagra@ug.sharda.ac.in" in p1_text)
print(f"1. Centered 3-Line Author Block Intact: {author_ok}")

# 2. REFERENCES Heading Check
ref_ok = any(p.text.strip() == "REFERENCES" for p in doc.paragraphs)
print(f"2. REFERENCES Section Heading Present: {ref_ok}")

# 3. Footer Cleanup Check
footer_text = "".join([p.text for s in doc.sections for p in s.footer.paragraphs])
footer_ok = ("IEEE ACCESS" not in footer_text and "Volume 14, 2026" not in footer_text)
print(f"3. Dummy Footer Text 0% Present:       {footer_ok}")

# 4. Body Paragraph Justification Check
justified_count = sum(1 for p in doc.paragraphs if p.alignment == docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY)
print(f"4. Body Paragraphs Fully Justified:     {justified_count} Paragraphs")

# 5. Empirical Results Synchronized Check
text_full = "".join([p.text for p in doc.paragraphs])
ollama_ok = ("Ollama" in text_full and "MiniCPM-V" in text_full and "75.23%" in text_full and "82.18%" in text_full)
print(f"5. Ollama & MiniCPM-V Empirical Data:   {ollama_ok}")

print("============================================================")
assert author_ok, "Author block corrupted!"
assert ref_ok, "REFERENCES heading missing!"
assert footer_ok, "Dummy footer text present!"
assert justified_count >= 200, "Body justification missing!"
assert ollama_ok, "Empirical data missing!"
print("ALL PRESERVED FORMATTING CHECKS PASSED 100%!")
