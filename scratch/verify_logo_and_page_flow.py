import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" VERIFYING LOGO PAGE 1, CONTENTS PAGE 2, FRONT MATTER PAGE 3")
print("============================================================")

p0 = doc.paragraphs[0]
has_drawing = "<w:drawing>" in p0._p.xml or "graphic" in p0._p.xml

contents_p = None
title_p = None
intro_p = None

for i, p in enumerate(doc.paragraphs):
    text_str = p.text.strip()
    if text_str == "CONTENTS":
        contents_p = p
        contents_idx = i
    elif text_str.startswith("ADBG v1.0 & AU DIC Benchmark"):
        title_p = p
        title_idx = i
    elif text_str == "1. Introduction":
        intro_p = p
        intro_idx = i

assert contents_p is not None, "CONTENTS paragraph missing!"
assert title_p is not None, "Title paragraph missing!"
assert intro_p is not None, "1. Introduction paragraph missing!"

print(f"1. Paragraph P0 Image Drawing XML Found:           {has_drawing} (PAGE 1 LOGO)")
print(f"2. CONTENTS Paragraph P{contents_idx} (Page Break Before: {contents_p.paragraph_format.page_break_before}) -> PAGE 2")
print(f"3. Title Paragraph P{title_idx} (Page Break Before:    {title_p.paragraph_format.page_break_before}) -> PAGE 3")
print(f"4. 1. Intro Paragraph P{intro_idx} (Page Break Before:    {intro_p.paragraph_format.page_break_before}) -> PAGE 4")

# Flow verification
flow_ok = (has_drawing and contents_p.paragraph_format.page_break_before is True and title_p.paragraph_format.page_break_before is True and intro_p.paragraph_format.page_break_before is True)

print("============================================================")
assert flow_ok, "Page flow hierarchy incorrect!"
print("EMPIRICAL VERIFICATION SUCCESSFUL: Logo Page 1, CONTENTS Page 2, Front Matter Page 3, Main Paper Page 4!")
