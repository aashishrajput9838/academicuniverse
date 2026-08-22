import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

contents_idx = None
toc_field_found = False

for i, p in enumerate(doc.paragraphs[:20]):
    if p.text.strip() == "CONTENTS":
        contents_idx = i
    if "TOC" in p._p.xml:
        toc_field_found = True

print("============================================================")
print(" TABLE OF CONTENTS (CONTENTS) VERIFICATION REPORT")
print("============================================================")
print(f"CONTENTS Heading Paragraph Index: P{contents_idx}")
print(f"TOC Field XML Found:              {toc_field_found}")
print(f"Index Terms Precedes CONTENTS:   {doc.paragraphs[contents_idx-1].text.strip().startswith('Index Terms')}")
print(f"CONTENTS Precedes Intro:          {'1. Introduction' in doc.paragraphs[contents_idx+2].text or '1. Introduction' in doc.paragraphs[contents_idx+3].text}")
print("============================================================")

assert contents_idx is not None, "Error: CONTENTS heading missing!"
assert toc_field_found, "Error: Word TOC field missing!"
