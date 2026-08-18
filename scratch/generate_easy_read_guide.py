import os
import sys
import docx
import subprocess
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
easy_read_dir = workspace / "docs" / "paper" / "easy_read"
md_source_path = easy_read_dir / "Easy_Read_Research_Guide.md"
target_docx_path = easy_read_dir / "Easy_Read_Research_Guide.docx"
target_pdf_path = easy_read_dir / "Easy_Read_Research_Guide.pdf"
dt_cm_path = workspace / "results" / "confusion_matrices" / "dt_composite.png"
rf_cm_path = workspace / "results" / "confusion_matrices" / "rf_composite.png"

print("============================================================")
print(" GENERATING EASY-READ RESEARCH GUIDE (DOCX & PDF)")
print("============================================================")

assert md_source_path.exists(), f"Source markdown missing at {md_source_path}"

# Force kill lingering Word/WPS background processes
try:
    subprocess.run(["taskkill", "/F", "/IM", "wps.exe", "/IM", "wpscenter.exe", "/IM", "WINWORD.EXE"], capture_output=True)
except Exception:
    pass

doc = docx.Document()

# Set page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Dark Navy
    r.bold = True
    return p

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x4A, 0x55, 0x68) # Slate Gray
    r.font.italic = True
    return p

def add_heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    r.bold = True
    return p

def add_heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) # Royal Blue
    r.bold = True
    return p

def add_heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    r.bold = True
    return p

def add_body(text, bold_prefix=None, italic=False, justify=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = "Segoe UI"
        r_pre.font.size = Pt(10.5)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
        
    r = p.add_run(text)
    r.font.name = "Segoe UI"
    r.font.size = Pt(10.5)
    r.font.italic = italic
    r.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return p

def add_callout(text, title=None, bg_color="F0F4F8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)
    
    # Set background color
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), bg_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.1)
    p.paragraph_format.right_indent = Inches(0.1)
    
    if title:
        r_t = p.add_run(title + "\n")
        r_t.font.name = "Segoe UI"
        r_t.font.size = Pt(11)
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        
    r_b = p.add_run(text)
    r_b.font.name = "Segoe UI"
    r_b.font.size = Pt(10)
    r_b.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

# Parse Markdown lines into styled Word document
with open(md_source_path, "r", encoding="utf-8") as f:
    lines = f.readlines()

in_table = False
table_lines = []

for line in lines:
    line_s = line.strip()
    if not line_s:
        continue
        
    if line_s.startswith("# Easy-Read"):
        add_title(line_s.lstrip("#").strip())
    elif line_s.startswith("## Smart Academic"):
        add_subtitle(line_s.lstrip("#").strip())
    elif line_s.startswith("## "):
        add_heading1(line_s.lstrip("#").strip())
    elif line_s.startswith("### "):
        add_heading2(line_s.lstrip("#").strip())
    elif line_s.startswith("#### "):
        add_heading3(line_s.lstrip("#").strip())
    elif line_s.startswith("> 💡 **Think of it like this:**"):
        callout_text = line_s.replace("> 💡 **Think of it like this:**", "").strip()
        add_callout(callout_text, title="💡 Think of it like this:", bg_color="FEFCBF") # Soft Yellow tint
    elif line_s.startswith(">"):
        add_callout(line_s.lstrip(">").strip(), bg_color="EBF8FF") # Soft Blue tint
    elif line_s.startswith("|"):
        # Table row handling
        if not in_table:
            in_table = True
            table_lines = [line_s]
        else:
            table_lines.append(line_s)
    else:
        if in_table:
            in_table = False
            # Render stored table
            rows = [r.split("|")[1:-1] for r in table_lines if "---" not in r]
            if rows:
                num_cols = len(rows[0])
                tbl = doc.add_table(rows=len(rows), cols=num_cols)
                tbl.autofit = True
                for r_idx, r_data in enumerate(rows):
                    for c_idx, val in enumerate(r_data):
                        cell = tbl.cell(r_idx, c_idx)
                        cell.text = val.strip()
                        p_c = cell.paragraphs[0]
                        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                        if p_c.runs:
                            r = p_c.runs[0]
                            r.font.name = "Segoe UI"
                            r.font.size = Pt(9)
                            if r_idx == 0 or c_idx == 0:
                                r.bold = True
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            table_lines = []
            
        # Normal body paragraph
        if line_s.startswith("- **") or line_s.startswith("1. **") or line_s.startswith("2. **") or line_s.startswith("3. **") or line_s.startswith("4. **") or line_s.startswith("5. **") or line_s.startswith("6. **"):
            parts = line_s.split("**", 2)
            if len(parts) >= 3:
                prefix = parts[0] + parts[1] + ": "
                body_txt = parts[2].lstrip(": ").strip()
                add_body(body_txt, bold_prefix=prefix)
            else:
                add_body(line_s)
        else:
            add_body(line_s)

# Embed Composite Confusion Matrix Figures
if dt_cm_path.exists():
    add_heading2("Decision Tree Confusion Matrix Grid (60:40, 70:30, 80:20)")
    p_fig_dt = doc.add_paragraph()
    p_fig_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_fig_dt.add_run()
    r_f.add_picture(str(dt_cm_path), width=Inches(6.2))
    p_cap_dt = doc.add_paragraph()
    p_cap_dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c = p_cap_dt.add_run("Figure E1. Decision Tree Classification Confusion Matrices across Train-Test Splits")
    r_c.font.name = "Segoe UI"
    r_c.font.size = Pt(9)
    r_c.font.italic = True
    p_cap_dt.paragraph_format.space_after = Pt(12)

if rf_cm_path.exists():
    add_heading2("Random Forest Confusion Matrix Grid (60:40, 70:30, 80:20)")
    p_fig_rf = doc.add_paragraph()
    p_fig_rf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f2 = p_fig_rf.add_run()
    r_f2.add_picture(str(rf_cm_path), width=Inches(6.2))
    p_cap_rf = doc.add_paragraph()
    p_cap_rf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_c2 = p_cap_rf.add_run("Figure E2. Random Forest Classification Confusion Matrices across Train-Test Splits")
    r_c2.font.name = "Segoe UI"
    r_c2.font.size = Pt(9)
    r_c2.font.italic = True
    p_cap_rf.paragraph_format.space_after = Pt(12)

# Save Easy_Read_Research_Guide.docx
doc.save(target_docx_path)
print(f"[SUCCESS] Saved Easy-Read Guide Word DOCX: {target_docx_path.relative_to(workspace)}")

# Export Easy_Read_Research_Guide.pdf via Word COM Automation
print(f"Exporting PDF via Word COM Automation: {target_pdf_path.name}...")
abs_docx = os.path.abspath(str(target_docx_path))
abs_pdf = os.path.abspath(str(target_pdf_path))

word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_word = word.Documents.Open(abs_docx)
    doc_word.SaveAs(abs_pdf, FileFormat=17) # wdFormatPDF = 17
    page_count = doc_word.ComputeStatistics(2)
    doc_word.Close()
    print(f"[SUCCESS] Exported Easy-Read Guide PDF: {target_pdf_path.relative_to(workspace)} ({page_count} Pages!)")
except Exception as e:
    print(f"Word COM PDF Export Error: {e}")
finally:
    if word:
        try: word.Quit()
        except: pass
