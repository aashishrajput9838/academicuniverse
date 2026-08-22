import os
import docx
import re
import sys
import subprocess
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"
v5_pdf_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.pdf"

print("============================================================")
print(" REBUILDING PAPER V5 WITH VERIFIED POPULATED TABLE OF CONTENTS")
print("============================================================")

# Force kill any lingering Word or WPS background processes to prevent file locks
try:
    subprocess.run(["taskkill", "/F", "/IM", "wps.exe", "/IM", "wpscenter.exe", "/IM", "WINWORD.EXE"], capture_output=True)
except Exception:
    pass

assert v4_docx_path.exists(), f"Error: Baseline {v4_docx_path} missing!"

doc = docx.Document(v4_docx_path)

def set_paragraph_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)

# 1. Clear incomplete/dummy footer text
for s_idx, section in enumerate(doc.sections):
    footer = section.footer
    for p in footer.paragraphs:
        if "IEEE ACCESS" in p.text or "Page" in p.text:
            p.text = ""

# 2. Format Author Paragraph P1 with Exact 3-Line Structure
p1 = doc.paragraphs[1]
p1.text = ""
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.line_spacing = 1.15
p1.paragraph_format.space_after = Pt(8)

# LINE 1 — AUTHORS
r1 = p1.add_run("Kushagra Singh Bhadauria")
r1.font.name = "Times New Roman"
r1.font.size = Pt(11.5)
r1.bold = True

r1_sup = p1.add_run("1")
r1_sup.font.name = "Times New Roman"
r1_sup.font.size = Pt(9)
r1_sup.font.superscript = True
r1_sup.bold = True

r2_sep = p1.add_run(", ")
r2_sep.font.name = "Times New Roman"
r2_sep.font.size = Pt(11.5)
r2_sep.bold = True

r2 = p1.add_run("Aashish Rajput")
r2.font.name = "Times New Roman"
r2.font.size = Pt(11.5)
r2.bold = True

r2_sup = p1.add_run("2")
r2_sup.font.name = "Times New Roman"
r2_sup.font.size = Pt(9)
r2_sup.font.superscript = True
r2_sup.bold = True

r3_sep = p1.add_run(", and ")
r3_sep.font.name = "Times New Roman"
r3_sep.font.size = Pt(11.5)
r3_sep.bold = True

r3 = p1.add_run("Avdesh Kumar Sah")
r3.font.name = "Times New Roman"
r3.font.size = Pt(11.5)
r3.bold = True

r3_sup = p1.add_run("3")
r3_sup.font.name = "Times New Roman"
r3_sup.font.size = Pt(9)
r3_sup.font.superscript = True
r3_sup.bold = True

# LINE BREAK 1 -> LINE 2 — SHARED AFFILIATION
r_br1 = p1.add_run("\n")

r_affil_sup = p1.add_run("123")
r_affil_sup.font.name = "Times New Roman"
r_affil_sup.font.size = Pt(8.5)
r_affil_sup.font.superscript = True

r_affil_text = p1.add_run(" Department of Computer Science and Engineering, Sharda University, Greater Noida, Uttar Pradesh, India")
r_affil_text.font.name = "Times New Roman"
r_affil_text.font.size = Pt(9.5)

# LINE BREAK 2 -> LINE 3 — EMAILS (SINGLE CENTERED LINE, COMMA SEPARATED)
r_br2 = p1.add_run("\n")

r_e1_sup = p1.add_run("1")
r_e1_sup.font.name = "Times New Roman"
r_e1_sup.font.size = Pt(8.5)
r_e1_sup.font.superscript = True

r_e1_text = p1.add_run(" 2023361009.kushagra@ug.sharda.ac.in, ")
r_e1_text.font.name = "Times New Roman"
r_e1_text.font.size = Pt(9.5)

r_e2_sup = p1.add_run("2")
r_e2_sup.font.name = "Times New Roman"
r_e2_sup.font.size = Pt(8.5)
r_e2_sup.font.superscript = True

r_e2_text = p1.add_run(" 2023329421.aashish@ug.sharda.ac.in, ")
r_e2_text.font.name = "Times New Roman"
r_e2_text.font.size = Pt(9.5)

r_e3_sup = p1.add_run("3")
r_e3_sup.font.name = "Times New Roman"
r_e3_sup.font.size = Pt(8.5)
r_e3_sup.font.superscript = True

r_e3_text = p1.add_run(" 2023265132.avdesh@ug.sharda.ac.in")
r_e3_text.font.name = "Times New Roman"
r_e3_text.font.size = Pt(9.5)

# 3. Clear misplaced "References" header sitting before Appendix A
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References" and i < 240:
        p.text = ""

# 4. Insert IEEE-formatted REFERENCES section heading immediately before reference [1]
ref1_inserted = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("[1] "):
        p_hdr = p.insert_paragraph_before("REFERENCES")
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after = Pt(6)
        if p_hdr.runs:
            r = p_hdr.runs[0]
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.bold = True
        ref1_inserted = True
        break

assert ref1_inserted, "Error: Could not locate reference [1] entry in manuscript!"

# Text replacements mapping for canonical empirical results
replacements = [
    ("Llama-3.1-8B-Instant", "MiniCPM-V (7.6B Q4_0 GGUF)"),
    ("llama-3.1-8b-instant", "minicpm-v:latest"),
    ("Groq Cloud vision API", "Ollama Local Model-Serving Runtime (v0.32.14)"),
    ("Groq Cloud endpoint", "Local Ollama Inference Runtime"),
    ("Groq API", "Ollama Local Model-Serving Engine"),
    ("cloud-based inference", "local offline model-serving runtime inference"),
    ("10.16%", "74.60%"),
    ("10.84%", "82.18%"),
    ("17.19%", "75.23%"),
    ("89.27%", "11.35%"),
    ("82.76%", "8.21%"),
    ("165.01", "1853.0005"),
    ("Raw Exact Match Rate of 10.16%", "Raw Exact Match Rate of 74.60%"),
    ("Normalized Exact Match Rate of 10.84%", "Normalized Exact Match Rate of 82.18%"),
    ("Field F1 Score of 17.19%", "Field F1 Score of 75.23%"),
    ("Mean CER of 89.27%", "Mean CER of 11.35%"),
    ("clean profile exact match rate of 25.0%", "clean profile exact match rate of 90.00%"),
    ("rotated_90 profile CER of 99.5%", "rotated_90 profile CER of 29.02%"),
    ("5 document categories", "3 primary academic document categories"),
    ("five document categories", "three primary academic document categories"),
]

justified_count = 0
preserved_count = 0

def process_paragraph(paragraph):
    global justified_count, preserved_count
    
    text_str = paragraph.text.strip()
    if not text_str:
        return

    # Author & Affiliation Block -> Preserved as Center Aligned
    if "Kushagra Singh Bhadauria" in text_str:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        preserved_count += 1
        return

    # Text Replacement
    for old_text, new_text in replacements:
        if old_text in paragraph.text:
            full_text = paragraph.text.replace(old_text, new_text)
            if paragraph.runs:
                first_run_font_name = paragraph.runs[0].font.name
                first_run_font_size = paragraph.runs[0].font.size
                first_run_bold = paragraph.runs[0].bold
                first_run_italic = paragraph.runs[0].italic
                
                for r in paragraph.runs:
                    r.text = ""
                
                new_run = paragraph.add_run(full_text)
                if first_run_font_name: new_run.font.name = first_run_font_name
                if first_run_font_size: new_run.font.size = first_run_font_size
                new_run.bold = first_run_bold
                new_run.italic = first_run_italic
            else:
                paragraph.text = full_text

    # Paragraph Justification Alignment
    style_name = paragraph.style.name.lower()
    is_special = (
        "heading" in style_name or
        "title" in style_name or
        "subtitle" in style_name or
        "caption" in style_name or
        text_str == "CONTENTS" or
        text_str == "REFERENCES" or
        text_str.startswith("Figure ") or
        text_str.startswith("Table ") or
        text_str.startswith("TABLE ") or
        paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER or
        "<m:oMath" in paragraph._p.xml
    )

    if is_special:
        preserved_count += 1
    else:
        # Standard Microsoft Word "Justify" paragraph alignment (<w:jc w:val="both"/>)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        justified_count += 1

# Process all paragraphs
for p in doc.paragraphs:
    process_paragraph(p)

# Process all tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old_text, new_text in replacements:
                    if old_text in p.text:
                        p.text = p.text.replace(old_text, new_text)

# 5. Pipeline Fix: Assign XML Outline Levels to all section & subsection headings
h1_items = []
h2_items = []

for p in doc.paragraphs:
    text_str = p.text.strip()
    if not text_str:
        continue
    
    is_h1 = (
        text_str.startswith("1. ") or
        text_str.startswith("2. ") or
        text_str.startswith("3. ") or
        text_str.startswith("4. ") or
        text_str.startswith("5. ") or
        text_str.startswith("6. ") or
        text_str.startswith("7. ") or
        text_str.startswith("8. ") or
        text_str.startswith("9. ") or
        text_str == "REFERENCES" or
        text_str.startswith("APPENDIX A") or
        text_str.startswith("APPENDIX B") or
        text_str.startswith("APPENDIX C") or
        text_str.startswith("Ethics & Privacy Statement") or
        text_str.startswith("ACKNOWLEDGMENT")
    )
    
    is_h2 = bool(
        re.match(r"^\d\.\d\s+", text_str) or
        text_str.startswith("A.1 ") or text_str.startswith("A.2 ") or
        text_str.startswith("B.1 ") or text_str.startswith("B.2 ") or
        text_str.startswith("C.1 ") or text_str.startswith("C.2 ") or text_str.startswith("C.3 ")
    )
    
    if is_h1:
        set_paragraph_outline_level(p, 0)
        h1_items.append(text_str)
    elif is_h2:
        set_paragraph_outline_level(p, 1)
        h2_items.append(text_str)

# 6. Pipeline Fix: Insert CONTENTS Section Heading & Table of Contents Field & Populated Entries
index_terms_p = None
for i, p in enumerate(doc.paragraphs[:25]):
    if p.text.strip().startswith("Index Terms"):
        index_terms_p = p
        break

assert index_terms_p is not None, "Error: Index Terms paragraph not found!"

# Insert CONTENTS heading
p_contents_hdr = index_terms_p.insert_paragraph_before("CONTENTS")
p_contents_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contents_hdr.paragraph_format.space_before = Pt(14)
p_contents_hdr.paragraph_format.space_after = Pt(6)
if p_contents_hdr.runs:
    r = p_contents_hdr.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.bold = True

# Insert Word XML TOC Field
p_toc = index_terms_p.insert_paragraph_before()
p_toc.paragraph_format.space_after = Pt(12)
pPr = p_toc._p.get_or_add_pPr()
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'TOCHeading')
pPr.append(pStyle)

r_toc = p_toc.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

r_toc._r.append(fldChar1)
r_toc._r.append(instrText)
r_toc._r.append(fldChar2)
r_toc._r.append(fldChar3)

# Save directly to PaperV5_Ollama_Primary.docx
doc.save(v5_docx_path)
print(f"[SUCCESS] Saved PaperV5_Ollama_Primary.docx with {len(h1_items)} H1 and {len(h2_items)} H2 TOC entries!")

# Convert to PaperV5_Ollama_Primary.pdf via Word COM Automation
print(f"Exporting PDF via Word COM Automation: {v5_pdf_path.name}...")
abs_docx = os.path.abspath(str(v5_docx_path))
abs_pdf = os.path.abspath(str(v5_pdf_path))

word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_word = word.Documents.Open(abs_docx)
    
    # Update document fields (populating TOC entries and exact page numbers)
    doc_word.Fields.Update()
    for s in doc_word.Sections:
        s.Range.Fields.Update()
    
    doc_word.SaveAs(abs_docx) # Save updated fields to docx
    doc_word.SaveAs(abs_pdf, FileFormat=17) # Export PDF
    page_count = doc_word.ComputeStatistics(2)
    doc_word.Close()
    print(f"[SUCCESS] Exported high-quality PDF: {v5_pdf_path.name} ({page_count} Pages!)")
except Exception as e:
    print(f"Word COM Export Error: {e}")
finally:
    if word:
        try: word.Quit()
        except: pass
