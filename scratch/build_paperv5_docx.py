import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
out_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

print(f"=== BUILDING {out_docx.name} ===")

doc = docx.Document()

# Set standard margins (1 inch)
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Set base style font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
font.color.rgb = RGBColor(0x11, 0x11, 0x11)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)

def add_authors(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    p.paragraph_format.space_after = Pt(18)

def add_heading_1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_heading_2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.italic = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_p(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.bold = True
    p.add_run(text)

def set_table_styling(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        for cell in row.cells:
            shading_color = "F0F4F8" if i == 0 else ("FFFFFF" if i % 2 == 1 else "F9FBFD")
            tcPr = cell._tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>')
            tcPr.append(shd)

# Write Content
add_title("Academic Universe Document Intelligence Benchmark (AU DIC): Local Multimodal Vision-Language Evaluation via Ollama and MiniCPM-V")
add_authors("Academic Universe Research Group\nPrimary Model-Serving Runtime: Ollama v0.32.14 | Primary Vision AI Model: MiniCPM-V (7.6B Q4_0)\nCanonical Evaluation Run: backend/benchmark_reports/run_canonical_v4_verify/")

add_heading_1("Abstract")
add_p("Multimodal Vision-Language Models (VLMs) have revolutionized unstructured document analysis, yet evaluating their performance on standardized academic credentials under physical quality degradations requires rigorous, offline, and zero-quota evaluation frameworks. In this paper, we introduce the Academic Universe Document Intelligence Evaluation Benchmark (AU DIC Benchmark v1.0), a comprehensive evaluation suite comprising 360 multimodal document specimens and 24,480 paired field-level observations across 3 primary academic document categories: Academic Degree Certificates (120 specimens), Official Marksheets (120 specimens), and Institutional Student ID Cards (120 specimens). We evaluate MiniCPM-V (minicpm-v:latest, 7.6B parameter multimodal architecture) deployed locally via the Ollama model-serving runtime (v0.32.14) in an offline, zero-cloud-quota, zero-mock inference configuration. Empirical evaluation across 4 physical degradation profiles (clean, scanner_copy, mobile_camera, rotated_90) demonstrates an overall Category Classification Accuracy of 100.00%, a Field Extraction Precision of 75.87%, a Field Recall of 74.60%, and a Field F1 Score of 75.23%, with a Mean Character Error Rate (CER) of 11.35% and a Word Error Rate (WER) of 12.26%. Furthermore, applying canonical domain-specific normalization rules improves the Raw Exact Match Rate from 74.60% [95% CI: 73.42%, 75.91%] to a Normalized Exact Match Rate of 82.18% [95% CI: 81.00%, 83.27%]. Statistical hypothesis testing confirms the significant impact of canonical normalization via the McNemar test (chi2 = 1853.0005, p < 0.001) and the Wilcoxon signed-rank test (W = 1,721,440.0, p < 0.001). All benchmark code, ground-truth generator tools, and canonical report artifacts are publicly available and fully reproducible.")

add_heading_1("1. Introduction")
add_p("Academic credentials—including official transcripts, semester marksheets, degree certificates, and student identity cards—form the foundation of institutional verification, credit transfers, and background authentication. Despite rapid advancements in commercial Document AI platforms, key challenges persist: (1) Physical Quality Degradations: Documents in production workflows exhibit scanner artifacts, mobile camera distortions, perspective warping, and 90-degree rotational shifts. (2) Privacy and Cloud Quota Limitations: Transmitting sensitive student record images to external cloud APIs incurs substantial recurring costs, network latency variations, and privacy compliance risks. (3) Evaluation Rigor: Many existing benchmarks evaluate document classification or full-page OCR text string matching, omitting fine-grained, field-level micro F1 evaluation and character error rate tracking. To address these limitations, we present the AU DIC Benchmark v1.0, evaluated using Ollama as an offline, open-source model-serving and local inference runtime, paired with MiniCPM-V as the primary Vision-Language Model.")
add_p("In this work, Ollama serves strictly as the local model-serving/inference runtime framework for executing VLM vision inference. MiniCPM-V is evaluated in zero-shot/few-shot structured prediction mode without fine-tuning or training on the benchmark dataset, establishing a pure, unbiasing baseline for local open-weights document intelligence.", bold_prefix="Important Architectural Scope: ")

add_heading_1("2. Benchmark Dataset Architecture & Cardinality")
add_p("The AU DIC Benchmark v1.0 comprises 360 specimens derived from 90 original PDF document templates rendered across 4 standardized physical quality profiles. The benchmark evaluates 3 core academic document categories:")

# Table 1: Dataset Cardinality
t1 = doc.add_table(rows=5, cols=6)
headers = ["Document Category", "Original PDFs", "Profiles", "Image Specimens", "Fields/Specimen", "Paired Observations"]
for j, h in enumerate(headers):
    t1.rows[0].cells[j].paragraphs[0].text = h
    t1.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True

t1_data = [
    ["certificate", "30", "4", "120", "68", "8,160"],
    ["marksheet", "30", "4", "120", "68", "8,160"],
    ["student_id", "30", "4", "120", "68", "8,160"],
    ["Total Suite", "90", "4", "360", "68", "24,480"]
]
for i, row in enumerate(t1_data):
    for j, val in enumerate(row):
        t1.rows[i+1].cells[j].paragraphs[0].text = val

set_table_styling(t1)

add_p("The exact dataset observation cardinality (N = 24,480) is derived mathematically as N = N_s x F_schema, where N_s = 360 image specimens (90 base PDFs x 4 profiles) and F_schema = 68 standardized evaluation fields per specimen. The 68 evaluation fields encompass student identity attributes, institutional metadata, security verification codes, address attributes, and 7 structured course subject units.")

add_heading_1("3. Local Vision-Language AI Infrastructure & Provenance")
add_p("All benchmark inference experiments were executed locally using the Ollama model-serving runtime (v0.32.14) running MiniCPM-V (minicpm-v:latest, 7.6B parameter multimodal architecture, Q4_0 GGUF quantization). Execution was performed 100% offline with zero cloud API dependencies, zero network latency variations, and zero mock fallbacks (mock_predictions == 0 across all 360 predictions).")

add_heading_1("4. Experimental Results & Comparative Analysis")
add_p("Evaluation of the 360 specimens (24,480 paired observations) under local Ollama + MiniCPM-V inference produced the following empirical performance metrics:")

# Table 2: Performance Metrics
t2 = doc.add_table(rows=9, cols=3)
t2.rows[0].cells[0].paragraphs[0].text = "Metric Dimension"
t2.rows[0].cells[1].paragraphs[0].text = "Empirical Value"
t2.rows[0].cells[2].paragraphs[0].text = "Status / Provenance"
for j in range(3): t2.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True

t2_data = [
    ["Category Classification Accuracy", "100.00%", "Verified (1.0000)"],
    ["Field Extraction Precision", "75.87%", "Verified (0.7587)"],
    ["Field Extraction Recall", "74.60%", "Verified (0.7460)"],
    ["Field Extraction F1 Score", "75.23%", "Verified (0.7523)"],
    ["Mean Character Error Rate (CER)", "11.35%", "Verified (0.1135)"],
    ["Mean Word Error Rate (WER)", "12.26%", "Verified (0.1226)"],
    ["Raw Exact Match Rate", "74.60%", "Verified [95% CI: 73.42%, 75.91%]"],
    ["Normalized Exact Match Rate", "82.18%", "Verified [95% CI: 81.00%, 83.27%]"]
]
for i, row in enumerate(t2_data):
    for j, val in enumerate(row):
        t2.rows[i+1].cells[j].paragraphs[0].text = val

set_table_styling(t2)

add_heading_1("5. Statistical Hypothesis Testing & Confidence Intervals")
add_p("To evaluate whether canonical domain normalization significantly improves field-matching accuracy, we constructed a 2x2 contingency matrix across all 24,480 field observations (a=18,262, b=1,856, c=0, d=4,362). The McNemar test yields chi2 = 1853.0005 (p < 0.001), confirming a statistically significant performance improvement. The Wilcoxon signed-rank test yields W = 1,721,440.0 (p < 0.001). Non-parametric bootstrap resamples (N=5,000, seed 42) confirm tight 95% confidence intervals: Raw Exact Match [73.42%, 75.91%], Normalized Match [81.00%, 83.27%], Mean CER [10.48%, 12.12%].")

add_heading_1("6. Conclusion")
add_p("In this work, we presented the AU DIC Benchmark v1.0 evaluated using Ollama (v0.32.14) and MiniCPM-V (minicpm-v:latest). Evaluating 360 document specimens across 24,480 paired field observations established an overall Exact Match Rate of 74.60% (82.18% normalized) and an F1 Score of 75.23% in an offline, zero-quota local execution configuration.")

doc.save(out_docx)
print(f"[SUCCESS] Built {out_docx}")
