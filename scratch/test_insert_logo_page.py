import os
import docx
import re
import subprocess
import win32com.client
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
logo_path = workspace / "public" / "new_logo_2.png"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_Logo.docx"
test_pdf_path = workspace / "docs" / "paper" / "PaperV5_Test_Logo.pdf"

assert v4_docx_path.exists(), "Baseline docx missing!"
assert logo_path.exists(), "Logo image missing!"

doc = docx.Document(v4_docx_path)
p_title = doc.paragraphs[0]

# PAGE 1: Sharda University Logo ONLY
p_logo = p_title.insert_paragraph_before()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_logo.paragraph_format.space_before = Pt(230) # Vertical centering on Page 1
p_logo.paragraph_format.space_after = Pt(0)
r_logo = p_logo.add_run()
r_logo.add_picture(str(logo_path), width=Inches(4.5))

# PAGE 2: CONTENTS Heading & Table of Contents Field
p_contents_hdr = p_title.insert_paragraph_before("CONTENTS")
p_contents_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contents_hdr.paragraph_format.page_break_before = True # Starts CONTENTS on Page 2
p_contents_hdr.paragraph_format.space_before = Pt(14)
p_contents_hdr.paragraph_format.space_after = Pt(6)
if p_contents_hdr.runs:
    r = p_contents_hdr.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.bold = True

p_toc = p_title.insert_paragraph_before()
p_toc.paragraph_format.space_after = Pt(6)
pPr = p_toc._p.get_or_add_pPr()
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'TOCHeading')
pPr.append(pStyle)

r_toc = p_toc.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

r_toc._r.append(fldChar1)
r_toc._r.append(instrText)
r_toc._r.append(fldChar2)
r_toc._r.append(fldChar3)

# Populate TOC Entries with updated page numbers (+1 shift for Page 1 Logo)
toc_entries = [
    ("1. Introduction", "4", 1),
    ("1.1 Background & Motivation", "4", 2),
    ("1.2 Research Objectives", "4", 2),
    ("1.3 Main Research Contributions", "4", 2),
    ("2. Related Work", "6", 1),
    ("2.1 Vision-Language Models & Document Intelligence", "6", 2),
    ("2.2 Document Image Classification & Parsing Benchmarks", "6", 2),
    ("2.3 Information Extraction & Optical Character Recognition", "7", 2),
    ("2.4 Synthetic Data Generation & Benchmark Suites", "7", 2),
    ("2.5 Semantic Normalization & Post-Processing", "8", 2),
    ("2.6 Error Taxonomies & Diagnostic Evaluation", "8", 2),
    ("3. ADBG v1.0 & AU DIC Benchmark System Architecture", "9", 1),
    ("3.1 ADBG v1.0 Synthetic Generator Subsystem", "9", 2),
    ("3.2 Optical Degradation Profile Processor", "9", 2),
    ("3.3 AU DIC Decoupled Benchmark Execution Subsystem", "10", 2),
    ("4. Methodology", "10", 1),
    ("4.1 Six-Stage Semantic Canonical Normalization Subsystem", "10", 2),
    ("4.2 Nine-Class Structured OCR Error Taxonomy", "13", 2),
    ("4.3 Mathematical Formulation of Evaluation Metrics", "14", 2),
    ("5. Results & Empirical Validation", "16", 1),
    ("5.1 Distinction Between Framework Validation, Benchmark Validation, and Model Performance", "16", 2),
    ("5.2 Framework Execution & System Verification Metrics", "16", 2),
    ("5.3 System Throughput & Execution Latency", "17", 2),
    ("5.4 Empirical Live Neural Model Evaluation Results", "17", 2),
    ("5.5 Empirical Ablation Study of Semantic Canonical Normalization", "19", 2),
    ("5.6 Statistical Significance Analysis (p < 0.0001)", "21", 2),
    ("5.7 Empirical Evaluation Scope & Methodological Limitations", "22", 2),
    ("5.8 Error Taxonomy Distribution Shift Analysis", "22", 2),
    ("6. Discussion & Threats to Validity", "23", 1),
    ("6.1 Scientific Contributions and Methodological Novelty", "23", 2),
    ("6.2 Discussion of Empirical Findings", "23", 2),
    ("6.3 Threats to Validity", "24", 2),
    ("7. Limitations Analysis", "24", 1),
    ("7.1 Methodological Limitations", "24", 2),
    ("8. Future Work", "25", 1),
    ("9. Conclusion", "25", 1),
    ("Ethics & Privacy Statement", "25", 1),
    ("ACKNOWLEDGMENT", "26", 1),
    ("APPENDIX A: REPRODUCIBILITY & SYSTEM SPECIFICATIONS", "26", 1),
    ("A.1 Reproducibility & System Environment Matrix", "26", 2),
    ("A.2 Technical Clarifications & Reviewer Inquiries", "26", 2),
    ("APPENDIX B: FIELD SPECIFICATION & OBSERVATION COUNT DERIVATION", "27", 1),
    ("B.1 Document Category Field Structure", "27", 2),
    ("B.2 Mathematical Derivation of 24,480 Paired Observations", "27", 2),
    ("APPENDIX C: EMPIRICAL STATISTICAL METHODOLOGY & BENCHMARKS", "28", 1),
    ("C.1 Empirical Category Confusion Matrix (360 Specimens)", "28", 2),
    ("C.2 McNemar Contingency Test & Normalization Rescues", "28", 2),
    ("C.3 Non-Parametric Bootstrap Confidence Intervals (B = 10,000)", "29", 2),
    ("REFERENCES", "29", 1),
]

for title, page_str, level in toc_entries:
    p_entry = p_title.insert_paragraph_before()
    p_entry.paragraph_format.line_spacing = 1.15
    p_entry.paragraph_format.space_after = Pt(2)
    p_entry.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    
    indent_prefix = "   " if level == 2 else ""
    full_title = f"{indent_prefix}{title}"
    
    r_title = p_entry.add_run(full_title)
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(9.5)
    if level == 1:
        r_title.bold = True
    
    r_page = p_entry.add_run(f"\t{page_str}")
    r_page.font.name = "Times New Roman"
    r_page.font.size = Pt(9.5)
    r_page.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

# PAGE 3: Title, Author Block, Abstract, and Index Terms
p_title.paragraph_format.page_break_before = True

# PAGE 4: 1. Introduction
for p in doc.paragraphs:
    if p.text.strip() == "1. Introduction":
        p.paragraph_format.page_break_before = True
        break

doc.save(test_docx_path)
print(f"[SUCCESS] Saved test DOCX: Page 1 = Logo, Page 2 = CONTENTS, Page 3 = Title/Abstract, Page 4 = Intro!")
