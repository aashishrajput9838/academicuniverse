import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_ieee_docx():
    md_path = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\902d6e6c-b35a-4d3e-a983-59d2e321b9fc\paper2_academic_universe_growth_intelligence.md"
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final.docx"
    fig1_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure1_Architecture_Pipeline.png"
    fig2_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure2_Decay_Sensitivity_Curves.png"

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()

    # Page Setup - 1 inch margins
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def process_table(t_lines):
        rows = []
        for line in t_lines:
            if '|' in line:
                parts = [p.strip() for p in line.split('|')[1:-1]]
                if parts and not all(set(p) <= {'-', ':', ' '} for p in parts):
                    rows.append(parts)
        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                if c_idx < len(table.columns):
                    cell = table.cell(r_idx, c_idx)
                    cell.text = val
                    set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

                    if r_idx == 0:
                        set_cell_background(cell, "1E293B")
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    else:
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "F8FAFC")
                        for run in p.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()

    while i < len(lines):
        line = lines[i].rstrip('\r\n')

        # Code Block handling (Check if Figure 1 or Figure 2 ASCII blocks)
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                block_content = '\n'.join(code_block_lines)
                
                # If ASCII Architecture block, replace with High-Res Image Figure 1
                if 'GitHub API' in block_content or 'Evidence Intelligence Layer' in block_content:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(10)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.add_run().add_picture(fig1_path, width=Inches(6.2))
                # If ASCII Decay block, replace with High-Res Image Figure 2
                elif 'Proficiency S(t)' in block_content or 'λ=0.01' in block_content:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(10)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.add_run().add_picture(fig2_path, width=Inches(5.8))
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    run = p.add_run(block_content)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and not line.startswith('>'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            process_table(table_lines)
            table_lines = []

        # Blockquotes
        if line.startswith('> '):
            quote_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:])
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[3:])
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[5:])
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        elif line.startswith('$$') and line.endswith('$$'):
            # Centered Display Math
            eq_text = line[2:-2].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(eq_text)
            run.font.name = 'Cambria Math'
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif line.strip() != '':
            p = doc.add_paragraph()
            text = line
            
            # Check for Figure/Table Captions
            if text.startswith('*Figure') or text.startswith('*Table'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(10)

            # Inline formatting parser
            pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$|`.*?`)')
            tokens = pattern.split(text)
            for token in tokens:
                if not token:
                    continue
                if token.startswith('**') and token.endswith('**'):
                    p.add_run(token[2:-2]).bold = True
                elif token.startswith('*') and token.endswith('*'):
                    r = p.add_run(token[1:-1])
                    r.italic = True
                    if text.startswith('*Figure') or text.startswith('*Table'):
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                elif token.startswith('$') and token.endswith('$'):
                    r = p.add_run(token[1:-1])
                    r.font.name = 'Cambria Math'
                    r.font.bold = True
                elif token.startswith('`') and token.endswith('`'):
                    r = p.add_run(token[1:-1])
                    r.font.name = 'Courier New'
                    r.font.size = Pt(10)
                else:
                    p.add_run(token)

        i += 1

    if in_table:
        process_table(table_lines)

    # Reference Enhancement Additions
    ref_heading = False
    for paragraph in doc.paragraphs:
        if '## References' in paragraph.text:
            ref_heading = True
            break
            
    if ref_heading:
        enhanced_refs = [
            "5. Corbett, A. T., & Anderson, J. R. (1994). Knowledge tracing: Modeling the acquisition of procedural knowledge. User Modeling and User-Adapted Interaction, 4(4), 253-278.",
            "6. Piech, C., Bassen, J., Huang, J., Ganguli, S., Sahami, M., Guibas, L. J., & Sohl-Dickstein, J. (2015). Deep knowledge tracing. Advances in Neural Information Processing Systems (NeurIPS), 28, 505-513.",
            "7. Bodily, R., & Verbert, K. (2017). Review of research on student-facing learning analytics dashboards. Journal of Learning Analytics, 4(3), 67-89.",
            "8. Ferguson, R. (2012). Learning analytics: Drivers, developments and challenges. International Journal of Technology Enhanced Learning, 4(5-6), 304-317.",
            "9. Jivet, I., Scheffel, M., Specht, M., & Drachsler, H. (2018). License to evaluate: Preparing learning analytics dashboards for educational practice. Proceedings of the 8th International Conference on Learning Analytics & Knowledge (LAK '18), 31-40."
        ]
        doc.add_paragraph()
        p_note = doc.add_paragraph()
        r_note = p_note.add_run("### Foundational Literature Additions (IEEE / ACM Learning Analytics Corpus):")
        r_note.bold = True
        r_note.font.size = Pt(11)
        r_note.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        
        for ref in enhanced_refs:
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.left_indent = Inches(0.25)
            p_ref.paragraph_format.space_after = Pt(4)
            r = p_ref.add_run(ref)
            r.font.size = Pt(10)

    doc.save(docx_path)
    print(f"Successfully compiled IEEE Final DOCX to {docx_path}")

if __name__ == '__main__':
    build_ieee_docx()
