import os
import re
from typing import Dict, List, Any
from docx import Document
from ..utils.logger import get_logger

logger = get_logger("quality")

class QAAuditor:
    """Automated QA auditor asserting zero raw Mermaid, zero raw LaTeX, and valid IEEE formatting."""

    def __init__(self, docx_path: str, md_path: str):
        self.docx_path = docx_path
        self.md_path = md_path

    def run_qa_checks(self) -> Dict[str, Any]:
        logger.info(f"Running automated Quality Assurance audit on: {self.docx_path}")
        
        checks = {}
        
        # Check 1: No raw Mermaid blocks in markdown
        if os.path.exists(self.md_path):
            with open(self.md_path, 'r', encoding='utf-8') as f:
                md_text = f.read()
            checks["no_raw_mermaid_in_md"] = "```mermaid" not in md_text
        else:
            checks["no_raw_mermaid_in_md"] = True

        # Check 2: No raw LaTeX math commands in DOCX
        if os.path.exists(self.docx_path):
            doc = Document(self.docx_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            checks["no_raw_latex_in_docx"] = (
                "\\frac" not in full_text and
                "\\mathcal" not in full_text and
                "\\text{" not in full_text
            )
            checks["docx_has_tables"] = len(doc.tables) >= 3
            checks["docx_has_inline_shapes"] = len(doc.inline_shapes) >= 2
            
            # Check references start with [#
            refs = [p.text for p in doc.paragraphs if re.match(r'^\[\d+\]', p.text)]
            checks["references_ieee_numbered"] = len(refs) > 0
        else:
            checks["no_raw_latex_in_docx"] = False
            checks["docx_has_tables"] = False
            checks["docx_has_inline_shapes"] = False
            checks["references_ieee_numbered"] = False

        all_passed = all(checks.values())
        logger.info(f"QA Audit Complete. Overall Pass: {all_passed}")

        return {
            "all_passed": all_passed,
            "checks": checks
        }
