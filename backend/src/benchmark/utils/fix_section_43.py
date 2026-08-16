import os
import shutil
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_section_43(doc_path):
    print(f"Fixing Section 4.3 in: {doc_path}")
    doc = Document(doc_path)
    
    # Locate Section 4.3 heading and Section 5 heading
    sec43_idx = -1
    sec5_idx = -1
    
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "4.3 Evaluation Metrics" or t.startswith("4.3 Evaluation Metrics"):
            sec43_idx = i
        elif t == "5. Results & Empirical Validation" or t.startswith("5. Results & Empirical Validation"):
            sec5_idx = i
            break
            
    print(f"sec43_idx: {sec43_idx}, sec5_idx: {sec5_idx}")
    
    if sec43_idx != -1 and sec5_idx != -1:
        # Remove old messy paragraphs between sec43_idx+1 and sec5_idx-1
        num_to_delete = sec5_idx - sec43_idx - 1
        print(f"Deleting {num_to_delete} old paragraphs in Section 4.3")
        
        for _ in range(num_to_delete):
            p = doc.paragraphs[sec43_idx + 1]
            p._element.getparent().remove(p._element)
            
        sec43_p = doc.paragraphs[sec43_idx]
        
        # Format the 4.3 heading
        sec43_p.style = doc.styles['Normal']
        sec43_p.paragraph_format.space_before = Pt(14)
        sec43_p.paragraph_format.space_after = Pt(6)
        sec43_p.paragraph_format.keep_with_next = True
        sec43_p.text = ""
        r_head = sec43_p.add_run("4.3 Evaluation Metrics")
        r_head.font.name = "Calibri"
        r_head.font.size = Pt(12)
        r_head.font.bold = True
        r_head.font.color.rgb = RGBColor(31, 73, 125)
        
        curr_elem = sec43_p._element
        
        def insert_paragraph(text, bold_prefix="", indent=False, space_after=4, is_eq=False):
            nonlocal curr_elem
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2 if not is_eq else 4)
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.line_spacing = 1.15
            
            if indent:
                p.paragraph_format.left_indent = Inches(0.25)
                
            if is_eq:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Inches(0.4)
                
            if bold_prefix:
                r_b = p.add_run(bold_prefix)
                r_b.font.name = "Calibri"
                r_b.font.size = Pt(10)
                r_b.font.bold = True
                r_b.font.color.rgb = RGBColor(38, 38, 38)
                
            r_t = p.add_run(text)
            r_t.font.name = "Calibri" if not is_eq else "Cambria Math"
            r_t.font.size = Pt(10 if not is_eq else 10.5)
            if is_eq:
                r_t.font.italic = True
                r_t.font.color.rgb = RGBColor(31, 73, 125)
            else:
                r_t.font.color.rgb = RGBColor(51, 51, 51)
                
            curr_elem.addnext(p._element)
            curr_elem = p._element
            return p

        # 1. Introductory text
        insert_paragraph("The performance of the document intelligence pipeline and evaluated models is quantified across six formal evaluation metrics spanning category classification, field-level information extraction, edit distance, joint record completeness, and runtime throughput:", space_after=6)

        # Metric 1
        insert_paragraph("The proportion of specimens where the predicted document category (ŷ_i) matches the ground-truth category (y_i):", bold_prefix="1. Category Classification Accuracy (Acc_cat): ", space_after=2)
        insert_paragraph("Acc_cat = (1 / N) * ∑_{i=1}^N  I(ŷ_i = y_i)   (1)", is_eq=True, space_after=2)
        insert_paragraph("where N = 360 total evaluated specimens, and I(·) is the binary indicator function.", indent=True, space_after=6)

        # Metric 2
        insert_paragraph("Macro-averaged key-value field extraction metrics across all extracted target entities:", bold_prefix="2. Field Extraction Precision, Recall, and F1-Score (Precision, Recall, F1): ", space_after=2)
        insert_paragraph("Precision = TP / (TP + FP),    Recall = TP / (TP + FN)\nF1 = (2 * Precision * Recall) / (Precision + Recall)   (2)", is_eq=True, space_after=6)

        # Metric 3
        insert_paragraph("Levenshtein character edit distance D_char(ŝ, s) [21] between canonically normalized predicted field strings (ŝ) and expected ground-truth field strings (s), normalized by total ground-truth character length (|s|):", bold_prefix="3. Character Error Rate (CER): ", space_after=2)
        insert_paragraph("CER = D_char(ŝ, s) / |s|   (3)", is_eq=True, space_after=6)

        # Metric 4
        insert_paragraph("Tokenized word-level edit distance D_word(ŵ, w) [20] between predicted field strings (ŵ) and ground-truth field strings (w), normalized by total ground-truth word count (|w|):", bold_prefix="4. Word Error Rate (WER): ", space_after=2)
        insert_paragraph("WER = D_word(ŵ, w) / |w|   (4)", is_eq=True, space_after=6)

        # Metric 5
        insert_paragraph("The percentage of specimens that achieve both 100% key-value field extraction (F1_i = 1.0) AND correct top-level category classification (ŷ_i = y_i) simultaneously:", bold_prefix="5. Joint Record Exact Match Rate (EM): ", space_after=2)
        insert_paragraph("EM = (1 / N) * ∑_{i=1}^N  I(ŷ_i = y_i  AND  F1_i = 1.0)   (5)", is_eq=True, space_after=6)

        # Metric 6
        insert_paragraph("Mean execution latency per specimen (L_proc) in milliseconds and framework processing throughput (TH) in specimens per second:", bold_prefix="6. Processing Latency & System Throughput (L_proc, TH): ", space_after=2)
        insert_paragraph("L_proc = T_total / N,    TH = N / T_total   (6)", is_eq=True, space_after=10)

        out_path = doc_path.replace(".docx", "_Fixed.docx")
        doc.save(out_path)
        print(f"Successfully saved structured Section 4.3 to: {out_path}")

if __name__ == "__main__":
    docs = [
        r"docs\paper\PaperV4_Final_Submission.docx",
        r"docs\paper\PaperV4_Final_Submission.docx.tmp"
    ]
    for d in docs:
        if os.path.exists(d):
            fix_section_43(d)

