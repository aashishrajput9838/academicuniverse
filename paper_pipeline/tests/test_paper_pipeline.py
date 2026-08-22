import os
import pytest
from docx import Document
from paper_pipeline.parser import MarkdownParser
from paper_pipeline.validator import StructureValidator

from pathlib import Path

WORKSPACE = Path(__file__).resolve().parents[2]
MD_PATH = str(WORKSPACE / "docs" / "paper" / "Paper_V12.md")
DOCX_PATH = str(WORKSPACE / "docs" / "paper" / "PaperV12_Ollama_Primary.docx")

def test_markdown_parser():
    parser = MarkdownParser(MD_PATH)
    data = parser.parse()
    assert data["total_lines"] > 100
    assert len(data["headings"]) > 5

def test_structure_validator():
    with open(MD_PATH, 'r', encoding='utf-8') as f:
        content = f.read()
    assert "Abstract" in content
    assert "Introduction" in content
    assert ("Related Work" in content or "Literature" in content)
    assert "Methodology" in content
    assert "Experimental" in content

def test_no_raw_mermaid_in_docx():
    if os.path.exists(DOCX_PATH):
        doc = Document(DOCX_PATH)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "```mermaid" not in full_text

def test_no_raw_latex_in_docx():
    if os.path.exists(DOCX_PATH):
        doc = Document(DOCX_PATH)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "\\frac{" not in full_text
        assert "\\mathcal{" not in full_text

def test_ieee_references_numbered():
    if os.path.exists(DOCX_PATH):
        doc = Document(DOCX_PATH)
        refs = [p.text for p in doc.paragraphs if p.text.startswith("[")]
        assert len(refs) > 0
