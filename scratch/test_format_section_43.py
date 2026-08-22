import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"

assert v4_docx_path.exists(), "Baseline docx missing!"

doc = docx.Document(v4_docx_path)

def set_paragraph_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)

print("=== TESTING SECTION 4.3 FORMATTING TRANSFORMATION ===")

sec43_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("4.3 Evaluation Metrics") or p.text.strip().startswith("4.3 Mathematical Formulation"):
        sec43_idx = i
        break

assert sec43_idx is not None, "Section 4.3 not found!"
print(f"Found Section 4.3 at Paragraph P{sec43_idx}")

# Inspect existing items from P{sec43_idx} to Section 5
for j in range(sec43_idx, sec43_idx + 15):
    print(f"P{j}: '{doc.paragraphs[j].text.strip()}'")
