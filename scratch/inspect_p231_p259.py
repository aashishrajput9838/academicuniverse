import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"

doc = docx.Document(v4_docx_path)

p231 = doc.paragraphs[231]
p259 = doc.paragraphs[259]

print("=== P231 Properties ===")
print(f"Text: '{p231.text}'")
print(f"Style: '{p231.style.name}'")
print(f"Alignment: {p231.alignment}")
print(f"Space Before: {p231.paragraph_format.space_before}")
print(f"Space After: {p231.paragraph_format.space_after}")
if p231.runs:
    r = p231.runs[0]
    print(f"Run Font: name={r.font.name}, size={r.font.size}, bold={r.bold}, italic={r.italic}")

print("\n=== P259 Properties ===")
print(f"Text: '{p259.text[:80]}'")
print(f"Style: '{p259.style.name}'")
print(f"Alignment: {p259.alignment}")
