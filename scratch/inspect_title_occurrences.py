import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

print("=== SEARCHING TITLE OCCURRENCES IN DOCX FILES ===")

doc4 = docx.Document(v4_docx_path)
print(f"V4 Baseline Title Paragraph P0: '{doc4.paragraphs[0].text.strip()}'")

doc5 = docx.Document(v5_docx_path)
for i, p in enumerate(doc5.paragraphs):
    text_str = p.text.strip()
    if "ADBG v1.0" in text_str or "Framework: Reproducible" in text_str or "Smart Academic Document Intelligence" in text_str:
        print(f"V5 Paragraph P{i}: '{text_str}'")
