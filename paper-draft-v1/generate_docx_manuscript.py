import os
import re
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.path.append(os.path.join(os.getcwd(), 'paper-draft-v1'))
from latex_to_omml import convert_latex_to_omml_element

def create_journal_manuscript():
    doc = docx.Document()
    
    # 1. Page Setup: Standard 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Styling helper colors
    NAVY_HEX = "1B365D"
    DARK_TEXT_HEX = "2C3E50"
    LIGHT_BG_HEX = "F4F6F9"
    BORDER_HEX = "D0D7DE"

    NAVY_COLOR = RGBColor(0x1B, 0x36, 0x5D)
    DARK_COLOR = RGBColor(0x2C, 0x3E, 0x50)
    BODY_COLOR = RGBColor(0x22, 0x22, 0x22)
    MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)

    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def set_table_borders(table, color=BORDER_HEX, sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
        tblPr.append(borders)

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_font.color.rgb = BODY_COLOR
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Read research_paper.md
    base_dir = os.path.join(os.getcwd(), 'paper-draft-v1')
    paper_md_path = os.path.join(base_dir, 'research_paper.md')

    with open(paper_md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Read Table Markdown files
    tables_dir = os.path.join(base_dir, 'tables')
    table_files = {}
    if os.path.exists(tables_dir):
        for t_file in os.listdir(tables_dir):
            if t_file.endswith('.md'):
                t_path = os.path.join(tables_dir, t_file)
                with open(t_path, 'r', encoding='utf-8') as tf:
                    table_files[t_file] = tf.read()

    # Figures directory
    figures_dir = os.path.join(base_dir, 'figures')

    # Mapping of sections to figures and tables
    fig_map = {
        '4.1 Academic Universe Ecosystem & Module Scope': ('fig0_ecosystem_architecture.png', 'Figure 1: Academic Universe Platform Ecosystem Architecture showing the parent student-centric multi-tenant SaaS platform and highlighting AU DIC as the target research module evaluated in this paper.'),
        '9.1 System Overview': ('fig1_overall_architecture.png', 'Figure 2: AU DIC Four-Tier Microservice Architecture showing Presentation, API, Service, and Data layers with Multi-Tenant Organization Isolation.'),
        '9.3 Document Processing Pipeline': ('fig2_dic_pipeline.png', 'Figure 3: End-to-End Document Intelligence Pipeline from Upload Validation through AI Orchestration to HITL Staging.'),
        '9.4 Transaction-Safe Soft Deletion': ('fig4_transaction_sequence.png', 'Figure 5: Transaction-Safe Soft Deletion Sequence with MongoDB Session ACID Commit and Fallback Rollback.'),
        '12.2 Benchmark Configuration': ('fig8_benchmark_workflow.png', 'Figure 9: Benchmark Execution and Certification Workflow for Reproducible Experiment Evaluation.'),
        '14.2 Precision Comparison': ('fig5_accuracy_comparison.png', 'Figure 6: Field Extraction Precision Comparison Across Evaluation Systems per Synthetic Document.'),
        '14.3 Latency Breakdown': ('fig6_latency_breakdown.png', 'Figure 7: Mean Latency Decomposition (Upload, AI Inference, DB Staging) per System.'),
        '15.2 HITL Staging Value': ('fig3_hitl_review.png', 'Figure 4: HITL Staging Reviewer Interface with Per-Field Confidence Scores and Inline Correction.'),
        '8. Research Objectives': ('fig7_research_workflow.png', 'Figure 8: Complete Research Paper & Benchmark Infrastructure Artifact Generation Workflow.')
    }

    table_map = {
        '6.1 Enterprise Document Management Systems': ('table1_system_feature_matrix.md', 'Table 1: System Feature Matrix — AU DIC vs. Enterprise Alternatives'),
        '10.1 Layered Architecture': ('table2_tech_stack.md', 'Table 2: Technology Stack and Architecture Specifications'),
        '11.2 Document Categories and Quality Profiles': ('table3_dataset_summary.md', 'Table 3: Synthetic Dataset Summary and Document Quality Profiles'),
        '13.4 System-Level Aggregate Metrics': ('table4_evaluation_metrics.md', 'Table 4: Evaluation Metrics Formal Definitions'),
        '14.1 Overall Performance': ('table6_aggregate_metrics.md', 'Table 6: System-Level Aggregate Performance (EXP-VAL-20260729)'),
        '14. Document-Level Evaluation Results': ('table5_benchmark_results.md', 'Table 5: Document-Level Evaluation Results (EXP-VAL-20260729)'),
        '14.6 Category Breakdown': ('table7_category_breakdown.md', 'Table 7: Category Breakdown and Quality Profile Impact (SYS-PROP Only)'),
        '16. Threats to Validity': ('table8_threats_to_validity.md', 'Table 8: Comprehensive Threats to Validity Matrix'),
        '17. Limitations': ('table9_limitations.md', 'Table 9: System, Data, and Methodological Limitations'),
        '18. Future Work': ('table10_future_work.md', 'Table 10: Multi-Horizon Research Agenda for AU DIC v2.0')
    }

    rendered_figures = set()
    rendered_tables = set()

    # ---- Per-table column width ratios (proportional, summing to 1.0) ----
    # Manually tuned for optimal readability in 6.5" text width (8.5" - 2" margins).
    TABLE_COL_WIDTHS = {
        'table1_system_feature_matrix.md': [0.22, 0.22, 0.19, 0.19, 0.18],   # 5 cols: Feature + 4 systems
        'table2_tech_stack.md':            [0.10, 0.16, 0.18, 0.36, 0.20],    # 5 cols: Layer, Comp, Tech, Source, Ref
        'table3_dataset_summary.md': {
            2: [0.30, 0.70],                                                   # 2 cols: Attribute | Value
            10: [0.14, 0.11, 0.12, 0.10, 0.12, 0.08, 0.07, 0.07, 0.10, 0.09],# 10 cols: inventory
            3: [0.25, 0.40, 0.35],                                             # 3 cols: Profile | Desc | Degradation
        },
        'table4_evaluation_metrics.md': {
            3: [0.22, 0.38, 0.40],                                             # 3 cols: Metric | Formula | Desc
            2: [0.30, 0.70],                                                   # 2 cols: Notation | Meaning
        },
        'table5_benchmark_results.md':     [0.14, 0.11, 0.06, 0.06, 0.06, 0.12, 0.08, 0.11, 0.13, 0.07],  # 10 cols
        'table6_aggregate_metrics.md': {
            13: [0.09, 0.06, 0.07, 0.07, 0.07, 0.07, 0.07, 0.07, 0.07, 0.09, 0.07, 0.09, 0.07, 0.07],
            5:  [0.20, 0.20, 0.20, 0.20, 0.20],                                # 5 cols: latency breakdown
        },
        'table7_category_breakdown.md': {
            9:  [0.13, 0.06, 0.10, 0.10, 0.10, 0.10, 0.08, 0.10, 0.13],       # 9 cols: category
            6:  [0.16, 0.16, 0.10, 0.16, 0.16, 0.16],                          # 6 cols: quality profile
            2:  [0.45, 0.55],                                                   # 2 cols: HITL impact
        },
        'table8_threats_to_validity.md':   [0.07, 0.30, 0.10, 0.53],           # 4 cols: ID, Threat, Severity, Mitigation
        'table9_limitations.md':           [0.15, 0.30, 0.25, 0.30],           # 4 cols: Limitation, Desc, Impact, Resolution
        'table10_future_work.md': {
            4:  [0.07, 0.22, 0.53, 0.10],                                      # 4 cols: ID, Item, Description, Priority
            3:  [0.10, 0.30, 0.60],                                            # 3 cols: ID, Item, Description (research)
        },
    }

    USABLE_WIDTH_INCHES = 6.5  # 8.5" page - 2" total margins

    def _get_col_widths(table_filename, num_cols):
        """Resolve column widths for a given table. Returns list of Inches."""
        spec = TABLE_COL_WIDTHS.get(table_filename)
        if spec is None:
            # Fallback: equal widths
            w = USABLE_WIDTH_INCHES / num_cols
            return [Inches(w)] * num_cols
        if isinstance(spec, dict):
            ratios = spec.get(num_cols)
            if ratios is None:
                w = USABLE_WIDTH_INCHES / num_cols
                return [Inches(w)] * num_cols
            return [Inches(r * USABLE_WIDTH_INCHES) for r in ratios]
        # Direct list
        if len(spec) == num_cols:
            return [Inches(r * USABLE_WIDTH_INCHES) for r in spec]
        w = USABLE_WIDTH_INCHES / num_cols
        return [Inches(w)] * num_cols

    def _apply_autofit_window(table):
        """Set table layout to AutoFit to Window via OOXML."""
        tblPr = table._tbl.tblPr
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        existing = tblPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
        for e in existing:
            tblPr.remove(e)
        tblPr.append(tblW)

    def _set_cell_no_wrap(cell):
        """Prevent awkward mid-word splitting inside a cell."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(f'<w:noWrap {nsdecls("w")}/>'))

    def _set_paragraph_keep_lines(paragraph):
        """Keep all lines of a paragraph on the same page."""
        pPr = paragraph._p.get_or_add_pPr()
        pPr.append(parse_xml(f'<w:keepLines {nsdecls("w")}/>'))

    def _set_cell_width(cell, width_inches):
        """Explicitly set preferred cell width."""
        tcPr = cell._tc.get_or_add_tcPr()
        tw = int(width_inches / Inches(1) * 1440)  # convert to twips
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{tw}" w:type="dxa"/>')
        existing = tcPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
        for e in existing:
            tcPr.remove(e)
        tcPr.append(tcW)

    def _apply_professional_borders(table):
        """IEEE-style three-line table borders: heavy top/bottom, thin inside horizontal, no vertical."""
        tblPr = table._tbl.tblPr
        existing = tblPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        for e in existing:
            tblPr.remove(e)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:space="0" w:color="{NAVY_HEX}"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="{NAVY_HEX}"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{BORDER_HEX}"/>'
            f'<w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

    def _apply_header_bottom_border(row, num_cols):
        """Add a heavier bottom border to the header row to separate it from data."""
        for c_idx in range(num_cols):
            cell = row.cells[c_idx]
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="{NAVY_HEX}"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    def render_markdown_table(doc, table_md_content, title_caption, table_filename=""):
        """Build a professional journal-quality table from Markdown source.
        
        Features:
        - Per-table optimized column widths
        - AutoFit to Window
        - Vertical centering in all cells
        - No-wrap / keep-lines to prevent awkward word splitting
        - Header row repeats across pages
        - IEEE-style three-line borders with navy header band
        - Alternating row shading for readability
        - Generous cell padding
        """
        lines = [line.strip() for line in table_md_content.strip().split('\n') if line.strip()]

        # ---- Parse all sub-tables and notes ----
        all_sub_tables = []
        current_rows = []
        note_text = ""
        sub_headers = []

        for line in lines:
            if line.startswith('#'):
                # Sub-table header — flush previous
                if current_rows:
                    all_sub_tables.append((sub_headers[-1] if sub_headers else "", current_rows))
                    current_rows = []
                sub_headers.append(line.lstrip('#').strip())
                continue
            if line.startswith('|') and '|' in line[1:]:
                if re.match(r'^\|[\s\:\-\|]+\|$', line):
                    continue
                cells = [c.strip() for c in line.split('|')[1:-1]]
                current_rows.append(cells)
            elif line.startswith('>') or 'Note:' in line:
                note_text += " " + line.lstrip('>').strip()
            # Skip non-table lines (bullet points, plain text descriptions)

        if current_rows:
            all_sub_tables.append((sub_headers[-1] if sub_headers else "", current_rows))

        if not all_sub_tables:
            return

        # ---- Table Caption ----
        p_cap = doc.add_paragraph()
        p_cap.paragraph_format.space_before = Pt(16)
        p_cap.paragraph_format.space_after = Pt(5)
        p_cap.paragraph_format.keep_with_next = True
        p_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Split caption into "Table N:" (bold) + rest (regular italic)
        cap_match = re.match(r'^(Table\s+\d+[\.\:]?)\s*(.*)', title_caption)
        if cap_match:
            run_label = p_cap.add_run(cap_match.group(1) + " ")
            run_label.font.name = 'Calibri'
            run_label.font.size = Pt(10)
            run_label.font.bold = True
            run_label.font.color.rgb = NAVY_COLOR

            run_desc = p_cap.add_run(cap_match.group(2))
            run_desc.font.name = 'Calibri'
            run_desc.font.size = Pt(10)
            run_desc.font.italic = True
            run_desc.font.color.rgb = DARK_COLOR
        else:
            run_cap = p_cap.add_run(title_caption)
            run_cap.font.name = 'Calibri'
            run_cap.font.size = Pt(10)
            run_cap.font.bold = True
            run_cap.font.color.rgb = DARK_COLOR

        # ---- Render each sub-table ----
        for st_idx, (sub_title, table_rows) in enumerate(all_sub_tables):
            if not table_rows:
                continue

            # Sub-table heading (if multi-section table files like table3, table4, etc.)
            if sub_title and len(all_sub_tables) > 1:
                p_sub = doc.add_paragraph()
                p_sub.paragraph_format.space_before = Pt(10)
                p_sub.paragraph_format.space_after = Pt(3)
                p_sub.paragraph_format.keep_with_next = True
                run_sub = p_sub.add_run(sub_title)
                run_sub.font.name = 'Calibri'
                run_sub.font.size = Pt(9.5)
                run_sub.font.bold = True
                run_sub.font.italic = True
                run_sub.font.color.rgb = DARK_COLOR

            num_rows = len(table_rows)
            num_cols = max(len(r) for r in table_rows)
            col_widths = _get_col_widths(table_filename, num_cols)

            # Create table
            tbl = doc.add_table(rows=num_rows, cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Professional borders and AutoFit
            _apply_professional_borders(tbl)
            _apply_autofit_window(tbl)

            # ---- Populate cells ----
            for r_idx, row in enumerate(table_rows):
                tr = tbl.rows[r_idx]
                trPr = tr._tr.get_or_add_trPr()

                # Prevent row splitting across pages
                trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

                is_header = (r_idx == 0)
                if is_header:
                    # Header row repeats on every page
                    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

                for c_idx in range(num_cols):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    cell = tr.cells[c_idx]

                    # Vertical centering
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                    # Set explicit column width
                    _set_cell_width(cell, col_widths[c_idx])

                    # Cell padding
                    set_cell_margins(cell, top=80, bottom=80, left=120, right=120)

                    # Paragraph formatting
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing = 1.10
                    _set_paragraph_keep_lines(p)

                    # Clean markdown bold/italic markers
                    has_bold = '**' in cell_text
                    clean_text_no_md = re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text)
                    clean_text_no_md = re.sub(r'\*(.*?)\*', r'\1', clean_text_no_md)
                    clean_text_no_md = re.sub(r'`(.*?)`', r'\1', clean_text_no_md)
                    clean_text = re.sub(r'\$(.*?)\$', r'\1', clean_text_no_md)

                    # Determine if cell contains short numeric data (center it)
                    is_numeric = bool(re.match(r'^[\d\.\,\%\-\s]+$', clean_text.strip()))

                    if is_header:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif is_numeric:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif c_idx == 0 and num_cols > 3:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Render cell runs and math objects
                    if '$' in cell_text:
                        tokens = re.split(r'(\$.*?\$)', cell_text)
                        for tok in tokens:
                            if not tok:
                                continue
                            if tok.startswith('$') and tok.endswith('$'):
                                try:
                                    omml_elem = convert_latex_to_omml_element(tok, is_display=False)
                                    p._p.append(omml_elem)
                                except Exception:
                                    r = p.add_run(tok.strip('$'))
                                    r.font.name = 'Calibri'
                                    r.font.size = Pt(8.5) if (is_header and num_cols >= 8) else Pt(9.5)
                            else:
                                clean_t = re.sub(r'\*\*(.*?)\*\*', r'\1', tok)
                                clean_t = re.sub(r'\*(.*?)\*', r'\1', clean_t)
                                clean_t = re.sub(r'`(.*?)`', r'\1', clean_t)
                                if clean_t:
                                    r = p.add_run(clean_t)
                                    r.font.name = 'Calibri'
                                    r.font.size = Pt(8.5) if (is_header and num_cols >= 8) else Pt(9.5)
                                    if is_header:
                                        r.font.bold = True
                                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    elif has_bold:
                                        r.font.bold = True
                                        r.font.color.rgb = DARK_COLOR
                                    else:
                                        r.font.color.rgb = BODY_COLOR
                    else:
                        run = p.add_run(clean_text)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9) if num_cols >= 8 else Pt(9.5)

                        if is_header:
                            set_cell_background(cell, NAVY_HEX)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(8.5) if num_cols >= 8 else Pt(9.5)
                        else:
                            if r_idx % 2 == 0:
                                set_cell_background(cell, LIGHT_BG_HEX)
                            else:
                                set_cell_background(cell, "FFFFFF")

                            if has_bold:
                                run.font.bold = True
                                run.font.color.rgb = DARK_COLOR
                            else:
                                run.font.color.rgb = BODY_COLOR

                        if len(clean_text) < 25:
                            _set_cell_no_wrap(cell)

                    if is_header:
                        set_cell_background(cell, NAVY_HEX)

                # Add heavier border below header row
                if is_header:
                    _apply_header_bottom_border(tr, num_cols)

        # ---- Table footnote ----
        if note_text:
            p_note = doc.add_paragraph()
            p_note.paragraph_format.space_before = Pt(3)
            p_note.paragraph_format.space_after = Pt(14)
            run_n = p_note.add_run(note_text.strip())
            run_n.font.name = 'Calibri'
            run_n.font.size = Pt(8.5)
            run_n.font.italic = True
            run_n.font.color.rgb = MUTED_COLOR

    def add_figure_image(doc, png_filename, caption_text):
        png_path = os.path.join(figures_dir, png_filename)
        if os.path.exists(png_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(14)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.paragraph_format.keep_with_next = True
            
            run_img = p_img.add_run()
            run_img.add_picture(png_path, width=Inches(6.2))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(12)
            run_c = p_cap.add_run(caption_text)
            run_c.font.name = 'Calibri'
            run_c.font.size = Pt(9.5)
            run_c.font.italic = True
            run_c.font.color.rgb = DARK_COLOR

    # Header / Banner
    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_banner.paragraph_format.space_after = Pt(4)
    run_b = p_banner.add_run("ACADEMIC UNIVERSE DOCUMENT INTELLIGENCE CORE (AU DIC)")
    run_b.font.name = 'Calibri'
    run_b.font.size = Pt(10)
    run_b.font.bold = True
    run_b.font.color.rgb = NAVY_COLOR

    p_subbanner = doc.add_paragraph()
    p_subbanner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_subbanner.paragraph_format.space_after = Pt(16)
    run_sb = p_subbanner.add_run("Official Research Manuscript • Version 1.0.0 Certified Baseline • EXP-VAL-20260729")
    run_sb.font.name = 'Calibri'
    run_sb.font.size = Pt(9)
    run_sb.font.italic = True
    run_sb.font.color.rgb = MUTED_COLOR

    # Title
    title = "Academic Universe: AU DIC—A Human-in-the-Loop Multimodal Document Intelligence Framework for Verifiable Academic Credential Parsing"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(16)
    run_t = p_title.add_run(title)
    run_t.font.name = 'Calibri Light'
    run_t.font.size = Pt(22)
    run_t.font.bold = True
    run_t.font.color.rgb = NAVY_COLOR

    # Author Metadata Block
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_auth.paragraph_format.space_after = Pt(20)
    run_a = p_auth.add_run("Academic Universe Research & Infrastructure Team\nDepartment of Computer Science & Engineering • Academic Universe Inc.\nContact: research@academicuniverse.com")
    run_a.font.name = 'Calibri'
    run_a.font.size = Pt(10)
    run_a.font.color.rgb = DARK_COLOR

    # Parse sections
    sections = re.split(r'\n(?=##\s+)', md_text)

    for sec in sections:
        sec_str = sec.strip()
        if not sec_str:
            continue

        lines = sec_str.split('\n')
        header_line = lines[0].strip()

        if header_line.startswith('# '):
            continue
        elif header_line.startswith('## '):
            h_text = header_line.replace('## ', '').strip()
            
            if h_text.lower() in ['1. title', 'title']:
                continue

            # Heading 1
            p_h1 = doc.add_paragraph()
            p_h1.paragraph_format.space_before = Pt(16)
            p_h1.paragraph_format.space_after = Pt(6)
            p_h1.paragraph_format.keep_with_next = True
            run_h1 = p_h1.add_run(h_text)
            run_h1.font.name = 'Calibri Light'
            run_h1.font.size = Pt(14)
            run_h1.font.bold = True
            run_h1.font.color.rgb = NAVY_COLOR

            body_lines = lines[1:]
            
            # Box formatting for Abstract & Keywords
            if 'abstract' in h_text.lower() or 'keywords' in h_text.lower():
                full_body = " ".join([l.strip() for l in body_lines if l.strip() and not l.strip().startswith('---')])
                tbl_box = doc.add_table(rows=1, cols=1)
                tbl_box.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell_box = tbl_box.cell(0, 0)
                set_cell_background(cell_box, LIGHT_BG_HEX)
                set_cell_margins(cell_box, top=140, bottom=140, left=200, right=200)
                
                tcPr = cell_box._tc.get_or_add_tcPr()
                borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{NAVY_HEX}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
                tcPr.append(borders)

                p_b = cell_box.paragraphs[0]
                p_b.paragraph_format.space_before = Pt(2)
                p_b.paragraph_format.space_after = Pt(2)
                p_b.paragraph_format.line_spacing = 1.15
                
                run_bx = p_b.add_run(full_body)
                run_bx.font.name = 'Calibri'
                run_bx.font.size = Pt(10.5)
                if 'keywords' in h_text.lower():
                    run_bx.font.italic = True
                    run_bx.font.color.rgb = DARK_COLOR
                else:
                    run_bx.font.color.rgb = BODY_COLOR
                continue

            # Process body lines
            for b_line in body_lines:
                b_str = b_line.strip()
                if not b_str or b_str == '---':
                    continue

                # Subheadings (###)
                if b_str.startswith('### '):
                    sub_h = b_str.replace('### ', '').strip()
                    p_sub = doc.add_paragraph()
                    p_sub.paragraph_format.space_before = Pt(12)
                    p_sub.paragraph_format.space_after = Pt(4)
                    p_sub.paragraph_format.keep_with_next = True
                    run_sub = p_sub.add_run(sub_h)
                    run_sub.font.name = 'Calibri'
                    run_sub.font.size = Pt(12)
                    run_sub.font.bold = True
                    run_sub.font.color.rgb = DARK_COLOR

                    # Check figures & tables
                    for key in fig_map:
                        if key not in rendered_figures and key.lower() in sub_h.lower():
                            png_f, cap_f = fig_map[key]
                            add_figure_image(doc, png_f, cap_f)
                            rendered_figures.add(key)
                    for key in table_map:
                        if key not in rendered_tables and key.lower() in sub_h.lower():
                            tbl_f, tbl_cap = table_map[key]
                            if tbl_f in table_files:
                                render_markdown_table(doc, table_files[tbl_f], tbl_cap, table_filename=tbl_f)
                                rendered_tables.add(key)
                    continue

                # Math equations ($$...$$)
                if b_str.startswith('$$') and b_str.endswith('$$'):
                    p_eq = doc.add_paragraph()
                    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_eq.paragraph_format.space_before = Pt(8)
                    p_eq.paragraph_format.space_after = Pt(8)
                    try:
                        omml_elem = convert_latex_to_omml_element(b_str, is_display=True)
                        p_eq._p.append(omml_elem)
                    except Exception:
                        eq_text = b_str.strip('$').strip()
                        run_eq = p_eq.add_run(eq_text)
                        run_eq.font.name = 'Cambria Math'
                        run_eq.font.size = Pt(11)
                        run_eq.font.italic = True
                        run_eq.font.color.rgb = NAVY_COLOR
                    continue

                # Bullet points
                if b_str.startswith('- ') or b_str.startswith('* '):
                    p_bullet = doc.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_before = Pt(1)
                    p_bullet.paragraph_format.space_after = Pt(3)
                    p_bullet.paragraph_format.line_spacing = 1.15
                    
                    bullet_text = b_str[2:].strip()
                    tokens = re.split(r'(\$.*?\$|\*\*.*?\*\*|\*.*?\*|`.*?`)', bullet_text)
                    for tok in tokens:
                        if not tok:
                            continue
                        if tok.startswith('$') and tok.endswith('$'):
                            try:
                                omml_elem = convert_latex_to_omml_element(tok, is_display=False)
                                p_bullet._p.append(omml_elem)
                            except Exception:
                                r = p_bullet.add_run(tok.strip('$'))
                                r.font.name = 'Cambria Math'
                        elif tok.startswith('**') and tok.endswith('**'):
                            r = p_bullet.add_run(tok[2:-2])
                            r.font.name = 'Calibri'
                            r.font.size = Pt(11)
                            r.font.bold = True
                            r.font.color.rgb = DARK_COLOR
                        elif tok.startswith('*') and tok.endswith('*'):
                            r = p_bullet.add_run(tok[1:-1])
                            r.font.name = 'Calibri'
                            r.font.size = Pt(11)
                            r.font.italic = True
                        elif tok.startswith('`') and tok.endswith('`'):
                            r = p_bullet.add_run(tok[1:-1])
                            r.font.name = 'Consolas'
                            r.font.size = Pt(10)
                            r.font.color.rgb = NAVY_COLOR
                        else:
                            r = p_bullet.add_run(tok)
                            r.font.name = 'Calibri'
                            r.font.size = Pt(11)
                            r.font.color.rgb = BODY_COLOR
                    continue

                # References (Section 20 hanging indent)
                if 'references' in h_text.lower():
                    p_ref = doc.add_paragraph()
                    p_ref.paragraph_format.left_indent = Inches(0.3)
                    p_ref.paragraph_format.first_line_indent = Inches(-0.3)
                    p_ref.paragraph_format.space_before = Pt(2)
                    p_ref.paragraph_format.space_after = Pt(4)
                    r_ref = p_ref.add_run(b_str)
                    r_ref.font.name = 'Calibri'
                    r_ref.font.size = Pt(10)
                    r_ref.font.color.rgb = BODY_COLOR
                    continue

                # Standard Paragraph
                p_p = doc.add_paragraph()
                p_p.paragraph_format.space_before = Pt(2)
                p_p.paragraph_format.space_after = Pt(6)
                p_p.paragraph_format.line_spacing = 1.15

                tokens = re.split(r'(\$.*?\$|\*\*.*?\*\*|\*.*?\*|`.*?`)', b_str)
                for tok in tokens:
                    if not tok:
                        continue
                    if tok.startswith('$') and tok.endswith('$'):
                        try:
                            omml_elem = convert_latex_to_omml_element(tok, is_display=False)
                            p_p._p.append(omml_elem)
                        except Exception:
                            r = p_p.add_run(tok.strip('$'))
                            r.font.name = 'Cambria Math'
                    elif tok.startswith('**') and tok.endswith('**'):
                        r = p_p.add_run(tok[2:-2])
                        r.font.bold = True
                        r.font.color.rgb = DARK_COLOR
                    elif tok.startswith('*') and tok.endswith('*'):
                        r = p_p.add_run(tok[1:-1])
                        r.font.italic = True
                    elif tok.startswith('`') and tok.endswith('`'):
                        r = p_p.add_run(tok[1:-1])
                        r.font.name = 'Consolas'
                        r.font.size = Pt(10)
                        r.font.color.rgb = NAVY_COLOR
                    else:
                        r = p_p.add_run(tok)

                # Check figure/table maps for main sections
                for key in fig_map:
                    if key not in rendered_figures and key.lower() in h_text.lower():
                        png_f, cap_f = fig_map[key]
                        add_figure_image(doc, png_f, cap_f)
                        rendered_figures.add(key)
                for key in table_map:
                    if key not in rendered_tables and key.lower() in h_text.lower():
                        tbl_f, tbl_cap = table_map[key]
                        if tbl_f in table_files:
                            render_markdown_table(doc, table_files[tbl_f], tbl_cap, table_filename=tbl_f)
                            rendered_tables.add(key)

    # Ensure any remaining unrendered tables/figures are appended at their relevant locations
    for key, (png_f, cap_f) in fig_map.items():
        if key not in rendered_figures:
            add_figure_image(doc, png_f, cap_f)
            rendered_figures.add(key)

    for key, (tbl_f, tbl_cap) in table_map.items():
        if key not in rendered_tables and tbl_f in table_files:
            render_markdown_table(doc, table_files[tbl_f], tbl_cap, table_filename=tbl_f)
            rendered_tables.add(key)

    out_docx_path = os.path.join(base_dir, 'AU_DIC_Research_Paper_v1.docx')
    try:
        doc.save(out_docx_path)
        print(f"Successfully generated publication-ready manuscript at: {out_docx_path}")
    except PermissionError:
        alt_path = os.path.join(base_dir, 'AU_DIC_Research_Paper_v1_rebuilt.docx')
        try:
            doc.save(alt_path)
            print(f"File locked. Successfully generated manuscript at: {alt_path}")
        except PermissionError:
            import time
            alt_path2 = os.path.join(base_dir, f'AU_DIC_Research_Paper_v1_{int(time.time())}.docx')
            doc.save(alt_path2)
            print(f"Primary files locked. Saved manuscript to: {alt_path2}")

if __name__ == '__main__':
    create_journal_manuscript()
