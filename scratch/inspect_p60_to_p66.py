import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

doc = docx.Document(v5_docx)

print("=== PARAGRAPHS P60 TO P66 ===")

for i in range(60, min(67, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f"P{i:2d} (Style: {p.style.name}): '{p.text.strip()}'")
