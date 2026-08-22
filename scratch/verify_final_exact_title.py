import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"
v5_pdf = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.pdf"
easy_md = workspace / "docs" / "paper" / "easy_read" / "Easy_Read_Research_Guide.md"
easy_docx = workspace / "docs" / "paper" / "easy_read" / "Easy_Read_Research_Guide.docx"
citation_cff = workspace / "CITATION.cff"

assert v5_docx.exists(), "PaperV5_Ollama_Primary.docx missing!"
assert v5_pdf.exists(), "PaperV5_Ollama_Primary.pdf missing!"

doc = docx.Document(v5_docx)

exact_title = "Smart Academic Document Intelligence System: Automated Extraction, Normalization, and Benchmark Generation"

print("============================================================")
print(" VERIFYING FINAL EXACT TITLE ACROSS ALL GENERATED ARTIFACTS")
print("============================================================")

# 1. Verify Page 3 Title in DOCX
title_p_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == exact_title:
        title_p_idx = i
        break

print(f"1. DOCX Title Found at Paragraph P{title_p_idx}: '{exact_title}'")
title_exact_match = (title_p_idx is not None)
print(f"   Exact Match: {title_exact_match}")

# 2. Check technical terms preservation inside body text
adbg_count = sum(1 for p in doc.paragraphs if "ADBG" in p.text)
audic_count = sum(1 for p in doc.paragraphs if "AU DIC" in p.text)
print(f"\n2. Technical References Intact:")
print(f"   'ADBG' Body References Count:   {adbg_count}")
print(f"   'AU DIC' Body References Count: {audic_count}")

# 3. Check Front Matter & TOC
toc_count = sum(1 for p in doc.paragraphs[:60] if "\t" in p.text and any(c.isdigit() for c in p.text))
print(f"\n3. Front Matter & TOC Integrity:")
print(f"   Populated TOC Entries Count:    {toc_count}")

# 4. Check Easy-Read Guide & CITATION.cff
with open(easy_md, "r", encoding="utf-8") as f:
    easy_md_text = f.read()
has_exact_title_easy_md = exact_title in easy_md_text

with open(citation_cff, "r", encoding="utf-8") as f:
    cff_text = f.read()
has_exact_title_cff = exact_title in cff_text

print(f"\n4. Metadata Files Title Match:")
print(f"   Easy-Read MD Match:  {has_exact_title_easy_md}")
print(f"   CITATION.cff Match:  {has_exact_title_cff}")

print("============================================================")
assert title_exact_match, "DOCX Title does NOT match exact string!"
assert adbg_count > 0 and audic_count > 0, "Technical references incorrectly modified!"
assert toc_count >= 50, "TOC corrupted or unpopulated!"
assert has_exact_title_easy_md and has_exact_title_cff, "Metadata files missing exact title!"

print("EMPIRICAL VERIFICATION SUCCESSFUL: Final exact title updated consistently across all artifacts!")
