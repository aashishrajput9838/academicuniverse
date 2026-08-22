import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" VERIFYING SWAPPED PAGE ORDER IN PaperV5_Ollama_Primary.docx")
print("============================================================")

p0 = doc.paragraphs[0]
print(f"P0 (Page 1 Start): '{p0.text.strip()}'")

title_p = None
intro_p = None

for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("ADBG v1.0 & AU DIC Benchmark"):
        title_p = p
        title_idx = i
    elif p.text.strip() == "1. Introduction":
        intro_p = p
        intro_idx = i

assert title_p is not None, "Error: Title paragraph not found!"
assert intro_p is not None, "Error: 1. Introduction paragraph not found!"

print(f"Title Paragraph Index:        P{title_idx} (Page Break Before: {title_p.paragraph_format.page_break_before})")
print(f"1. Introduction Paragraph:    P{intro_idx} (Page Break Before: {intro_p.paragraph_format.page_break_before})")

# Order verification
page1_is_contents = (p0.text.strip() == "CONTENTS")
page2_is_title = (title_idx > 0 and title_p.paragraph_format.page_break_before is True)
page3_is_intro = (intro_idx > title_idx and intro_p.paragraph_format.page_break_before is True)

print(f"1. Page 1 is CONTENTS:                         {page1_is_contents}")
print(f"2. Page 2 starts with Title / Authors / Abstract: {page2_is_title}")
print(f"3. Page 3 starts with 1. Introduction:          {page3_is_intro}")

print("============================================================")
assert page1_is_contents, "Page 1 is not CONTENTS!"
assert page2_is_title, "Page 2 does not start with Title!"
assert page3_is_intro, "Page 3 does not start with Introduction!"

print("EMPIRICAL VERIFICATION SUCCESSFUL: Page order swapped cleanly!")
