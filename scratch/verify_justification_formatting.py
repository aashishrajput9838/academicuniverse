import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
justified_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary_Justified.docx"

assert justified_docx.exists(), "Justified docx missing!"

doc = docx.Document(justified_docx)

justified_count = 0
left_count = 0
center_count = 0

for p in doc.paragraphs:
    if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        justified_count += 1
    elif p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        center_count += 1
    elif p.alignment == WD_ALIGN_PARAGRAPH.LEFT:
        left_count += 1

print("============================================================")
print(" TYPOGRAPHY JUSTIFICATION VERIFICATION REPORT")
print("============================================================")
print(f"Total Paragraphs Inspected:       {len(doc.paragraphs)}")
print(f"Fully Justified Body Paragraphs:  {justified_count}")
print(f"Centered Headings/Titles/Shapes: {center_count}")
print(f"Left-Aligned Elements/Code:       {left_count}")
print(f"Table Count:                      {len(doc.tables)}")
print(f"Inline Shape Diagrams:            {len(doc.inline_shapes)}")
print("============================================================")
