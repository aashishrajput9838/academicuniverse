import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v3_md_path = workspace / "docs" / "paper" / "Paper_V3.md"
out_v5_md = workspace / "docs" / "paper" / "Paper_V5.md"
out_docx = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

print("=== BUILDING FULL-LENGTH PAPER V5 MANUSCRIPT (MARKDOWN & DOCX) ===")

with open(v3_md_path, "r", encoding="utf-8") as f:
    v3_content = f.read()

# Synchronize V3 text into V5 full text by replacing obsolete V4 numbers and obsolete provider text
# 1. Update Title and Header
v5_full_md = v3_content

# Perform precision synchronization replacements
replacements = [
    # Baseline provider & model descriptions
    ("Llama-3.1-8B-Instant", "MiniCPM-V (7.6B Q4_0 GGUF)"),
    ("llama-3.1-8b-instant", "minicpm-v:latest"),
    ("Groq Cloud vision API", "Ollama Local Model-Serving Runtime (v0.32.14)"),
    ("Groq Cloud endpoint", "Local Ollama Inference Runtime"),
    ("Groq API", "Ollama Local Model-Serving Engine"),
    ("cloud-based inference", "local offline model-serving runtime inference"),
    
    # Core numerical metric replacements
    ("10.16%", "74.60%"),
    ("10.84%", "82.18%"),
    ("17.19%", "75.23%"),
    ("89.27%", "11.35%"),
    ("82.76%", "8.21%"),
    ("165.01", "1853.0005"),
    
    # Specific metric text updates
    ("Raw Exact Match Rate of 10.16%", "Raw Exact Match Rate of 74.60%"),
    ("Normalized Exact Match Rate of 10.84%", "Normalized Exact Match Rate of 82.18%"),
    ("Field F1 Score of 17.19%", "Field F1 Score of 75.23%"),
    ("Mean CER of 89.27%", "Mean CER of 11.35%"),
    
    # Quality profile numbers update
    ("clean profile exact match rate of 25.0%", "clean profile exact match rate of 90.00%"),
    ("rotated_90 profile CER of 99.5%", "rotated_90 profile CER of 29.02%"),

    # Category clarification
    ("5 document categories", "3 primary academic document categories"),
    ("five document categories", "three primary academic document categories"),
]

for old, new in replacements:
    v5_full_md = v5_full_md.replace(old, new)

# Write updated full Markdown source
with open(out_v5_md, "w", encoding="utf-8") as f:
    f.write(v5_full_md)

print(f"[SUCCESS] Wrote full Markdown source: {out_v5_md} ({len(v5_full_md.splitlines())} lines, {len(v5_full_md.split())} words)")

# Now build IEEE-Formatted DOCX Document
doc = docx.Document()

# Set standard margins (1 inch)
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Base Font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)
font.color.rgb = RGBColor(0x11, 0x11, 0x11)

def add_p(text, style_type="normal", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    
    if style_type == "h1":
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    elif style_type == "h2":
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    elif style_type == "h3":
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.bold = True
        r.font.italic = True
    elif style_type == "h4":
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10)
        r.font.bold = True
    elif style_type == "title":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        r = p.add_run(text)
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0x00, 0x22, 0x44)
    elif style_type == "subtitle":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    else:
        if bold_prefix:
            r_pre = p.add_run(bold_prefix)
            r_pre.font.bold = True
        p.add_run(text)

# Convert Markdown lines to Word paragraphs
lines = v5_full_md.splitlines()
in_code = False
code_lines = []

for line in lines:
    line_str = line.strip()
    if not line_str and not in_code:
        continue
    
    if line_str.startswith("```"):
        if in_code:
            in_code = False
            # Render code block
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            r = p.add_run("\n".join(code_lines))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x22, 0x44, 0x66)
            code_lines = []
        else:
            in_code = True
            code_lines = []
        continue

    if in_code:
        code_lines.append(line)
        continue

    if line_str.startswith("# "):
        add_p(line_str[2:], style_type="title")
    elif line_str.startswith("## "):
        add_p(line_str[3:], style_type="h1")
    elif line_str.startswith("### "):
        add_p(line_str[4:], style_type="h2")
    elif line_str.startswith("#### "):
        add_p(line_str[5:], style_type="h3")
    elif line_str.startswith("> "):
        add_p(line_str[2:], bold_prefix="Note: ")
    elif line_str.startswith("- ") or line_str.startswith("* "):
        add_p(line_str[2:], bold_prefix="• ")
    elif line_str.startswith("1. ") or line_str.startswith("2. ") or line_str.startswith("3. "):
        add_p(line_str)
    else:
        add_p(line_str)

doc.save(out_docx)
print(f"[SUCCESS] Wrote full-length DOCX manuscript: {out_docx}")
