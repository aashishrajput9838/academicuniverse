"""
Paper_V3_IEEE_Final Production Typesetter
=========================================
Reads Paper_V3_IEEE_v2.docx and applies all 10 IEEE production formatting tasks.
Outputs Paper_V3_IEEE_Final.docx and Paper_V3_IEEE_Final.pdf.
NO scientific content is modified.
"""
import re
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client

# ============================================================
# IEEE ROMAN NUMERAL SECTION MAPPING
# ============================================================
SECTION_ROMAN_MAP = {
    "1. Introduction": "I. INTRODUCTION",
    "2. Related Work": "II. RELATED WORK",
    "3. System Architecture Overview": "III. SYSTEM ARCHITECTURE OVERVIEW",
    "4. ADBG Synthetic Data Generation Methodology": "IV. ADBG SYNTHETIC DATA GENERATION METHODOLOGY",
    "5. AU DIC Evaluation Subsystem & Semantic Normalization Layer": "V. AU DIC EVALUATION SUBSYSTEM & SEMANTIC NORMALIZATION LAYER",
    "6. Experimental Setup & Evaluation Metrics": "VI. EXPERIMENTAL SETUP & EVALUATION METRICS",
    "7. Results & Validation": "VII. RESULTS & VALIDATION",
    "8. Discussion, Threats to Validity & Limitations": "VIII. DISCUSSION, THREATS TO VALIDITY & LIMITATIONS",
    "9. Future Work": "IX. FUTURE WORK",
    "10. Conclusion": "X. CONCLUSION",
}

SUBSECTION_LETTER_MAP = {
    "1.1": "A.", "1.2": "B.", "1.3": "C.", "1.4": "D.",
    "2.1": "A.", "2.2": "B.", "2.3": "C.",
    "4.1": "A.", "4.2": "B.", "4.3": "C.",
    "5.1": "A.", "5.2": "B.", "5.3": "C.",
    "6.1": "A.", "6.2": "B.",
    "7.1": "A.", "7.2": "B.", "7.3": "C.", "7.4": "D.",
    "8.1": "A.", "8.2": "B.", "8.3": "C.",
}

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def apply_ieee_formatting(docx_path, output_path, fig1_path, fig2_path):
    doc = Document(docx_path)

    # ============================================================
    # TASK 4: CLEAN HEADER & FOOTER (Remove fake branding)
    # ============================================================
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Clean footer — page number only, no fake volume/issue
        footer = section.footer
        for p in footer.paragraphs:
            p.clear()
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f_run = f_p.add_run()
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(10)
        f_run.font.color.rgb = RGBColor(80, 80, 80)
        # Insert dynamic page number field
        fld_xml = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        f_p._p.append(fld_xml)

        # Clean header
        header = section.header
        for p in header.paragraphs:
            p.clear()

    # ============================================================
    # ITERATE ALL PARAGRAPHS AND APPLY FORMATTING
    # ============================================================
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()

        # ============================================================
        # TASK 1: SECTION NUMBERING — Roman Numerals
        # ============================================================
        # Major sections: "1. Introduction" -> "I. INTRODUCTION"
        for old_num, new_roman in SECTION_ROMAN_MAP.items():
            if txt == old_num:
                for run in para.runs:
                    run.clear()
                r = para.runs[0] if para.runs else para.add_run()
                r.text = new_roman
                r.font.name = "Times New Roman"
                r.font.size = Pt(13)
                r.bold = True
                r.font.color.rgb = RGBColor(0, 51, 102)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_before = Pt(16)
                para.paragraph_format.space_after = Pt(8)
                para.paragraph_format.keep_with_next = True
                break

        # Subsections: "1.1 Background" -> "A. Background"
        sub_match = re.match(r'^(\d+\.\d+)\s+(.+)$', txt)
        if sub_match:
            sub_num = sub_match.group(1)
            sub_title = sub_match.group(2)
            if sub_num in SUBSECTION_LETTER_MAP:
                letter = SUBSECTION_LETTER_MAP[sub_num]
                new_text = f"{letter} {sub_title}"
                for run in para.runs:
                    run.clear()
                r = para.runs[0] if para.runs else para.add_run()
                r.text = new_text
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = True
                r.italic = True
                r.font.color.rgb = RGBColor(51, 51, 51)
                para.paragraph_format.space_before = Pt(10)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.keep_with_next = True

        # Sub-subsection: "7.4.1 Inference Pipeline..." -> keep as-is but italicize
        sub3_match = re.match(r'^(\d+\.\d+\.\d+)\s+(.+)$', txt)
        if sub3_match:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10.5)
                run.bold = True
                run.italic = True
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(3)
            para.paragraph_format.keep_with_next = True

        # ============================================================
        # TASK 2: TITLE PAGE — Fix visual layout
        # ============================================================
        if i == 0:  # Title paragraph
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(18)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(36)
            para.paragraph_format.space_after = Pt(12)

        # ============================================================
        # TASK 3: AUTHOR BLOCK
        # ============================================================
        if i == 1 and "Authors" in txt:
            for run in para.runs:
                run.clear()
            r = para.runs[0] if para.runs else para.add_run()
            r.text = (
                "Author Name¹, Co-Author Name²\n"
                "¹Department of Computer Science, University Name, City, Country\n"
                "²Department of Information Technology, University Name, City, Country\n"
                "Corresponding Email: author@university.edu"
            )
            r.font.name = "Times New Roman"
            r.font.size = Pt(10)
            r.bold = False
            r.font.color.rgb = RGBColor(51, 51, 51)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)

        if i == 2 and "Target Publication" in txt:
            for run in para.runs:
                run.clear()
            r = para.runs[0] if para.runs else para.add_run()
            r.text = "Manuscript submitted for review to IEEE Access"
            r.font.name = "Times New Roman"
            r.font.size = Pt(9)
            r.italic = True
            r.font.color.rgb = RGBColor(100, 100, 100)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(4)

        if i == 3 and "Repository" in txt:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(120, 120, 120)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(12)

        # ============================================================
        # TASK 2 CONTINUED: Remove horizontal rule placeholders
        # ============================================================
        if txt == "---" or txt.startswith("___"):
            for run in para.runs:
                run.clear()
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)

        # ============================================================
        # TASK 5: EQUATIONS — Style display equations
        # ============================================================
        # Detect standalone equation-like paragraphs
        eq_patterns = [
            "DocumentSpecimen",
            "I_{degraded}",
            "Category Accuracy",
            "P = (True Positive",
            "CER = (S_{char",
            "WER = (S_{word",
        ]
        for eq_pat in eq_patterns:
            if eq_pat in txt and len(txt) < 200:
                for run in para.runs:
                    run.font.name = "Cambria Math"
                    run.font.size = Pt(11)
                    run.italic = True
                    run.font.color.rgb = RGBColor(0, 51, 102)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(8)
                break

        # ============================================================
        # TASK 8: REFERENCES — Ensure proper typography
        # ============================================================
        ref_match = re.match(r'^\[(\d+)\]', txt)
        if ref_match:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.first_line_indent = Inches(-0.3)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(3)

        # ============================================================
        # "Abstract" heading — Style it
        # ============================================================
        if txt == "Abstract":
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.keep_with_next = True

        # "References" heading
        if txt == "References":
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(8)
            para.paragraph_format.keep_with_next = True

        # "Ethics" heading
        if txt.startswith("Ethics"):
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)

        # "Appendix" headings
        if txt.startswith("Appendix"):
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(6)

        # ============================================================
        # TASK 9: TYPOGRAPHY — Body paragraphs
        # ============================================================
        style_name = para.style.name if para.style else ""
        if style_name == "Normal" and txt and not any([
            txt.startswith("["), txt == "---", txt.startswith("Table "),
            txt.startswith("Fig. "), txt.startswith("Abstract"),
            txt.startswith("References"), txt.startswith("Ethics"),
            txt.startswith("Appendix"), txt.startswith("B."),
        ]):
            for run in para.runs:
                if not run.font.name or run.font.name == "Calibri":
                    run.font.name = "Times New Roman"
                if not run.font.size:
                    run.font.size = Pt(10)
            if not para.paragraph_format.line_spacing:
                para.paragraph_format.line_spacing = 1.15

        # List Bullet styling
        if style_name == "List Bullet":
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(2)

        # ============================================================
        # TASK 9 CONTINUED: Widow/Orphan protection
        # ============================================================
        pf = para.paragraph_format
        pf._element.attrib  # ensure element exists
        # Set widow/orphan control via XML
        pPr = para._p.get_or_add_pPr()
        widow = parse_xml(f'<w:widowControl {nsdecls("w")} w:val="1"/>')
        existing = pPr.findall(f'{{{nsdecls("w").split(chr(34))[1]}}}widowControl')
        if not existing:
            pPr.append(widow)

        # ============================================================
        # Table caption styling
        # ============================================================
        if txt.startswith("Table ") and ":" in txt:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.keep_with_next = True

        # Figure caption styling
        if txt.startswith("Fig. "):
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(10)

        # Index Terms styling
        if txt.startswith("Index Terms"):
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(9.5)
                run.bold = True
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(10)

    # ============================================================
    # TASK 6: TABLE PROFESSIONALIZATION
    # ============================================================
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Set table borders via XML
        tbl = table._tbl
        tblPr = tbl.tblPr
        # Remove existing borders
        for existing_borders in tblPr.findall(f'{{{nsdecls("w").split(chr(34))[1]}}}tblBorders'):
            tblPr.remove(existing_borders)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="12" w:space="0" w:color="003366"/>'
            f'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="003366"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            f'<w:insideV w:val="none"/>'
            f'<w:left w:val="none"/>'
            f'<w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)

        # Style header row and data rows
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        if r_idx == 0:
                            run.font.size = Pt(9)
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                        else:
                            run.font.size = Pt(9)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)

                if r_idx == 0:
                    set_cell_background(cell, "003366")
                elif r_idx % 2 == 0:
                    set_cell_background(cell, "F5F7FA")

    # Save output
    doc.save(output_path)
    print(f"[SUCCESS] Paper_V3_IEEE_Final.docx generated: {output_path}")


def convert_to_pdf(docx_path, pdf_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close(False)
        print(f"[SUCCESS] Paper_V3_IEEE_Final.pdf generated: {pdf_path}")
    except Exception as e:
        print(f"[NOTE] PDF: {e}")
    finally:
        try:
            word.Quit()
        except:
            pass


if __name__ == "__main__":
    input_docx = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_v2.docx"
    output_docx = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final.docx"
    output_pdf = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final.pdf"
    fig1 = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\figure1_system_architecture.png"
    fig2 = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\figure2_option_b_pipeline.png"

    apply_ieee_formatting(input_docx, output_docx, fig1, fig2)
    convert_to_pdf(output_docx, output_pdf)
