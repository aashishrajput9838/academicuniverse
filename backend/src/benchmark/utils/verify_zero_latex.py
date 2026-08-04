import docx

def verify_docx():
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Final_Corrected.docx"
    doc = docx.Document(docx_path)
    
    latex_targets = [
        r'\mu', r'\lambda', r'\mathcal', r'\frac', r'\sum', r'\text{', 
        r'\xrightarrow', r'\cdot', r'\bar', r'\epsilon', r'\Delta', 
        r'_{', r'^{', r'\left', r'\right'
    ]
    
    findings = {target: 0 for target in latex_targets}
    
    # Scan Paragraphs
    for p in doc.paragraphs:
        for target in latex_targets:
            findings[target] += p.text.count(target)
            
    # Scan Tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for target in latex_targets:
                    findings[target] += cell.text.count(target)
                    
    total_remaining = sum(findings.values())
    print("===========================================================")
    print("             LATEX REMOVAL VERIFICATION SCANNER            ")
    print("===========================================================")
    for target, count in findings.items():
        print(f"Target '{target}': {count} occurrences remaining")
    print("-----------------------------------------------------------")
    print(f"TOTAL RAW LATEX COMMANDS REMAINING: {total_remaining}")
    print("===========================================================")
    
    return total_remaining

if __name__ == '__main__':
    verify_docx()
