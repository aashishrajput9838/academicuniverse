import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

print("============================================================")
print(" EMPIRICAL VERIFICATION OF GENERATED PaperV5_Ollama_Primary.docx")
print("============================================================")

contents_idx = None
intro_idx = None
abstract_idx = None
index_terms_idx = None
toc_entries_found = []

for i, p in enumerate(doc.paragraphs):
    text_str = p.text.strip()
    if text_str == "Abstract":
        abstract_idx = i
    elif text_str.startswith("Index Terms"):
        index_terms_idx = i
    elif text_str == "CONTENTS":
        contents_idx = i
    elif text_str == "1. Introduction":
        intro_idx = i
    
    if contents_idx is not None and intro_idx is None and text_str != "CONTENTS" and text_str:
        toc_entries_found.append((i, text_str))

print(f"A. Abstract Paragraph Index:         P{abstract_idx}")
print(f"B. Index Terms Paragraph Index:      P{index_terms_idx}")
print(f"C. CONTENTS Paragraph Index:         P{contents_idx}")
print(f"D. 1. Introduction Paragraph Index:  P{intro_idx}")
print(f"E. Total Visible TOC Entry Paragraphs Found: {len(toc_entries_found)}")
print("\nFirst 10 Visible TOC Entries in DOCX:")
for idx, text in toc_entries_found[:10]:
    print(f"   P{idx}: '{text}'")

h1_toc_count = sum(1 for _, t in toc_entries_found if not t.startswith("."))
h2_toc_count = sum(1 for _, t in toc_entries_found if t.startswith("."))

print(f"\nF. Number of TOC Level-1 Entries: {h1_toc_count}")
print(f"G. Number of TOC Level-2 Entries: {h2_toc_count}")

# Error check
error_found = any("No table of contents entries found" in p.text for p in doc.paragraphs)
print(f"H. 'Error. No table of contents entries found.' Absent: {not error_found}")

# Position check
position_correct = (index_terms_idx < contents_idx < intro_idx)
print(f"I. TOC Placement Correct (Index Terms < CONTENTS < 1. Introduction): {position_correct}")

print("============================================================")
assert not error_found, "Error message found in DOCX!"
assert position_correct, "TOC position incorrect!"
assert len(toc_entries_found) >= 30, f"Expected >= 30 TOC entries, found {len(toc_entries_found)}"
print("EMPIRICAL VERIFICATION SUCCESSFUL: Generated DOCX contains populated TOC!")
