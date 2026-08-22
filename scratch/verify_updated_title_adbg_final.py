import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

assert v5_docx_path.exists(), "PaperV5_Ollama_Primary.docx missing!"

doc = docx.Document(v5_docx_path)

expected_title = "Smart Academic Document Intelligence System with the Academic Document Benchmark Generator for Automated Academic Document Extraction and Normalization"

print("============================================================")
print(" VERIFYING UPDATED TITLE IN PaperV5_Ollama_Primary.docx")
print("============================================================")

title_p = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == expected_title:
        title_p = p
        title_idx = i
        break

assert title_p is not None, f"Error: Expected title '{expected_title}' not found in document!"

print(f"1. New Title Paragraph Found at P{title_idx}: '{title_p.text.strip()}'")
print(f"2. Page Break Before on Title Paragraph: {title_p.paragraph_format.page_break_before} (PAGE 3)")

# Forbidden terms check in title paragraph
forbidden_terms = ["AU DIC", "DIC", "Reproducible", "Evaluation and Normalization Framework"]
forbidden_found = [term for term in forbidden_terms if term in title_p.text]

print(f"3. Forbidden Terms in Title:              {forbidden_found} (Should be empty [])")
print(f"4. Contains 'Academic Document Benchmark Generator': {'Academic Document Benchmark Generator' in title_p.text}")

# Page order verification
p0 = doc.paragraphs[0]
has_logo = "<w:drawing>" in p0._p.xml or "graphic" in p0._p.xml
contents_p = doc.paragraphs[1]
has_contents = (contents_p.text.strip() == "CONTENTS" and contents_p.paragraph_format.page_break_before is True)

print(f"5. Page 1 Sharda University Logo Present: {has_logo}")
print(f"6. Page 2 CONTENTS Header Present:        {has_contents}")

print("============================================================")
assert len(forbidden_found) == 0, "Forbidden terms found in title!"
assert "Academic Document Benchmark Generator" in title_p.text, "ADBG expansion missing!"
assert has_logo, "Logo missing!"
assert has_contents, "CONTENTS header missing!"

print("EMPIRICAL VERIFICATION SUCCESSFUL: Title with Academic Document Benchmark Generator updated perfectly!")
