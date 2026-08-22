import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_RefHeader.docx"

print("=== TESTING REFERENCES HEADING FIX ===")

doc = docx.Document(v4_docx_path)

# 1. Find misplaced P231 "References"
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "References" and i < 240:
        print(f"Clearing misplaced References header at P{i}")
        p.text = ""

# 2. Find reference [1] entry and insert REFERENCES header immediately before it
ref1_found = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("[1] "):
        print(f"Found reference [1] at P{i}. Inserting REFERENCES header immediately before it...")
        p_hdr = p.insert_paragraph_before("REFERENCES")
        p_hdr.paragraph_format.space_before = Pt(14)
        p_hdr.paragraph_format.space_after = Pt(6)
        if p_hdr.runs:
            r = p_hdr.runs[0]
            r.font.name = "Times New Roman"
            r.font.size = Pt(13)
            r.bold = True
        ref1_found = True
        break

assert ref1_found, "Error: Could not find reference [1] entry!"

doc.save(test_docx_path)
print(f"[SUCCESS] Saved test docx with REFERENCES heading fix to {test_docx_path.name}")
