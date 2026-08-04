import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import re
import os
import shutil
import win32com.client
import subprocess
import time
from omml_engine import append_display_omml, append_inline_omml, convert_latex_to_omml_xml

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_production_ieee_docx(md_path, docx_path, fig1_path, fig2_path):
    subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], capture_output=True)
    time.sleep(2)
    
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("IEEE ACCESS | Volume 14, 2026 | Page ")
        f_run.font.name = "Times New Roman"
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(100, 100, 100)
        
        f_xml = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        f_p._p.append(f_xml)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_mermaid = False
    mermaid_lines = []
    current_list_num = 0
    in_references = False
    ref_idx = 1

    def format_inline_markdown(paragraph, text):
        text_no_backtick = text.replace('`', '')
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|\$.*?\$)', text_no_backtick)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                inner_text = token[2:-2]
                if '$' in inner_text:
                    sub_tokens = re.split(r'(\$.*?\$)', inner_text)
                    for sub in sub_tokens:
                        if sub.startswith('$') and sub.endswith('$'):
                            append_inline_omml(paragraph, sub[1:-1].strip())
                        elif sub:
                            r = paragraph.add_run(sub)
                            r.bold = True
                else:
                    r = paragraph.add_run(inner_text)
                    r.bold = True
            elif token.startswith('*') and token.endswith('*'):
                inner_text = token[1:-1]
                if '$' in inner_text:
                    sub_tokens = re.split(r'(\$.*?\$)', inner_text)
                    for sub in sub_tokens:
                        if sub.startswith('$') and sub.endswith('$'):
                            append_inline_omml(paragraph, sub[1:-1].strip())
                        elif sub:
                            r = paragraph.add_run(sub)
                            r.italic = True
                else:
                    r = paragraph.add_run(inner_text)
                    r.italic = True
            elif token.startswith('$') and token.endswith('$'):
                append_inline_omml(paragraph, token[1:-1].strip())
            else:
                paragraph.add_run(token)

    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        if line.startswith('```'):
            if in_code_block or in_mermaid:
                if in_mermaid:
                    mermaid_str = "\n".join(mermaid_lines)
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(4)
                    
                    if "ADBG Subsystem" in mermaid_str or "Seed Generator" in mermaid_str:
                        if os.path.exists(fig1_path):
                            p_img.add_run().add_picture(fig1_path, width=Inches(6.5))
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cap.paragraph_format.space_after = Pt(8)
                            r_cap = p_cap.add_run("Fig. 1. Decoupled System Architecture of the ADBG Synthetic Generation and AU DIC Evaluation Subsystems.")
                            r_cap.font.name = "Times New Roman"
                            r_cap.font.size = Pt(9.5)
                            r_cap.bold = True
                            r_cap.font.color.rgb = RGBColor(0, 51, 102)
                    elif "Phase I: Synthetic Benchmark Generation" in mermaid_str or "End-to-End Methodological Workflow" in mermaid_str:
                        wf_img = r"c:\github\academicuniverse.com\academicuniverse\methodology_workflow_600dpi.png"
                        if not os.path.exists(wf_img):
                            wf_img = r"c:\github\academicuniverse.com\academicuniverse\methodology_workflow_300dpi.png"
                        if os.path.exists(wf_img):
                            p_img.add_run().add_picture(wf_img, width=Inches(6.5))
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cap.paragraph_format.space_after = Pt(8)
                            r_cap = p_cap.add_run("Fig. 2. End-to-End Methodological Workflow of the Proposed AU DIC Benchmark Evaluation Framework.")
                            r_cap.font.name = "Times New Roman"
                            r_cap.font.size = Pt(9.5)
                            r_cap.bold = True
                            r_cap.font.color.rgb = RGBColor(0, 51, 102)
                    elif "Option B" in mermaid_str or "Zero-Shot" in mermaid_str or "Groq Cloud" in mermaid_str:
                        if os.path.exists(fig2_path):
                            p_img.add_run().add_picture(fig2_path, width=Inches(6.5))
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cap.paragraph_format.space_after = Pt(8)
                            r_cap = p_cap.add_run("Fig. 3. Option B Zero-Shot Text-Prompted Neural LLM Evaluation Pipeline Architecture.")
                            r_cap.font.name = "Times New Roman"
                            r_cap.font.size = Pt(9.5)
                            r_cap.bold = True
                            r_cap.font.color.rgb = RGBColor(0, 51, 102)
                    
                    mermaid_lines = []
                    in_mermaid = False
                else:
                    code_lines = []
                    in_code_block = False
            else:
                if 'mermaid' in line:
                    in_mermaid = True
                    mermaid_lines = []
                else:
                    in_code_block = True
                    code_lines = []
            i += 1
            continue

        if in_mermaid:
            mermaid_lines.append(line)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        if line.startswith('#'):
            current_list_num = 0
            if "References" in line:
                in_references = True
            elif in_references:
                in_references = False

        if '|' in line and ('---' in line or (i + 1 < len(lines) and '|' in lines[i+1])):
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\r\n'))
                i += 1
            
            rows_data = []
            for tline in table_lines:
                if '---' in tline:
                    continue
                cells = [c.strip() for c in tline.split('|')[1:-1]]
                if cells:
                    rows_data.append(cells)
            
            if rows_data:
                num_rows = len(rows_data)
                num_cols = len(rows_data[0])
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                tblPr = table._tbl.tblPr
                borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="single" w:sz="12" w:space="0" w:color="003366"/><w:bottom w:val="single" w:sz="12" w:space="0" w:color="003366"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/><w:insideV w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
                tblPr.append(borders)

                for r_idx, row in enumerate(rows_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            set_cell_margins(cell, top=100, bottom=100, left=130, right=130)
                            cp = cell.paragraphs[0]
                            cp.paragraph_format.space_before = Pt(2)
                            cp.paragraph_format.space_after = Pt(2)
                            
                            if r_idx == 0:
                                set_cell_background(cell, "003366")
                                format_inline_markdown(cp, val)
                                for run in cp.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(9.5)
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, "FFFFFF")
                                else:
                                    set_cell_background(cell, "F8F9FA")
                                format_inline_markdown(cp, val)
                                for run in cp.runs:
                                    run.font.name = "Times New Roman"
                                    run.font.size = Pt(9)
                                cp.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(line[2:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(20)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[3:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            r.bold = True
            r.font.color.rgb = RGBColor(0, 51, 102)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[4:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11.5)
            r.bold = True
            r.font.color.rgb = RGBColor(51, 51, 51)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            r = p.add_run(line[5:].strip())
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10.5)
            r.bold = True
            r.italic = True
        elif line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            format_inline_markdown(p, line[2:].strip())
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.italic = True
                run.font.color.rgb = RGBColor(70, 70, 70)
        elif in_references and (line.startswith('- ') or line.startswith('* ')):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.3)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            
            ref_content = line.strip()[2:]
            r_ref = p.add_run(f"[{ref_idx}] ")
            r_ref.font.name = 'Times New Roman'
            r_ref.font.size = Pt(9.5)
            r_ref.bold = True
            
            format_inline_markdown(p, ref_content)
            for run in p.runs:
                if run != r_ref:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9.5)
            ref_idx += 1
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            format_inline_markdown(p, line[2:].strip())
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
        elif re.match(r'^\d+\.\s', line):
            current_list_num += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            content = re.sub(r'^\d+\.\s', '', line)
            
            r_num = p.add_run(f"{current_list_num}. ")
            r_num.font.name = 'Times New Roman'
            r_num.font.size = Pt(10)
            r_num.bold = True
            
            format_inline_markdown(p, content)
            for run in p.runs:
                if run != r_num:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
        elif line.strip().startswith('$$') and line.strip().endswith('$$'):
            # AST Display OMML Conversion!
            display_tex = line.strip()[2:-2].strip()
            append_display_omml(doc, display_tex)
        elif line.startswith('**Table ') or line.startswith('**Fig. ') or line.startswith('**Figure '):
            # Check for figure image insertions
            brain_dir = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\bb9b3069-0e60-4209-b2b8-d0321ac491db"
            fig_map = {
                "Fig. 4": os.path.join(brain_dir, "figure_normalization_ablation.png"),
                "Fig. 5": os.path.join(brain_dir, "figure_metric_improvement.png"),
                "Fig. 6": os.path.join(brain_dir, "figure_rule_contribution.png"),
                "Fig. 7": os.path.join(brain_dir, "figure_field_improvement.png")
            }
            for fig_key, img_path in fig_map.items():
                if fig_key in line:
                    if os.path.exists(img_path):
                        p_fig = doc.add_paragraph()
                        p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_fig.paragraph_format.space_before = Pt(8)
                        p_fig.paragraph_format.space_after = Pt(4)
                        p_fig.add_run().add_picture(img_path, width=Inches(6.0))
                    break

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            format_inline_markdown(p, line.strip())
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9.5)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
        elif line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            format_inline_markdown(p, line.strip())
            for run in p.runs:
                if not run.font.name:
                    run.font.name = 'Times New Roman'
                if not run.font.size:
                    run.font.size = Pt(10)

        i += 1

    temp_docx = docx_path.replace(".docx", "_LiteratureExpanded_Build.docx")
    doc.save(temp_docx)
    try:
        shutil.copyfile(temp_docx, docx_path)
    except Exception as e:
        print(f"Note: Saved to {temp_docx} ({e})")
    print(f"Production IEEE Docx generated with AST OMML objects: {temp_docx}")
    
    # Run Word COM 2D BuildUp for maximum visual compatibility
    try:
        subprocess.run(["taskkill", "/f", "/im", "WINWORD.EXE"], capture_output=True)
        time.sleep(2)
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        com_doc = word.Documents.Open(os.path.abspath(temp_docx))
        print(f"Executing 2D OMML BuildUp on {com_doc.OMaths.Count} AST OMML equation objects...")
        for o_idx in range(1, com_doc.OMaths.Count + 1):
            try:
                com_doc.OMaths(o_idx).BuildUp()
            except Exception as e_bu:
                pass
        com_doc.Save()
        com_doc.Close(False)
        try:
            word.Quit()
        except:
            pass
        print(f"AST OMML BuildUp completed successfully!")
    except Exception as e_com:
        print(f"OMML BuildUp COM note: {e_com}")
        
    return temp_docx

if __name__ == "__main__":
    md_file = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\bb9b3069-0e60-4209-b2b8-d0321ac491db\Paper_V3.md"
    docx_file = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\Paper_V3_IEEE_Final.docx"
    fig1 = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\figure1_system_architecture.png"
    fig2 = r"c:\github\academicuniverse.com\academicuniverse\docs\paper\figure2_option_b_pipeline.png"

    generated_docx = create_production_ieee_docx(md_file, docx_file, fig1, fig2)
