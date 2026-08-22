import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"

doc = docx.Document(v4_docx_path)

print("=== PARAGRAPHS 225 TO 265 IN BASELINE DOCX ===")
for i in range(225, min(265, len(doc.paragraphs))):
    p = doc.paragraphs[i]
    print(f"P{i} [{p.style.name}] (align={p.alignment}): '{p.text[:90]}'")
