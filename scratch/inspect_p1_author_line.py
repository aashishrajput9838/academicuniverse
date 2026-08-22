import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

print("=== INSPECTING AUTHOR PARAGRAPH P1 ===")

doc4 = docx.Document(v4_docx_path)
p1_4 = doc4.paragraphs[1]
print(f"V4 P1 Alignment: {p1_4.alignment}")
print(f"V4 P1 Style:     {p1_4.style.name}")
print(f"V4 P1 Text:      '{p1_4.text[:120]}'")

doc5 = docx.Document(v5_docx_path)
p1_5 = doc5.paragraphs[1]
print(f"\nV5 P1 Alignment: {p1_5.alignment}")
print(f"V5 P1 Style:     {p1_5.style.name}")
print(f"V5 P1 Text:      '{p1_5.text[:120]}'")
