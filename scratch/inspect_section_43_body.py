import docx
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v5_docx_path = workspace / "docs" / "paper" / "PaperV5_Ollama_Primary.docx"

doc = docx.Document(v5_docx_path)

print("=== SEARCHING BODY PARAGRAPHS FOR SECTION 4.3 ===")

sec43_found = False
for i, p in enumerate(doc.paragraphs):
    text_str = p.text.strip()
    if text_str.startswith("4.3 Mathematical Formulation"):
        if i > 50: # Body text
            sec43_found = True
            print(f"Body P{i}: '{text_str}'")
            # Print next 20 paragraphs
            for j in range(i+1, min(i+25, len(doc.paragraphs))):
                print(f"  P{j:3d}: '{doc.paragraphs[j].text.strip()}'")
            break
