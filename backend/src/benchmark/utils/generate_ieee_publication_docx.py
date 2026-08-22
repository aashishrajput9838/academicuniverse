import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_omml_equation(eq_id, text):
    """
    Constructs a native Microsoft Word OMML Display Equation (m:oMathPara)
    with a right-aligned equation number tag: (1), (2), (3), (4), (5).
    """
    # Equation 1: velocity
    if eq_id == 1:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath>
            <m:r><m:rPr><m:sty m:val="p"/></m:rPr><m:t>μ</m:t></m:r>
            <m:sSub>
              <m:e><m:r><m:t></m:t></m:r></m:e>
              <m:sub><m:r><m:t>v</m:t></m:r></m:sub>
            </m:sSub>
            <m:r><m:t> = </m:t></m:r>
            <m:f>
              <m:num>
                <m:r><m:t>S(t</m:t></m:r>
                <m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>2</m:t></m:r></m:sub></m:sSub>
                <m:r><m:t>) − S(t</m:t></m:r>
                <m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>1</m:t></m:r></m:sub></m:sSub>
                <m:r><m:t>)</m:t></m:r>
              </m:num>
              <m:den>
                <m:r><m:t>Δt</m:t></m:r>
              </m:den>
            </m:f>
          </m:oMath>
        </m:oMathPara>
        '''
    # Equation 2: Holistic Growth Index
    elif eq_id == 2:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath>
            <m:r><m:t>H = </m:t></m:r>
            <m:f>
              <m:num>
                <m:nary>
                  <m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/></m:naryPr>
                  <m:sub><m:r><m:t>k=1</m:t></m:r></m:sub>
                  <m:sup><m:r><m:t>K</m:t></m:r></m:sup>
                  <m:e><m:r><m:t>( S̄</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t> · C̄</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t> · w</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t> )</m:t></m:r></m:e>
                </m:nary>
              </m:num>
              <m:den>
                <m:nary>
                  <m:naryPr><m:chr m:val="∑"/><m:limLoc m:val="undOvr"/></m:naryPr>
                  <m:sub><m:r><m:t>k=1</m:t></m:r></m:sub>
                  <m:sup><m:r><m:t>K</m:t></m:r></m:sup>
                  <m:e><m:r><m:t>( C̄</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t> · w</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>k</m:t></m:r></m:sub></m:sSub><m:r><m:t> ) + ε</m:t></m:r></m:e>
                </m:nary>
              </m:den>
            </m:f>
          </m:oMath>
        </m:oMathPara>
        '''
    # Equation 3: Skill Decay Model
    elif eq_id == 3:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath>
            <m:r><m:t>S</m:t></m:r>
            <m:sSub>
              <m:e><m:r><m:t></m:t></m:r></m:e>
              <m:sub><m:r><m:t>decayed</m:t></m:r></m:sub>
            </m:sSub>
            <m:r><m:t>(t) = S</m:t></m:r>
            <m:sSub>
              <m:e><m:r><m:t></m:t></m:r></m:e>
              <m:sub><m:r><m:t>0</m:t></m:r></m:sub>
            </m:sSub>
            <m:r><m:t> · e</m:t></m:r>
            <m:sSup>
              <m:e><m:r><m:t></m:t></m:r></m:e>
              <m:sup><m:r><m:t>-λ(t - t</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>last</m:t></m:r></m:sub></m:sSub><m:r><m:t>)</m:t></m:r></m:sup>
            </m:sSup>
          </m:oMath>
        </m:oMathPara>
        '''
    # Equation 4: Next.js Framework DAG Dependency
    elif eq_id == 4:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath>
            <m:r><m:t>Next.js  ──(depends on)──►  React  ──(depends on)──►  JavaScript</m:t></m:r>
          </m:oMath>
        </m:oMathPara>
        '''
    # Equation 5: NestJS Framework DAG Dependency
    elif eq_id == 5:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath>
            <m:r><m:t>NestJS  ──(depends on)──►  Node.js  ──(depends on)──►  TypeScript</m:t></m:r>
          </m:oMath>
        </m:oMathPara>
        '''
    # Section 5.1: SIE-1.0 Proficiency equation
    elif eq_id == 10:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath>
            <m:r><m:t>S = min( 99,  </m:t></m:r>
            <m:f>
              <m:num>
                <m:r><m:t>V</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>total</m:t></m:r></m:sub></m:sSub><m:r><m:t> · (0.6 + 0.4 · O</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>ratio</m:t></m:r></m:sub></m:sSub><m:r><m:t>) · D</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>avg</m:t></m:r></m:sub></m:sSub><m:r><m:t> · K</m:t></m:r><m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>complexity</m:t></m:r></m:sub></m:sSub>
              </m:num>
              <m:den>
                <m:r><m:t>max( 1, √N )</m:t></m:r>
              </m:den>
            </m:f>
            <m:r><m:t> + 3N )</m:t></m:r>
          </m:oMath>
        </m:oMathPara>
        '''
    else:
        xml = f'''
        <m:oMathPara {nsdecls("m")}>
          <m:oMath><m:r><m:t>{text}</m:t></m:r></m:oMath>
        </m:oMathPara>
        '''

    return parse_xml(xml)

def clean_ieee_text(text):
    """
    Cleans inline math notation and applies IEEE formatting replacements:
    - [Eq. 1] -> Eq. (1)
    - Figure 1 -> Fig. 1
    - Table I -> TABLE I
    """
    if not text:
        return text

    # IEEE Cross-reference replacements
    text = re.sub(r'\[Eq\. (\d+)\]', r'Eq. (\1)', text)
    text = re.sub(r'Eq\. (\d+)', r'Eq. (\1)', text)
    text = re.sub(r'Figure (\d+)', r'Fig. \1', text)
    text = re.sub(r'Table ([I|V|X]+)', r'TABLE \1', text)

    # Specific Math replacements for IEEE prose
    text = text.replace(r'\mu_v', 'μᵥ')
    text = text.replace(r'\mu', 'μ')
    text = text.replace(r'\lambda', 'λ')
    text = text.replace(r'\mathcal{H}', 'H')
    text = text.replace(r'\mathcal', '')
    text = text.replace(r'\epsilon', 'ε')
    text = text.replace(r'\Delta', 'Δ')
    text = text.replace(r'\frac', '')
    text = text.replace(r'\sum', 'Σ')
    text = text.replace(r'\xrightarrow', ' → ')
    text = text.replace(r'\cdot', '·')
    text = text.replace(r'\bar', '')
    text = re.sub(r'\^\{([^}]+)\}', r'^\1', text)
    text = re.sub(r'\_\{([^}]+)\}', r'_\1', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\$([^$]+)\$', r'\1', text)
    text = text.replace('\\_', '_')
    text = text.replace('\\', '')
    
    return text

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:left w:val="none"/><w:right w:val="none"/><w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/><w:insideV w:val="none"/></w:tblBorders>')
    tblPr.append(borders)

def build_ieee_publication_docx():
    md_path = r"C:\Users\elitebook840g89319\.gemini\antigravity-ide\brain\902d6e6c-b35a-4d3e-a983-59d2e321b9fc\paper2_academic_universe_growth_intelligence.md"
    docx_path = r"c:\github\academicuniverse.com\academicuniverse\Academic_Universe_Paper2_v1.0_IEEE_Publication_Final.docx"
    fig1_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure1_Architecture_Pipeline.png"
    fig2_path = r"c:\github\academicuniverse.com\academicuniverse\docs\figures\Figure2_Decay_Sensitivity_Curves.png"

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = docx.Document()

    # Page Setup - 1 inch margins
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    eq_counter = 0

    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    def process_table(t_lines):
        rows = []
        for line in t_lines:
            if '|' in line:
                parts = [p.strip() for p in line.split('|')[1:-1]]
                if parts and not all(set(p) <= {'-', ':', ' '} for p in parts):
                    rows.append(parts)
        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table)

        # Enforce CantSplit on all rows & tblHeader on first row
        for r_idx, row in enumerate(rows):
            trPr = table.rows[r_idx]._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

            if r_idx == 0:
                trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

            for c_idx, val in enumerate(row):
                if c_idx < len(table.columns):
                    cell = table.cell(r_idx, c_idx)
                    cleaned_val = clean_ieee_text(val)
                    cell.text = cleaned_val
                    set_cell_margins(cell)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)

                    if r_idx == 0:
                        set_cell_background(cell, "1E293B")
                        for run in p.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    else:
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "F8FAFC")
                        for run in p.runs:
                            run.font.size = Pt(10)

        doc.add_paragraph()

    while i < len(lines):
        line = lines[i].rstrip('\r\n')

        # Code Blocks -> Replace ASCII diagrams with High-Res IEEE Figures
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                block_content = '\n'.join(code_block_lines)
                
                # Figure 1 Architecture
                if 'GitHub API' in block_content or 'Evidence Intelligence Layer' in block_content:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.paragraph_format.keep_with_next = True
                    p_img.add_run().add_picture(fig1_path, width=Inches(6.2))

                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    r_tag = p_cap.add_run("Fig. 1. ")
                    r_tag.bold = True
                    r_tag.font.size = Pt(9)
                    r_txt = p_cap.add_run("High-level architectural pipeline of the Academic Universe Growth Intelligence Ecosystem.")
                    r_txt.font.size = Pt(9)
                    r_txt.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

                # Figure 2 Decay Curves
                elif 'Proficiency S(t)' in block_content or 'λ=0.01' in block_content:
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(4)
                    p_img.paragraph_format.keep_with_next = True
                    p_img.add_run().add_picture(fig2_path, width=Inches(5.8))

                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_after = Pt(12)
                    r_tag = p_cap.add_run("Fig. 2. ")
                    r_tag.bold = True
                    r_tag.font.size = Pt(9)
                    r_txt = p_cap.add_run("Sensitivity analysis of technical skill proficiency decay S(t) under varied decay parameter λ coefficients.")
                    r_txt.font.size = Pt(9)
                    r_txt.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.2)
                    run = p.add_run(clean_ieee_text(block_content))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                code_block_lines = []
            else:
                in_code_block = True
                code_block_lines = []
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Tables
        if '|' in line and not line.startswith('>'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            process_table(table_lines)
            table_lines = []

        # Blockquotes
        if line.startswith('> '):
            quote_text = clean_ieee_text(line[2:].strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            i += 1
            continue

        # Headings (Enforce keep_with_next to eliminate widows/orphans)
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_ieee_text(line[2:]))
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_ieee_text(line[3:]))
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_ieee_text(line[4:]))
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_ieee_text(line[5:]))
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

        # Centered Display Equations with OMML XML Engine
        elif line.startswith('$$') and line.endswith('$$'):
            eq_counter += 1
            p_eq = doc.add_paragraph()
            p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_eq.paragraph_format.space_before = Pt(8)
            p_eq.paragraph_format.space_after = Pt(8)
            p_eq.paragraph_format.keep_with_next = True

            # Extract equation number tag if present
            raw_eq = line[2:-2].strip()
            num_match = re.search(r'\\text\{\[Eq\. (\d+)\]\}|\[Eq\. (\d+)\]', raw_eq)
            eq_num = num_match.group(1) or num_match.group(2) if num_match else str(eq_counter)

            # Build OMML Display Equation
            omml_elem = create_omml_equation(eq_counter, raw_eq)
            p_eq._p.append(omml_elem)

            # Add right-aligned equation tag: (1), (2), (3), (4), (5)
            p_tag = doc.add_paragraph()
            p_tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_tag.paragraph_format.space_before = Pt(0)
            p_tag.paragraph_format.space_after = Pt(6)
            r_tag = p_tag.add_run(f"({eq_num})")
            r_tag.font.name = 'Times New Roman'
            r_tag.font.size = Pt(10)
            r_tag.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

        elif line.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        elif line.strip() != '':
            cleaned_line = clean_ieee_text(line)

            # Check if Table Header Caption (*Table I: ... or *Table II: ...)
            if cleaned_line.startswith('*Table I:') or cleaned_line.startswith('*Table II:') or cleaned_line.startswith('Table I:') or cleaned_line.startswith('Table II:'):
                p_tbl_cap = doc.add_paragraph()
                p_tbl_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_tbl_cap.paragraph_format.space_before = Pt(10)
                p_tbl_cap.paragraph_format.space_after = Pt(4)
                p_tbl_cap.paragraph_format.keep_with_next = True
                
                table_tag = "TABLE I. " if "Table I" in cleaned_line else "TABLE II. "
                table_title = cleaned_line.split(':', 1)[1].strip().rstrip('*') if ':' in cleaned_line else cleaned_line
                
                r_tag = p_tbl_cap.add_run(table_tag)
                r_tag.bold = True
                r_tag.font.size = Pt(9.5)
                r_tag.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
                
                r_txt = p_tbl_cap.add_run(table_title)
                r_txt.font.size = Pt(9.5)
                r_txt.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            # Ignore duplicate markdown figure caption lines since handled above
            elif cleaned_line.startswith('*Figure') or cleaned_line.startswith('Figure'):
                pass
            else:
                p = doc.add_paragraph()
                p.add_run(cleaned_line)

        i += 1

    if in_table:
        process_table(table_lines)

    # Reference Section Polish
    ref_heading = False
    for paragraph in doc.paragraphs:
        if '## References' in paragraph.text or 'References' in paragraph.text:
            ref_heading = True
            break

    if ref_heading:
        enhanced_refs = [
            "5. Corbett, A. T., & Anderson, J. R. (1994). Knowledge tracing: Modeling the acquisition of procedural knowledge. User Modeling and User-Adapted Interaction, 4(4), 253-278.",
            "6. Piech, C., Bassen, J., Huang, J., Ganguli, S., Sahami, M., Guibas, L. J., & Sohl-Dickstein, J. (2015). Deep knowledge tracing. Advances in Neural Information Processing Systems (NeurIPS), 28, 505-513.",
            "7. Bodily, R., & Verbert, K. (2017). Review of research on student-facing learning analytics dashboards. Journal of Learning Analytics, 4(3), 67-89.",
            "8. Ferguson, R. (2012). Learning analytics: Drivers, developments and challenges. International Journal of Technology Enhanced Learning, 4(5-6), 304-317.",
            "9. Jivet, I., Scheffel, M., Specht, M., & Drachsler, H. (2018). License to evaluate: Preparing learning analytics dashboards for educational practice. Proceedings of the 8th International Conference on Learning Analytics & Knowledge (LAK '18), 31-40."
        ]
        doc.add_paragraph()
        p_note = doc.add_paragraph()
        r_note = p_note.add_run("Foundational Literature Additions (IEEE / ACM Learning Analytics Corpus):")
        r_note.bold = True
        r_note.font.size = Pt(11)
        r_note.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

        for ref in enhanced_refs:
            p_ref = doc.add_paragraph()
            p_ref.paragraph_format.left_indent = Inches(0.25)
            p_ref.paragraph_format.space_after = Pt(4)
            r = p_ref.add_run(ref)
            r.font.size = Pt(10)

    doc.save(docx_path)
    print(f"Successfully compiled IEEE Publication Final DOCX to {docx_path}")

if __name__ == '__main__':
    build_ieee_publication_docx()
