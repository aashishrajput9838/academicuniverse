import os
import docx
import re
import win32com.client
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

workspace = Path(__file__).resolve().parents[1]
v4_docx_path = workspace / "docs" / "paper" / "PaperV4_Final_Submission.docx"
test_docx_path = workspace / "docs" / "paper" / "PaperV5_Test_TOC_Fixed.docx"
test_pdf_path = workspace / "docs" / "paper" / "PaperV5_Test_TOC_Fixed.pdf"

print("=== TESTING TOC OUTLINE LEVEL FIX ===")

doc = docx.Document(v4_docx_path)

def set_paragraph_outline_level(p, level):
    pPr = p._p.get_or_add_pPr()
    # Remove existing outlineLvl if present
    existing = pPr.find(qn('w:outlineLvl'))
    if existing is not None:
        pPr.remove(existing)
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level))
    pPr.append(outlineLvl)

# 1. Assign Outline Levels to all section headings
h1_count = 0
h2_count = 0

for i, p in enumerate(doc.paragraphs):
    text_str = p.text.strip()
    if not text_str:
        continue
    
    # Level 1 (Outline Level 0)
    is_h1 = (
        text_str.startswith("1. ") or
        text_str.startswith("2. ") or
        text_str.startswith("3. ") or
        text_str.startswith("4. ") or
        text_str.startswith("5. ") or
        text_str.startswith("6. ") or
        text_str.startswith("7. ") or
        text_str.startswith("8. ") or
        text_str.startswith("9. ") or
        text_str == "REFERENCES" or
        text_str.startswith("APPENDIX A") or
        text_str.startswith("APPENDIX B") or
        text_str.startswith("APPENDIX C") or
        text_str.startswith("Ethics & Privacy Statement") or
        text_str.startswith("ACKNOWLEDGMENT")
    )
    
    # Level 2 (Outline Level 1)
    is_h2 = bool(
        re.match(r"^\d\.\d\s+", text_str) or
        text_str.startswith("A.1 ") or text_str.startswith("A.2 ") or
        text_str.startswith("B.1 ") or text_str.startswith("B.2 ") or
        text_str.startswith("C.1 ") or text_str.startswith("C.2 ") or text_str.startswith("C.3 ")
    )
    
    if is_h1:
        set_paragraph_outline_level(p, 0)
        h1_count += 1
    elif is_h2:
        set_paragraph_outline_level(p, 1)
        h2_count += 1

print(f"Outline Level Assignment Summary:")
print(f"  - Heading 1 (Outline Level 0): {h1_count}")
print(f"  - Heading 2 (Outline Level 1): {h2_count}")

# 2. Insert CONTENTS heading and TOC field
target_p = None
for i, p in enumerate(doc.paragraphs[:15]):
    if p.text.strip().startswith("Index Terms"):
        target_p = p
        break

assert target_p is not None, "Error: Index Terms paragraph not found!"

p_contents_hdr = target_p.insert_paragraph_before("CONTENTS")
p_contents_hdr.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
p_contents_hdr.paragraph_format.space_before = docx.shared.Pt(14)
p_contents_hdr.paragraph_format.space_after = docx.shared.Pt(6)
if p_contents_hdr.runs:
    r = p_contents_hdr.runs[0]
    r.font.name = "Times New Roman"
    r.font.size = docx.shared.Pt(13)
    r.bold = True

p_toc = target_p.insert_paragraph_before()
p_toc.paragraph_format.space_after = docx.shared.Pt(12)
pPr = p_toc._p.get_or_add_pPr()
pStyle = OxmlElement('w:pStyle')
pStyle.set(qn('w:val'), 'TOCHeading')
pPr.append(pStyle)

r_toc = p_toc.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

r_toc._r.append(fldChar1)
r_toc._r.append(instrText)
r_toc._r.append(fldChar2)
r_toc._r.append(fldChar3)

doc.save(test_docx_path)
print(f"[SUCCESS] Saved test docx with outline levels to {test_docx_path.name}")

# Test Word COM TOC update & PDF export
abs_docx = os.path.abspath(str(test_docx_path))
abs_pdf = os.path.abspath(str(test_pdf_path))

word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    doc_word = word.Documents.Open(abs_docx)
    
    # Update fields to populate TOC
    doc_word.Fields.Update()
    for s in doc_word.Sections:
        s.Range.Fields.Update()
    
    doc_word.SaveAs(abs_pdf, FileFormat=17) # 17 = wdFormatPDF
    page_count = doc_word.ComputeStatistics(2)
    doc_word.Close()
    print(f"[SUCCESS] Exported PDF with populated TOC: {test_pdf_path.name} ({page_count} Pages!)")
except Exception as e:
    print(f"Word COM TOC Export Error: {e}")
finally:
    if word:
        try: word.Quit()
        except: pass
